"""
Version Comparison Service

Compares two test plan versions and generates Was/Is diffs with track changes.
Uses difflib for text comparison and python-docx for DOCX export.
"""

import json
import difflib
from typing import List, Dict, Any, Tuple
from sqlalchemy.orm import Session
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_COLOR_INDEX
from io import BytesIO

from repositories.versioning_repository import TestPlanVersionRepository, TestPlanRepository
from integrations.chromadb_client import get_chroma_client


class VersionComparisonService:
    """Service for comparing test plan versions"""

    def __init__(self, db: Session):
        self.db = db
        self.version_repo = TestPlanVersionRepository(db)
        self.plan_repo = TestPlanRepository(db)

    def compare_versions(
        self, plan_id: int, was_version_id: int, is_version_id: int
    ) -> Dict[str, Any]:
        """
        Compare two versions and return structured diff.

        Args:
            plan_id: Test plan ID
            was_version_id: Previous version ID
            is_version_id: Current version ID

        Returns:
            CompareVersionsResponse dict
        """
        # Fetch test plan to get collection name
        test_plan = self.plan_repo.get(plan_id)
        if not test_plan:
            raise ValueError(f"Test plan {plan_id} not found")

        collection_name = test_plan.collection_name or "generated_test_plan"

        # Fetch versions
        was_version = self.version_repo.get(was_version_id)
        is_version = self.version_repo.get(is_version_id)

        if not was_version or was_version.plan_id != plan_id:
            raise ValueError(f"Was version {was_version_id} not found or doesn't belong to plan {plan_id}")
        if not is_version or is_version.plan_id != plan_id:
            raise ValueError(f"Is version {is_version_id} not found or doesn't belong to plan {plan_id}")

        # Load JSON content from ChromaDB
        was_content = self._load_version_content(was_version.document_id, collection_name)
        is_content = self._load_version_content(is_version.document_id, collection_name)

        # Generate diffs
        differences = self._generate_diffs(was_content, is_content)

        # Generate HTML preview
        html_preview = self._generate_html_preview(differences, was_content, is_content)

        # Convert ORM objects to dicts
        was_version_dict = {
            "id": was_version.id,
            "plan_id": was_version.plan_id,
            "version_number": was_version.version_number,
            "document_id": was_version.document_id,
            "status": was_version.status.value if hasattr(was_version.status, 'value') else str(was_version.status),
            "created_at": was_version.created_at.isoformat() if was_version.created_at else None,
            "updated_at": was_version.updated_at.isoformat() if hasattr(was_version, 'updated_at') and was_version.updated_at else None
        }

        is_version_dict = {
            "id": is_version.id,
            "plan_id": is_version.plan_id,
            "version_number": is_version.version_number,
            "document_id": is_version.document_id,
            "status": is_version.status.value if hasattr(is_version.status, 'value') else str(is_version.status),
            "created_at": is_version.created_at.isoformat() if is_version.created_at else None,
            "updated_at": is_version.updated_at.isoformat() if hasattr(is_version, 'updated_at') and is_version.updated_at else None
        }

        return {
            "was_version": was_version_dict,
            "is_version": is_version_dict,
            "differences": differences,
            "total_changes": len(differences),
            "html_preview": html_preview
        }

    def _load_version_content(self, document_id: str, collection_name: str) -> Dict[str, Any]:
        """Load JSON content from ChromaDB, trying multiple collections"""
        # Try multiple collections - drafts might be in test_plan_drafts
        collections_to_try = [collection_name, "test_plan_drafts", "generated_test_plan"]
        # Remove duplicates while preserving order
        seen = set()
        collections_to_try = [c for c in collections_to_try if not (c in seen or seen.add(c))]

        chroma_client = get_chroma_client()
        last_error = None

        for coll_name in collections_to_try:
            try:
                collection = chroma_client.get_collection(name=coll_name)

                # Get document by ID
                result = collection.get(ids=[document_id], include=["documents", "metadatas"])

                if result and result.get("documents") and len(result["documents"]) > 0:
                    # Parse the document content (first document in results)
                    doc_content = result["documents"][0]
                    if isinstance(doc_content, str):
                        try:
                            return json.loads(doc_content)
                        except json.JSONDecodeError:
                            # Not JSON, return as-is wrapped in a dict
                            return {"raw_content": doc_content}
                    return doc_content
            except Exception as e:
                last_error = e
                continue

        raise ValueError(f"Document {document_id} not found in any collection. Last error: {str(last_error)}")

    def _generate_diffs(
        self, was_content: Dict, is_content: Dict
    ) -> List[Dict[str, Any]]:
        """Generate structured diffs between versions"""
        differences = []

        # Check for raw content (non-JSON format)
        if "raw_content" in was_content or "raw_content" in is_content:
            # Fall back to simple text comparison
            was_text = was_content.get("raw_content", str(was_content))
            is_text = is_content.get("raw_content", str(is_content))

            if was_text != is_text:
                differences.append({
                    "section_id": "full_document",
                    "section_title": "Full Document",
                    "field": "content",
                    "change_type": "modified",
                    "old_value": was_text[:2000] if was_text else "",
                    "new_value": is_text[:2000] if is_text else ""
                })
            return differences

        # Extract test plan sections
        was_test_plan = was_content.get("test_plan", {})
        is_test_plan = is_content.get("test_plan", {})

        was_sections = {s.get("section_id", f"section_{i}"): s for i, s in enumerate(was_test_plan.get("sections", []))}
        is_sections = {s.get("section_id", f"section_{i}"): s for i, s in enumerate(is_test_plan.get("sections", []))}

        # Find deleted sections
        for section_id in was_sections:
            if section_id not in is_sections:
                differences.append({
                    "section_id": section_id,
                    "section_title": was_sections[section_id].get("section_title", "Untitled"),
                    "field": "section",
                    "change_type": "deleted",
                    "old_value": json.dumps(was_sections[section_id], indent=2)[:500],
                    "new_value": None
                })

        # Find added and modified sections
        for section_id, is_section in is_sections.items():
            if section_id not in was_sections:
                # New section
                differences.append({
                    "section_id": section_id,
                    "section_title": is_section.get("section_title", "Untitled"),
                    "field": "section",
                    "change_type": "added",
                    "old_value": None,
                    "new_value": json.dumps(is_section, indent=2)[:500]
                })
            else:
                # Compare section fields
                was_section = was_sections[section_id]
                section_diffs = self._compare_sections(section_id, was_section, is_section)
                differences.extend(section_diffs)

        return differences

    def _compare_sections(
        self, section_id: str, was: Dict, is_: Dict
    ) -> List[Dict[str, Any]]:
        """Compare individual section fields"""
        diffs = []

        # Compare text fields
        text_fields = ["section_title", "synthesized_rules"]
        for field in text_fields:
            was_val = was.get(field, "")
            is_val = is_.get(field, "")
            if was_val != is_val:
                diffs.append({
                    "section_id": section_id,
                    "section_title": is_.get("section_title", "Untitled"),
                    "field": field,
                    "change_type": "modified",
                    "old_value": str(was_val)[:500] if was_val else "",
                    "new_value": str(is_val)[:500] if is_val else ""
                })

        # Compare test procedures (complex comparison)
        was_procs = {p.get("id", f"proc_{i}"): p for i, p in enumerate(was.get("test_procedures", []))}
        is_procs = {p.get("id", f"proc_{i}"): p for i, p in enumerate(is_.get("test_procedures", []))}

        # Deleted procedures
        for proc_id in was_procs:
            if proc_id not in is_procs:
                diffs.append({
                    "section_id": section_id,
                    "section_title": is_.get("section_title", "Untitled"),
                    "field": f"test_procedure.{proc_id}",
                    "change_type": "deleted",
                    "old_value": json.dumps(was_procs[proc_id], indent=2)[:500],
                    "new_value": None
                })

        # Added/modified procedures
        for proc_id, is_proc in is_procs.items():
            if proc_id not in was_procs:
                diffs.append({
                    "section_id": section_id,
                    "section_title": is_.get("section_title", "Untitled"),
                    "field": f"test_procedure.{proc_id}",
                    "change_type": "added",
                    "old_value": None,
                    "new_value": json.dumps(is_proc, indent=2)[:500]
                })
            elif was_procs[proc_id] != is_proc:
                diffs.append({
                    "section_id": section_id,
                    "section_title": is_.get("section_title", "Untitled"),
                    "field": f"test_procedure.{proc_id}",
                    "change_type": "modified",
                    "old_value": json.dumps(was_procs[proc_id], indent=2)[:500],
                    "new_value": json.dumps(is_proc, indent=2)[:500]
                })

        return diffs

    def _generate_html_preview(
        self, differences: List[Dict], was_content: Dict, is_content: Dict
    ) -> str:
        """Generate HTML preview with track changes styling"""
        html = """
        <style>
            .deleted { color: red; text-decoration: line-through; }
            .added { color: green; text-decoration: underline; }
            .section { margin: 20px 0; border-left: 3px solid #ccc; padding-left: 15px; }
            .section-title { font-weight: bold; font-size: 1.2em; margin-bottom: 10px; }
            .field-name { font-weight: bold; color: #333; }
            .change-badge {
                display: inline-block;
                padding: 2px 8px;
                border-radius: 3px;
                font-size: 0.85em;
                margin-right: 8px;
            }
            .badge-added { background-color: #d4edda; color: #155724; }
            .badge-deleted { background-color: #f8d7da; color: #721c24; }
            .badge-modified { background-color: #fff3cd; color: #856404; }
        </style>
        <h2>Track Changes Preview</h2>
        """

        if not differences:
            html += "<p><em>No changes detected between versions.</em></p>"
            return html

        # Group by section
        sections_dict = {}
        for diff in differences:
            section_title = diff["section_title"]
            if section_title not in sections_dict:
                sections_dict[section_title] = []
            sections_dict[section_title].append(diff)

        # Render each section
        for section_title, section_diffs in sections_dict.items():
            html += f'<div class="section">'
            html += f'<div class="section-title">{section_title}</div>'

            for diff in section_diffs:
                change_type = diff["change_type"]
                field = diff["field"]

                # Change badge
                badge_class = f"badge-{change_type}"
                html += f'<div style="margin: 10px 0;">'
                html += f'<span class="change-badge {badge_class}">{change_type.upper()}</span>'
                html += f'<span class="field-name">{field}</span>'
                html += '<br/>'

                if change_type == "deleted":
                    old_val = diff["old_value"] or ""
                    preview = old_val[:200] + "..." if len(old_val) > 200 else old_val
                    html += f'<p class="deleted">{self._escape_html(preview)}</p>'
                elif change_type == "added":
                    new_val = diff["new_value"] or ""
                    preview = new_val[:200] + "..." if len(new_val) > 200 else new_val
                    html += f'<p class="added">{self._escape_html(preview)}</p>'
                elif change_type == "modified":
                    old_val = diff["old_value"] or ""
                    new_val = diff["new_value"] or ""
                    # Use word-level diff for better visualization
                    old_words = old_val.split()[:50]  # Limit words for preview
                    new_words = new_val.split()[:50]
                    diff_html = self._word_diff_html(old_words, new_words)
                    html += f'<p>{diff_html}</p>'

                html += '</div>'

            html += '</div>'

        return html

    def _word_diff_html(self, old_words: List[str], new_words: List[str]) -> str:
        """Generate word-level diff HTML"""
        matcher = difflib.SequenceMatcher(None, old_words, new_words)
        result = []

        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            if tag == 'equal':
                result.append(self._escape_html(' '.join(old_words[i1:i2])))
            elif tag == 'delete':
                result.append(f'<span class="deleted">{self._escape_html(" ".join(old_words[i1:i2]))}</span>')
            elif tag == 'insert':
                result.append(f'<span class="added">{self._escape_html(" ".join(new_words[j1:j2]))}</span>')
            elif tag == 'replace':
                result.append(f'<span class="deleted">{self._escape_html(" ".join(old_words[i1:i2]))}</span>')
                result.append(f'<span class="added">{self._escape_html(" ".join(new_words[j1:j2]))}</span>')

        return ' '.join(result)

    def _escape_html(self, text: str) -> str:
        """Escape HTML special characters"""
        return (
            text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace('"', "&quot;")
            .replace("'", "&#39;")
        )

    def export_comparison_docx(
        self, plan_id: int, was_version_id: int, is_version_id: int
    ) -> bytes:
        """Export comparison as DOCX with track changes formatting"""

        # Get comparison data
        comparison = self.compare_versions(plan_id, was_version_id, is_version_id)
        differences = comparison["differences"]
        was_version = comparison["was_version"]
        is_version = comparison["is_version"]

        # Create DOCX
        doc = Document()

        # Title
        title = doc.add_heading('Version Comparison: Was vs Is', 0)

        # Metadata section
        doc.add_paragraph(f"Previous Version: v{was_version['version_number']} ({was_version['status']})")
        doc.add_paragraph(f"Current Version: v{is_version['version_number']} ({is_version['status']})")
        doc.add_paragraph(f"Total Changes: {len(differences)}")
        doc.add_paragraph("")

        if not differences:
            doc.add_paragraph("No changes detected between versions.")
        else:
            # Group changes by section
            sections_dict = {}
            for diff in differences:
                section_title = diff["section_title"]
                if section_title not in sections_dict:
                    sections_dict[section_title] = []
                sections_dict[section_title].append(diff)

            # Add changes by section
            for section_title, section_diffs in sections_dict.items():
                doc.add_heading(section_title, level=1)

                for diff in section_diffs:
                    # Add field name
                    para = doc.add_paragraph()
                    run = para.add_run(f"{diff['field']}: ")
                    run.bold = True

                    change_type = diff["change_type"]

                    if change_type == "deleted":
                        # Red strikethrough for deletions
                        old_val = diff.get("old_value", "")
                        if old_val:
                            run = para.add_run(old_val[:1000])  # Limit length
                            run.font.color.rgb = RGBColor(255, 0, 0)  # Red
                            run.font.strike = True

                    elif change_type == "added":
                        # Green underline for additions
                        new_val = diff.get("new_value", "")
                        if new_val:
                            run = para.add_run(new_val[:1000])
                            run.font.color.rgb = RGBColor(0, 128, 0)  # Green
                            run.font.underline = True

                    elif change_type == "modified":
                        # Show old (strikethrough) then new (underline)
                        old_val = diff.get("old_value", "")
                        new_val = diff.get("new_value", "")

                        if old_val:
                            old_run = para.add_run(old_val[:500])
                            old_run.font.color.rgb = RGBColor(255, 0, 0)
                            old_run.font.strike = True

                        para.add_run(" → ")

                        if new_val:
                            new_run = para.add_run(new_val[:500])
                            new_run.font.color.rgb = RGBColor(0, 128, 0)
                            new_run.font.underline = True

                    # Add spacing
                    doc.add_paragraph("")

        # Convert to bytes
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        return doc_io.getvalue()
