# services/document_service.py
import io
import uuid
import base64
import requests
from typing import List, Optional, Union, Dict, Any
from docx import Document
import fitz
import os
from bs4 import BeautifulSoup
from services.rag_service import RAGService
from services.agent_service import AgentService
from services.llm_service import LLMService
from services.multi_agent_test_plan_service import MultiAgentTestPlanService
import re
import json
import time
import redis
import datetime
try:
    import tiktoken
    TIKTOKEN_AVAILABLE = True
except ImportError:
    TIKTOKEN_AVAILABLE = False

class TemplateParser:
    @staticmethod
    def extract_headings_from_docx(path: str) -> List[str]:
        doc = Document(path)
        return [p.text for p in doc.paragraphs
                if p.style.name.startswith("Heading") and p.text.strip()]

    @staticmethod
    def extract_headings_from_pdf(path: str, size_threshold: float = 16.0) -> List[str]:
        doc = fitz.open(path)
        headings = []
        for page in doc:
            for block in page.get_text("dict")["blocks"]:
                if block["type"] != 0: continue
                for line in block["lines"]:
                    size = line["spans"][0]["size"]
                    text = "".join(s["text"] for s in line["spans"]).strip()
                    if size >= size_threshold and text:
                        headings.append(text)
        return headings

    @staticmethod
    def extract_headings_from_html(html: str) -> List[str]:
        soup = BeautifulSoup(html, "html.parser")
        headings = []
        for level in range(1,7):
            for tag in soup.find_all(f"h{level}"):
                text = tag.get_text(strip=True)
                if text:
                    headings.append(text)
        return headings

    @staticmethod
    def extract_headings_from_markdown(md: str) -> List[str]:
        lines = md.splitlines()
        return [l.lstrip("# ").strip() for l in lines
                if l.startswith("#")]
        
    @staticmethod
    def extract_headings_from_text(text: str) -> List[str]:
        return [
            line.lstrip("# ").strip()
            for line in text.splitlines()
            if line.startswith("# ")
        ]

    @classmethod
    def extract(cls, path_or_str: str, is_path: bool=True) -> List[str]:
        if not is_path:
            return cls.extract_headings_from_text(path_or_str)

        # ext = os.path.splitext(path_or_str if is_path else "")[1].lower()
        ext = os.path.splitext(path_or_str)[1].lower()
        
        if ext == ".docx":
            return cls.extract_headings_from_docx(path_or_str)
        if ext == ".docx":
            return cls.extract_headings_from_docx(path_or_str)
        if ext == ".pdf":
            return cls.extract_headings_from_pdf(path_or_str)
        if ext in {".html", ".htm"}:
            html = open(path_or_str).read() if is_path else path_or_str
            return cls.extract_headings_from_html(html)
        if ext == ".md" or (not is_path and "\n" in path_or_str):
            md = open(path_or_str).read() if is_path else path_or_str
            return cls.extract_headings_from_markdown(md)
        raise ValueError(f"Unsupported template format: {ext!r}")
class DocumentService:
    def __init__(
        self,
        rag_service: RAGService,
        agent_service: AgentService,
        llm_service: LLMService,
        chroma_url: str,
        agent_api_url: str,
    ):
        self.rag = rag_service
        self.agent = agent_service
        self.llm = llm_service
        self.chroma_url = chroma_url.rstrip("/")
        self.agent_api = agent_api_url.rstrip("/")
        # Initialize multi-agent test plan service
        self.multi_agent_test_plan_service = MultiAgentTestPlanService(llm_service, chroma_url, agent_api_url)
        # Token encoder for counting tokens
        if TIKTOKEN_AVAILABLE:
            try:
                self._tokenizer = tiktoken.encoding_for_model("gpt-4")
            except:
                self._tokenizer = tiktoken.get_encoding("cl100k_base")
        else:
            self._tokenizer = None
        # Lightweight in-memory cache for reconstructed documents to avoid repeated fetches
        # Key: (collection_name, document_id) -> { 'content': str, 'metadata': dict }
        self._reconstructed_doc_cache: Dict[tuple, Dict[str, Any]] = {}
        self._reconstructed_doc_cache_order: List[tuple] = []  # simple LRU order
        self._reconstructed_doc_cache_max = 50

    def _cache_get_reconstructed(self, collection: str, doc_id: str) -> Optional[Dict[str, Any]]:
        key = (collection, doc_id)
        item = self._reconstructed_doc_cache.get(key)
        if item is not None:
            # move to end (MRU)
            try:
                self._reconstructed_doc_cache_order.remove(key)
            except ValueError:
                pass
            self._reconstructed_doc_cache_order.append(key)
        return item

    def _cache_set_reconstructed(self, collection: str, doc_id: str, value: Dict[str, Any]) -> None:
        key = (collection, doc_id)
        self._reconstructed_doc_cache[key] = value
        # maintain order
        try:
            self._reconstructed_doc_cache_order.remove(key)
        except ValueError:
            pass
        self._reconstructed_doc_cache_order.append(key)
        # evict LRU if over capacity
        while len(self._reconstructed_doc_cache_order) > self._reconstructed_doc_cache_max:
            lru_key = self._reconstructed_doc_cache_order.pop(0)
            self._reconstructed_doc_cache.pop(lru_key, None)

    def count_tokens(self, text: str) -> int:
        """Count tokens in text using tiktoken"""
        if TIKTOKEN_AVAILABLE and self._tokenizer:
            try:
                return len(self._tokenizer.encode(text))
            except:
                pass
        # Fallback: roughly estimate 4 chars per token
        return len(text) // 4

    def _extract_document_sections(self, source_collections: List[str], source_doc_ids: List[str], 
                                 use_rag: bool, top_k: int,
                                 sectioning_strategy: str = "auto",
                                 chunks_per_section: int = 5) -> Dict[str, str]:
        """Extract organized sections using stored chunk metadata where possible.

        Strategies:
        - by_metadata: group by metadata.section_title, else page_number
        - by_chunks: group sequentially in windows of N chunk indices
        - auto: prefer metadata grouping; fallback to by_chunks; then natural/size splits
        """
        sections = {}
        
        for coll in (source_collections or []):
            try:
                # Fast path: if explicit document IDs are provided, reconstruct each document
                if source_doc_ids:
                    print(f"DEBUG: Using reconstruct fast path for collection '{coll}' and {len(source_doc_ids)} document(s)")
                    for doc_id in source_doc_ids:
                        try:
                            cached = self._cache_get_reconstructed(coll, doc_id)
                            if cached is None:
                                r = requests.get(
                                    f"{self.chroma_url}/documents/reconstruct/{doc_id}",
                                    params={"collection_name": coll},
                                    timeout=120,
                                )
                                if not r.ok:
                                    print(f"ERROR: Reconstruct failed for doc_id '{doc_id}': {r.status_code} - {r.text}")
                                    continue
                                rec = r.json()
                                cached = {
                                    "content": rec.get("reconstructed_content", ""),
                                    "document_name": rec.get("document_name", doc_id),
                                    "metadata": rec.get("metadata", {}),
                                }
                                self._cache_set_reconstructed(coll, doc_id, cached)
                            doc_name = cached.get("document_name") or doc_id
                            full_document = cached.get("content") or ""
                            if not full_document:
                                print(f"DEBUG: Empty reconstructed content for '{doc_id}'")
                                continue
                            sections_before = len(sections)
                            self._create_document_sections(doc_name, full_document, sections)
                            print(f"DEBUG: Reconstructed '{doc_name}' added {len(sections) - sections_before} sections")
                        except Exception as e:
                            print(f"ERROR: Exception reconstructing doc '{doc_id}': {e}")
                    # Done with this collection
                    continue

                # Get ALL documents from collection - no limits
                resp = requests.get(f"{self.chroma_url}/documents", 
                                  params={"collection_name": coll}, timeout=120)
                if not resp.ok:
                    print(f"ERROR: Failed to get documents from collection '{coll}': {resp.status_code} - {resp.text}")
                    continue
                    
                data = resp.json()
                docs = data.get("documents", [])
                metas = data.get("metadatas", [])
                ids = data.get("ids", [])
                
                print(f"DEBUG: Collection '{coll}' has {len(docs)} total documents")
                
                # Group by document_name
                document_pages = {}
                document_chunks = {}
                metadata_grouped_sections = {}

                for i, (doc_id, doc, meta) in enumerate(zip(ids, docs, metas)):
                    doc_name = meta.get("document_name", meta.get("document_id", "unknown"))
                    # Handle ChromaDB metadata where -1 was used instead of None
                    page_num = meta.get("page", meta.get("page_number", -1))
                    page_num = None if page_num == -1 else page_num
                    chunk_index = meta.get("chunk_index", None)
                    section_title = meta.get("section_title", "")
                    section_type = meta.get("section_type", "chunk")
                    
                    # Filter by source_doc_ids if specified - check all possible identifiers
                    if source_doc_ids:
                        doc_identifier = meta.get("document_id", doc_name)
                        base_name = doc_name.split('.')[0] if '.' in doc_name else doc_name
                        
                        found_match = False
                        for source_id in source_doc_ids:
                            if (source_id in doc_identifier or source_id in doc_name or 
                                source_id in base_name or doc_name.startswith(source_id)):
                                found_match = True
                                break
                        
                        if not found_match:
                            continue
                    
                    print(f"DEBUG: Processing {doc_name} - Page: {page_num}, Section: '{section_title}', Type: {section_type}")
                    
                    # Build metadata-based sections (preferred if available)
                    if section_title:
                        key = f"{doc_name} - {section_title}"
                        metadata_grouped_sections.setdefault(key, [])
                        metadata_grouped_sections[key].append(doc)
                    elif page_num is not None or section_type == "page":
                        # Page-level data (like from notebook approach)
                        if doc_name not in document_pages:
                            document_pages[doc_name] = {}
                        page_key = page_num if page_num is not None else chunk_index
                        document_pages[doc_name][page_key] = {
                            "content": doc,
                            "section_title": section_title,
                            "section_type": section_type
                        }
                    elif section_type == "logical_section" and section_title:
                        # Structure-preserving sections - use directly
                        section_key = f"{doc_name} - {section_title}"
                        sections[section_key] = doc.strip()
                        print(f"DEBUG: Added logical section: {section_key}")
                    elif chunk_index is not None:
                        # Chunk-based data (from original ChromaDB implementation)
                        if doc_name not in document_chunks:
                            document_chunks[doc_name] = {}
                        document_chunks[doc_name][chunk_index] = doc
                    else:
                        # Single document
                        sections[f"Document: {doc_name}"] = doc.strip()

                # FIXED: Process ALL content types to ensure comprehensive coverage
                def process_metadata_sections():
                    """Process metadata-based sections"""
                    if metadata_grouped_sections and sectioning_strategy in ("auto", "by_metadata"):
                        for key, parts in metadata_grouped_sections.items():
                            sections[key] = "\n\n".join(part.strip() for part in parts if part)
                        return len(sections)
                    return 0

                def process_chunk_sections():
                    """Process chunk-based sections"""
                    if document_chunks and sectioning_strategy in ("auto", "by_chunks"):
                        for doc_name, chunks in document_chunks.items():
                            for chunk_idx, content in sorted(chunks.items()):
                                group = (chunk_idx // max(1, chunks_per_section)) + 1
                                key = f"{doc_name} - Chunk Group {group}"
                                sections.setdefault(key, [])
                                sections[key].append(content)
                        # join lists to strings
                        for key, vals in list(sections.items()):
                            if isinstance(vals, list):
                                sections[key] = "\n\n".join(v.strip() for v in vals if v)
                        return len([k for k in sections.keys() if "Chunk Group" in k])
                    return 0

                # Process metadata sections first
                metadata_sections_added = process_metadata_sections()
                
                # Process chunk sections 
                chunk_sections_added = process_chunk_sections()
                
                print(f"DEBUG: Added {metadata_sections_added} metadata sections, {chunk_sections_added} chunk sections")
                
                # Process pages and chunks separately - prioritize pages over chunks to avoid duplication
                page_sections_added = 0
                processed_documents = set()
                
                # First, process documents with page data
                if document_pages:
                    print(f"DEBUG: Processing {len(document_pages)} documents with page data")
                    for doc_name, pages in document_pages.items():
                        sorted_pages = sorted(pages.items())
                        full_document = "\n\n".join([page_data["content"] if isinstance(page_data, dict) else page_data 
                                                    for page_num, page_data in sorted_pages])
                        print(f"DEBUG: Reconstructed '{doc_name}' from {len(sorted_pages)} pages, total length: {len(full_document)}")
                        print(f"DEBUG: First 200 chars of full_document: {full_document[:200]}...")
                        print(f"DEBUG: Last 200 chars of full_document: ...{full_document[-200:]}")
                        sections_before = len(sections)
                        self._create_document_sections(doc_name, full_document, sections)
                        page_sections_added += len(sections) - sections_before
                        processed_documents.add(doc_name)
                
                # Then, process documents with chunk data (but skip those already processed as pages)
                if document_chunks:
                    print(f"DEBUG: Processing {len(document_chunks)} documents with chunk data")
                    for doc_name, chunks in document_chunks.items():
                        if doc_name in processed_documents:
                            print(f"DEBUG: Skipping '{doc_name}' chunks - already processed as pages")
                            continue
                        sorted_chunks = sorted(chunks.items())
                        full_document = "\n\n".join([content for chunk_idx, content in sorted_chunks])
                        print(f"DEBUG: Reconstructed '{doc_name}' from {len(sorted_chunks)} chunks, total length: {len(full_document)}")
                        print(f"DEBUG: First chunk content preview: {sorted_chunks[0][1][:200] if sorted_chunks else 'No chunks'}...")
                        sections_before = len(sections)
                        self._create_document_sections(doc_name, full_document, sections)
                        page_sections_added += len(sections) - sections_before
                
                print(f"DEBUG: TOTAL sections created - Metadata: {metadata_sections_added}, Chunks: {chunk_sections_added}, Pages: {page_sections_added}")
                print(f"DEBUG: Final total sections: {len(sections)}")
                        
            except Exception as e:
                print(f"Error extracting sections from collection {coll}: {e}")
                
        print(f"DEBUG: Final extraction result: {len(sections)} sections created")
        return sections

    def _create_document_sections(self, doc_name: str, full_document: str, sections: Dict[str, str]):
        """Create logical sections from a full document while preserving structure.

        - Prefer natural headers (numbered headings, ALL-CAPS, APPENDIX blocks)
        - If too few sections, fall back to overlapping size-based splits
        - Keys include original document name for traceability
        """
        # First, try to identify natural document sections based on headers
        natural_sections = self._extract_natural_sections(full_document)

        # Ensure APPENDIX blocks are captured distinctly
        import re
        appendix_pattern = re.compile(r"^APPENDIX\s+[A-Z](?:\s*[-–]\s*.*)?$", re.MULTILINE)
        appendix_matches = list(appendix_pattern.finditer(full_document))
        if appendix_matches:
            for idx, m in enumerate(appendix_matches):
                start = m.start()
                end = appendix_matches[idx+1].start() if idx+1 < len(appendix_matches) else len(full_document)
                title = m.group(0).strip()
                body = full_document[start:end].strip()
                # Only add if not already present from natural extraction
                if title not in natural_sections:
                    natural_sections[title] = body

        if len(natural_sections) > 1:
            print(f"DEBUG: Found {len(natural_sections)} natural/appendix sections in {doc_name}")
            # Ensure each section has substantial content for test plan generation
            for section_title, section_content in natural_sections.items():
                section_key = f"{doc_name} - {section_title}"
                content = (section_content or "").strip()
                
                # Split large sections into smaller testable units
                if len(content) > 8000:  # If section is very large, break it down
                    subsections = self._split_large_section_for_testing(content, section_title)
                    for subsection_title, subsection_content in subsections.items():
                        sections[f"{section_key} - {subsection_title}"] = subsection_content
                        print(f"DEBUG: Created large subsection '{subsection_title}': {len(subsection_content)} chars")
                else:
                    sections[section_key] = content
                    print(f"DEBUG: Created section '{section_key}': {len(content)} chars, preview: {content[:100]}...")
        else:
            # Fall back to comprehensive page-based splitting if no natural sections found
            print(f"DEBUG: No natural sections found in {doc_name}, using comprehensive splitting")
            if len(full_document) > 6000:  # Increased for more comprehensive content
                section_parts = self._split_into_sections(full_document, section_size=6000)
                for part_idx, section_text in enumerate(section_parts):
                    section_key = f"{doc_name} - Part {part_idx + 1}"
                    sections[section_key] = (section_text or "").strip()
            else:
                sections[f"{doc_name} - Complete Document"] = (full_document or "").strip()

    def _extract_natural_sections(self, document_text: str) -> Dict[str, str]:
        """Extract sections based on document structure (headers, numbered sections, etc.)"""
        sections = {}
        lines = document_text.split('\n')
        current_section_title = "Introduction"
        current_section_content = []
        
        for line in lines:
            line_clean = line.strip()
            
            # Check for section headers - various patterns
            is_section_header = False
            header_title = ""
            
            if line_clean:
                # Pattern 1: Numbered sections (1., 1.1, 1.1.1, 4.2.1, etc.) - MUST be true section headers
                if re.match(r'^\d+(\.\d+)*\.?\s+[A-Z][A-Z\s]+', line_clean):
                    is_section_header = True
                    header_title = line_clean
                # Pattern 2: ALL CAPS headers (but not full sentences - must be concise headers)
                elif (line_clean.isupper() and 
                      len(line_clean.split()) <= 8 and 
                      len(line_clean) > 5 and
                      not line_clean.endswith('.') and  # Not a sentence
                      not line_clean.startswith(('THE ', 'THIS ', 'THESE '))):  # Not a sentence start
                    is_section_header = True
                    header_title = line_clean
                # Pattern 3: APPENDIX headers
                elif line_clean.startswith(('APPENDIX', 'CHAPTER', 'SECTION', 'PART')):
                    is_section_header = True
                    header_title = line_clean
                # Pattern 4: Headers with specific keywords (must be short headers, not sentences)
                elif (any(keyword in line_clean.upper() for keyword in ['REQUIREMENTS', 'SPECIFICATIONS', 'PROCEDURES', 'TESTING', 'CONFIGURATION']) and
                      len(line_clean.split()) <= 6 and  # Shorter requirement for keyword headers
                      not line_clean.endswith('.') and  # Not a sentence
                      (line_clean.startswith(tuple('0123456789')) or line_clean.isupper())):  # Either numbered or all caps
                    is_section_header = True
                    header_title = line_clean
            
            if is_section_header and current_section_content:
                # Save previous section
                sections[current_section_title] = '\n'.join(current_section_content)
                current_section_title = header_title
                current_section_content = [line]  # Include the header line in the new section
            elif is_section_header and not current_section_content:
                # First section header - just set title and include the header line
                current_section_title = header_title
                current_section_content = [line]
            else:
                current_section_content.append(line)
        
        # Add the last section
        if current_section_content:
            sections[current_section_title] = '\n'.join(current_section_content)
        
        return sections
    
    def _split_into_sections(self, document_text: str, section_size: int = 4000) -> List[str]:
        """Split a document into logical sections with overlap to ensure no requirements are missed"""
        sections = []
        overlap_size = 500  # Overlap to ensure continuity
        
        # Look for section headers (numbered sections, chapters, etc.)
        lines = document_text.split('\n')
        current_section = []
        current_size = 0
        
        for line in lines:
            line_clean = line.strip()
            
            # Check if this looks like a section header
            is_section_header = (
                line_clean and (
                    re.match(r'^\d+(\.\d+)*\.?\s+[A-Z]', line_clean) or  # 4.1 REQUIREMENTS
                    line_clean.startswith(('Chapter', 'Section', 'Part', 'Article', 'APPENDIX')) or
                    (line_clean.isupper() and len(line_clean.split()) <= 8 and len(line_clean) > 5)
                )
            )
            
            # If we hit a section header and current section is substantial, save it
            if is_section_header and current_section and current_size > 1000:
                sections.append('\n'.join(current_section))
                current_section = [line]
                current_size = len(line)
            else:
                current_section.append(line)
                current_size += len(line)
                
                # If section gets too large, split it with overlap
                if current_size > section_size:
                    section_text = '\n'.join(current_section)
                    sections.append(section_text)
                    
                    # Keep last part for overlap
                    overlap_lines = current_section[-overlap_size//50:]  # Roughly 500 chars worth
                    current_section = overlap_lines
                    current_size = sum(len(line) for line in overlap_lines)
        
        # Add remaining content
        if current_section:
            sections.append('\n'.join(current_section))
            
        # If no logical sections found, create overlapping chunks
        if len(sections) <= 1 and len(document_text) > section_size:
            sections = []
            words = document_text.split()
            words_per_section = section_size // 5  # Average 5 chars per word
            overlap_words = overlap_size // 5
            
            start = 0
            while start < len(words):
                end = min(start + words_per_section, len(words))
                section_words = words[start:end]
                sections.append(' '.join(section_words))
                start = end - overlap_words  # Overlap
                
                if start >= len(words) - overlap_words:
                    break
                
        return sections if sections else [document_text]

    def _fetch_templates(self, collection: str) -> List[str]:
        resp = requests.get(f"{self.chroma_url}/documents",
                            params={"collection_name": collection})
        resp.raise_for_status()
        return resp.json().get("documents", [])

    def _retrieve_context(self, tmpl: str, sources: List[str], top_k: int) -> str:
        pieces = []
        for coll in sources:
            docs, ok = self.rag.get_relevant_documents(tmpl, coll)
            if ok:
                pieces += docs[:top_k]
        return "\n\n".join(pieces)

    
    def _fetch_templates(self, collection: str) -> List[str]:
        resp = requests.get(f"{self.chroma_url}/documents",
                            params={"collection_name": collection})
        resp.raise_for_status()
        return resp.json().get("documents", [])
    
    def generate_test_plan(
        self,
        source_collections: Optional[List[str]] = None,
        source_doc_ids: Optional[List[str]] = None,
        doc_title: Optional[str] = None,
        agent_set_id: int = None,
        pipeline_id: str = None,
    ) -> List[dict]:
        """
        Generate test plan using multi-agent architecture:
        1. Actor agents per section extract requirements
        2. Critic agent synthesizes actor outputs per section
        3. Final Critic agent consolidates all sections
        4. Redis pipeline for scalable processing
        5. Export to Word document and save to generated_documents

        Args:
            source_collections: List of ChromaDB collections
            source_doc_ids: List of specific document IDs
            doc_title: Title for the generated test plan
            agent_set_id: Required agent set ID for orchestration
        """
        try:
            print("===== STARTING MULTI-AGENT TEST PLAN GENERATION =====")

            # Use the multi-agent test plan service
            test_plan_result = self.multi_agent_test_plan_service.generate_multi_agent_test_plan(
                source_collections=source_collections or [],
                source_doc_ids=source_doc_ids or [],
                doc_title=doc_title or "Test Plan",
                agent_set_id=agent_set_id,
                pipeline_id=pipeline_id
            )
            
            print(f"Multi-agent pipeline generated: {test_plan_result.total_requirements} requirements, {test_plan_result.total_test_procedures} procedures from {test_plan_result.total_sections} sections")

            docx_b64 = None
            chromadb_result = {}
            processing_status = getattr(test_plan_result, 'processing_status', 'COMPLETED')

            # Save to ChromaDB for all statuses except ABORTED
            # This ensures we capture COMPLETED, FALLBACK, and FAILED test plans
            if processing_status != 'ABORTED':
                # Export to Word document
                docx_b64 = self.multi_agent_test_plan_service.export_to_word(test_plan_result)
                # Save to ChromaDB generated_test_plan collection
                session_id = str(uuid.uuid4())
                chromadb_result = self.multi_agent_test_plan_service.save_to_chromadb(
                    test_plan_result,
                    session_id,
                    pipeline_id=getattr(test_plan_result, 'pipeline_id', None)
                )

                if chromadb_result.get('saved'):
                    print(f"✓ Test plan saved to ChromaDB: {chromadb_result.get('document_id')} in collection '{chromadb_result.get('collection_name')}' (status: {processing_status})")
                else:
                    print(f"✗ Failed to save test plan to ChromaDB: {chromadb_result.get('error', 'Unknown error')} (status: {processing_status})")
            else:
                print(f"Skipping ChromaDB save for ABORTED test plan")
            
            return [{
                "title": test_plan_result.title,
                "content": test_plan_result.consolidated_markdown,
                "docx_b64": docx_b64,
                "document_id": chromadb_result.get("document_id"),
                "collection_name": chromadb_result.get("collection_name"),
                "generated_at": chromadb_result.get("generated_at"),
                "processing_status": test_plan_result.processing_status,
                "meta": {
                    "architecture": "multi_agent_gpt4_pipeline",
                    "total_sections": test_plan_result.total_sections,
                    "total_requirements": test_plan_result.total_requirements,
                    "total_test_procedures": test_plan_result.total_test_procedures,
                    "agent_configuration": "3x_gpt4_actors_1x_critic_1x_final_critic",
                    "redis_pipeline": True,
                    "scalable_processing": True,
                    "chromadb_saved": chromadb_result.get("saved", False),
                    "sections_processed": len(test_plan_result.sections),
                    "pipeline_id": pipeline_id
                }
            }]
            
        except Exception as e:
            print(f"Error in multi-agent test plan generation: {e}")
            import traceback
            traceback.print_exc()
            return [{
                "title": doc_title or "Test Plan",
                "content": f"Error in multi-agent generation: {str(e)}",
                "error": str(e),
                "processing_status": "ERROR"
            }]

    def _convert_markdown_to_docx(self, markdown_content: str, doc: Document):
        """Convert markdown content to Word document structure (like notebook markdown_to_docx)"""
        lines = markdown_content.split('\n')
        
        for line in lines:
            line_clean = line.strip()
            if not line_clean:
                continue
                
            # Handle different markdown elements
            if line_clean.startswith('# '):
                # H1 - Main title (skip since already added)
                continue
            elif line_clean.startswith('## '):
                # H2 - Major sections
                doc.add_heading(line_clean[3:].strip(), level=2)
            elif line_clean.startswith('### '):
                # H3 - Sub-sections
                doc.add_heading(line_clean[4:].strip(), level=3)
            elif line_clean.startswith('#### '):
                # H4 - Sub-sub-sections
                doc.add_heading(line_clean[5:].strip(), level=4)
            elif line_clean.startswith('**') and line_clean.endswith('**'):
                # Bold text as heading
                doc.add_heading(line_clean[2:-2].strip(), level=4)
            elif line_clean.startswith('- ') or line_clean.startswith('* '):
                # Bullet points
                doc.add_paragraph(line_clean[2:].strip(), style='List Bullet')
            elif re.match(r'^\d+\.', line_clean):
                # Numbered lists
                doc.add_paragraph(line_clean, style='List Number')
            elif line_clean.startswith('**') and ':' in line_clean:
                # Bold labels (like "Dependencies:")
                parts = line_clean.split(':', 1)
                p = doc.add_paragraph()
                run = p.add_run(parts[0].replace('*', '') + ':')
                run.bold = True
                if len(parts) > 1:
                    p.add_run(' ' + parts[1])
            else:
                # Regular paragraph
                doc.add_paragraph(line_clean)
    
    def _split_large_section_for_testing(self, content: str, section_title: str) -> Dict[str, str]:
        """Split large sections into smaller units suitable for comprehensive test plan generation"""
        subsections = {}
        
        # Try to split on subsection headers first
        import re
        
        # Look for numbered subsections (4.1, 4.2, etc.)
        subsection_pattern = re.compile(r'^(\d+\.\d+(?:\.\d+)*)\s+(.+)$', re.MULTILINE)
        matches = list(subsection_pattern.finditer(content))
        
        if len(matches) > 1:
            # Split by subsections
            for i, match in enumerate(matches):
                start = match.start()
                end = matches[i+1].start() if i+1 < len(matches) else len(content)
                subsection_num = match.group(1)
                subsection_name = match.group(2)[:50]  # Limit name length
                subsection_content = content[start:end].strip()
                
                if len(subsection_content) > 500:  # Only include substantial content
                    subsections[f"{subsection_num} {subsection_name}"] = subsection_content
        
        # If no subsections found, split by paragraph blocks
        if not subsections:
            paragraphs = content.split('\n\n')
            current_block = ""
            block_num = 1
            
            for paragraph in paragraphs:
                if len(current_block + paragraph) > 5000:  # Target size per block
                    if current_block.strip():
                        subsections[f"Block {block_num}"] = current_block.strip()
                        block_num += 1
                        current_block = paragraph + "\n\n"
                    else:
                        current_block += paragraph + "\n\n"
                else:
                    current_block += paragraph + "\n\n"
            
            # Add final block
            if current_block.strip():
                subsections[f"Block {block_num}"] = current_block.strip()
        
        print(f"DEBUG: Split large section '{section_title}' into {len(subsections)} testable units")
        return subsections
    
    def _extract_text_from_docx_object(self, doc) -> str:
        """Extract plain text from a python-docx Document object"""
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        return "\n\n".join(text_parts)
    
    def _save_to_vector_store(self, title: str, content: str, template_collection: str, 
                            agent_ids: List[int], session_id: str) -> Dict[str, Any]:
        """Save generated document to ChromaDB vector store"""
        try:
            from datetime import datetime
            import json
            
            # Create a collection name for generated documents
            generated_collection = "generated_documents"
            
            # First, ensure the collection exists
            try:
                # Check if collection exists by listing collections
                list_response = requests.get(
                    f"{self.chroma_url}/collections",
                    timeout=5
                )
                
                if list_response.ok:
                    collections = list_response.json().get("collections", [])
                    if generated_collection not in collections:
                        # Collection doesn't exist, create it using the proper endpoint
                        print(f"Creating collection '{generated_collection}'...")
                        create_response = requests.post(
                            f"{self.chroma_url}/collection/create",
                            params={"collection_name": generated_collection},
                            timeout=10
                        )
                        
                        if create_response.ok:
                            print(f"Collection '{generated_collection}' created successfully")
                        else:
                            print(f"Failed to create collection: {create_response.status_code} - {create_response.text}")
                    else:
                        print(f"Collection '{generated_collection}' already exists")
                else:
                    print(f"Failed to list collections: {list_response.status_code} - {list_response.text}")
                        
            except Exception as e:
                print(f"Collection setup error: {e}")
            
            # Create unique document ID
            doc_id = f"gen_doc_{session_id}_{title.replace(' ', '_').replace('/', '_')}"
            
            # Prepare metadata
            metadata = {
                "title": title,
                "type": "generated_document",
                "template_collection": template_collection,
                "agent_ids": json.dumps(agent_ids),
                "session_id": session_id,
                "generated_at": datetime.now().isoformat(),
                "word_count": len(content.split()),
                "char_count": len(content)
            }
            
            # Store in ChromaDB via API
            payload = {
                "collection_name": generated_collection,
                "documents": [content],
                "metadatas": [metadata],
                "ids": [doc_id]
            }
            
            response = requests.post(
                f"{self.chroma_url}/documents/add",
                json=payload,
                timeout=30
            )
            
            if response.ok:
                print(f"Saved generated document to vector store: {doc_id}")
                return {
                    "document_id": doc_id,
                    "collection_name": generated_collection,
                    "generated_at": metadata["generated_at"],
                    "saved": True
                }
            else:
                print(f"Failed to save to vector store: {response.status_code} - {response.text}")
                return {"saved": False, "error": response.text}
                
        except Exception as e:
            print(f"Error saving to vector store: {e}")
            return {"saved": False, "error": str(e)}
    
    # def _generate_test_plan_fallback_mode(self, doc_title: str, source_collections: Optional[List[str]], 
    #                                     source_doc_ids: Optional[List[str]], max_workers: int) -> List[dict]:
    #     """
    #     Fallback test plan generation when ChromaDB is empty or not accessible.
    #     Creates a comprehensive test plan template that can be customized based on common compliance patterns.
    #     """
    #     print("=== FALLBACK MODE: DIRECT TEST PLAN GENERATION ===")
    #     print("Generating comprehensive test plan template without document-specific content")
        
    #     try:
    #         # Create a comprehensive test plan based on common military/technical standards patterns
    #         fallback_sections = self._create_comprehensive_test_template()
            
    #         # Use the optimized test plan service with direct content instead of RAG
    #         document_id = f"fallback_testplan_{uuid.uuid4().hex[:12]}"
            
    #         # Temporarily configure the service for direct queries instead of RAG
    #         collection_name = source_collections[0] if source_collections else "fallback_collection"
    #         self.optimized_test_plan_service.collection_name = collection_name
    #         self.optimized_test_plan_service.rule_extraction_agent.collection_name = collection_name
    #         self.optimized_test_plan_service.test_step_generator.collection_name = collection_name
            
    #         print(f"Fallback mode: Processing {len(fallback_sections)} template sections")
            
    #         # Generate test plan using the template sections
    #         test_plan_result = self.optimized_test_plan_service.generate_comprehensive_test_plan_streaming(
    #             sections=fallback_sections,
    #             document_id=document_id,
    #             max_workers=max_workers
    #         )
            
    #         # Create Word document
    #         doc = Document()
    #         doc.add_heading(doc_title, level=1)
            
    #         # Add notice about fallback mode
    #         notice_para = doc.add_paragraph()
    #         notice_run = notice_para.add_run(" NOTICE: ")
    #         notice_run.bold = True
    #         notice_para.add_run("This test plan was generated in fallback mode due to missing source documents. "
    #                           "Please customize the test procedures based on your specific requirements and upload "
    #                           "source documents to ChromaDB for more accurate test plan generation.")
            
    #         # Add the generated content
    #         final_markdown = test_plan_result.consolidated_plan
    #         if final_markdown and final_markdown.startswith('#'):
    #             self._convert_markdown_to_docx(final_markdown, doc)
    #         else:
    #             doc.add_heading('Generated Test Sections', level=2)
    #             doc.add_paragraph("The following test sections were generated based on common compliance patterns:")
                
    #             # Add template sections directly if markdown conversion fails
    #             for section_title, section_content in fallback_sections.items():
    #                 doc.add_heading(section_title.replace("Template - ", ""), level=3)
    #                 doc.add_paragraph(section_content[:500] + "..." if len(section_content) > 500 else section_content)
            
    #         # Add customization guidance
    #         doc.add_heading('Customization Required', level=2)
    #         customization_items = [
    #             "Review and modify test procedures based on actual system requirements",
    #             "Add specific test data, configurations, and environment setups", 
    #             "Define acceptance criteria and pass/fail thresholds",
    #             "Specify test equipment, tools, and measurement methods",
    #             "Add traceability to specific requirement documents",
    #             "Include risk assessments and mitigation strategies"
    #         ]
    #         for item in customization_items:
    #             doc.add_paragraph(f"• {item}", style='List Bullet')
            
    #         # Serialize document
    #         buf = io.BytesIO()
    #         doc.save(buf)
            
    #         return [{
    #             "title": doc_title,
    #             "docx_b64": base64.b64encode(buf.getvalue()).decode("utf-8"),
    #             "document_id": document_id,
    #             "processing_status": "generated_fallback_mode",
    #             "meta": {
    #                 "architecture": "fallback_template_based",
    #                 "total_sections": len(fallback_sections),
    #                 "mode": "fallback_no_chromadb",
    #                 "customization_required": True,
    #                 "template_based": True,
    #                 "notice": "Generated without source documents - requires customization"
    #             }
    #         }]
            
    #     except Exception as e:
    #         print(f"Error in fallback mode test plan generation: {e}")
            
    #         # Ultra-simple fallback - just create a basic template
    #         doc = Document()
    #         doc.add_heading(doc_title, level=1)
            
    #         error_para = doc.add_paragraph()
    #         error_run = error_para.add_run("ERROR: ")
    #         error_run.bold = True
    #         error_para.add_run(f"Test plan generation failed: {str(e)}")
            
    #         doc.add_heading('Manual Test Plan Template', level=2)
    #         doc.add_paragraph("Please create a test plan manually using the following structure:")
            
    #         basic_sections = [
    #             "1. Test Scope and Objectives",
    #             "2. Test Environment Setup",
    #             "3. Functional Requirements Testing",
    #             "4. Performance Requirements Testing",
    #             "5. Interface Testing",
    #             "6. Compliance Testing",
    #             "7. Test Procedures and Expected Results",
    #             "8. Acceptance Criteria",
    #             "9. Risk Assessment",
    #             "10. Test Schedule and Resources"
    #         ]
            
    #         for section in basic_sections:
    #             doc.add_paragraph(section, style='List Number')
    #             doc.add_paragraph(f"[TODO: Define test procedures for {section.split('. ')[1]}]")
            
    #         buf = io.BytesIO()
    #         doc.save(buf)
            
    #         return [{
    #             "title": doc_title,
    #             "docx_b64": base64.b64encode(buf.getvalue()).decode("utf-8"),
    #             "document_id": "error_fallback",
    #             "processing_status": "error_basic_template",
    #             "error": str(e),
    #             "meta": {
    #                 "mode": "error_recovery",
    #                 "template_only": True
    #             }
    #         }]
