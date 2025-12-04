"""
Word Export Service
Provides Word document export capabilities for agents, chat history, and agent responses.
"""

import os
import base64
import uuid
from datetime import datetime, timezone
from typing import Dict, List, Any, Optional, Union
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from sqlalchemy.orm import Session
from models.agent import ComplianceAgent
from models.chat import ChatHistory
from models.response import AgentResponse
from models.session import DebateSession
import logging
import tempfile
import subprocess
import re

logger = logging.getLogger("WORD_EXPORT_SERVICE")


def sanitize_for_xml(text: str) -> str:
    """
    Sanitize text to be XML-compatible for Word documents.
    Removes NULL bytes and control characters that are invalid in XML.

    Valid XML characters per spec:
    #x9 | #xA | #xD | [#x20-#xD7FF] | [#xE000-#xFFFD] | [#x10000-#x10FFFF]
    """
    if not text:
        return ""

    # Pattern to match invalid XML characters
    # This removes NULL bytes (0x00) and control characters except tab (0x09), newline (0x0A), carriage return (0x0D)
    invalid_xml_chars = re.compile(
        r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F-\x9F]'
    )

    # Replace invalid characters with empty string
    sanitized = invalid_xml_chars.sub('', text)

    return sanitized


class WordExportService:
    """
    Service for exporting various application data to Word documents.
    Enhanced with streaming capabilities and performance optimizations.
    """
    
    def __init__(self, enable_streaming: bool = True, chunk_size: int = 50):
        self.temp_dir = os.path.join(os.getcwd(), "temp_exports")
        os.makedirs(self.temp_dir, exist_ok=True)
        self.enable_streaming = enable_streaming
        self.chunk_size = chunk_size  # Number of items to process per chunk
        self._doc_cache = {}  # Cache for frequently used document styles
    
    def export_agents_to_word(self, agents: List[Dict[str, Any]], export_format: str = "detailed") -> bytes:
        """
        Export agent configurations to a Word document.
        Enhanced with streaming for large datasets.
        
        Args:
            agents: List of agent dictionaries
            export_format: "summary" or "detailed"
        
        Returns:
            bytes: Word document content
        """
        try:
            start_time = datetime.now()
            logger.info(f"Starting Word export for {len(agents)} agents in {export_format} format")
            
            doc = Document()
            
            # Set up document styles with caching
            self._setup_document_styles_optimized(doc)
            
            # Title
            title = doc.add_heading('Agent Configurations Export', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Metadata
            doc.add_paragraph(f"Export Date: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S UTC')}")
            doc.add_paragraph(f"Total Agents: {len(agents)}")
            doc.add_paragraph(f"Export Format: {export_format.title()}")
            doc.add_paragraph("")
            
            # Agents section
            doc.add_heading('Agent Details', level=1)
            
            # Process agents in chunks for better memory management
            if self.enable_streaming and len(agents) > self.chunk_size:
                return self._export_agents_streaming(doc, agents, export_format)
            else:
                return self._export_agents_traditional(doc, agents, export_format)
                
        except Exception as e:
            logger.error(f"Failed to export agents to Word: {str(e)}")
            raise e
    
    def _export_agents_streaming(self, doc: Document, agents: List[Dict[str, Any]], export_format: str) -> bytes:
        """Export agents using streaming approach for better performance"""
        logger.info(f"Using streaming export for {len(agents)} agents")
        
        chunks = [agents[i:i + self.chunk_size] for i in range(0, len(agents), self.chunk_size)]
        
        for chunk_idx, chunk in enumerate(chunks):
            logger.debug(f"Processing chunk {chunk_idx + 1}/{len(chunks)} ({len(chunk)} agents)")
            
            for i, agent in enumerate(chunk, chunk_idx * self.chunk_size + 1):
                self._add_agent_to_document(doc, agent, i, export_format)
                
                # Add page break between agents (except last one)
                if i < len(agents):
                    doc.add_page_break()
        
        # Convert to bytes
        doc_bytes = self._document_to_bytes(doc)
        logger.info(f"Streaming export completed for {len(agents)} agents")
        return doc_bytes
    
    def _export_agents_traditional(self, doc: Document, agents: List[Dict[str, Any]], export_format: str) -> bytes:
        """Traditional export approach for smaller datasets"""
        logger.debug(f"Using traditional export for {len(agents)} agents")
        
        for i, agent in enumerate(agents, 1):
            self._add_agent_to_document(doc, agent, i, export_format)
            
            # Page break between agents (except last one)
            if i < len(agents):
                doc.add_page_break()
        
        # Convert to bytes
        doc_bytes = self._document_to_bytes(doc)
        logger.info(f"Traditional export completed for {len(agents)} agents")
        return doc_bytes
    
    def _add_agent_to_document(self, doc: Document, agent: Dict[str, Any], agent_num: int, export_format: str):
        """Add a single agent to the document (optimized)"""
        # Agent header
        agent_heading = doc.add_heading(f"{agent_num}. {agent.get('name', 'Unknown Agent')}", level=2)
        
        # Basic information table (optimized)
        self._add_agent_basic_info_optimized(doc, agent)
        
        if export_format == "detailed":
            # System prompt
            system_prompt = agent.get('system_prompt', 'No system prompt defined')
            if system_prompt and len(system_prompt.strip()) > 0:
                doc.add_heading('System Prompt', level=3)
                prompt_para = doc.add_paragraph(system_prompt[:1000] + "..." if len(system_prompt) > 1000 else system_prompt)
                prompt_para.style = 'Quote'
            
            # User prompt template
            user_prompt = agent.get('user_prompt_template', 'No user prompt template defined')
            if user_prompt and len(user_prompt.strip()) > 0:
                doc.add_heading('User Prompt Template', level=3)
                template_para = doc.add_paragraph(user_prompt[:1000] + "..." if len(user_prompt) > 1000 else user_prompt)
                template_para.style = 'Quote'
            
            # Performance metrics if available
            if agent.get('total_queries', 0) > 0:
                doc.add_heading('Performance Metrics', level=3)
                self._add_agent_performance_metrics_optimized(doc, agent)
    
    def export_chat_history_to_word(self, chat_history: List[Dict[str, Any]], session_filter: Optional[str] = None) -> bytes:
        """
        Export chat history to a Word document.
        
        Args:
            chat_history: List of chat history records
            session_filter: Optional session ID to filter by
        
        Returns:
            bytes: Word document content
        """
        try:
            doc = Document()
            self._setup_document_styles(doc)
            
            # Title
            title = doc.add_heading('Chat History Export', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Metadata
            doc.add_paragraph(f"Export Date: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S UTC')}")
            doc.add_paragraph(f"Total Conversations: {len(chat_history)}")
            if session_filter:
                doc.add_paragraph(f"Session Filter: {session_filter}")
            doc.add_paragraph("")
            
            # Group by session if multiple sessions
            sessions = {}
            for chat in chat_history:
                session_id = chat.get('session_id', 'unknown')
                if session_id not in sessions:
                    sessions[session_id] = []
                sessions[session_id].append(chat)
            
            # Export each session
            for session_id, chats in sessions.items():
                doc.add_heading(f'Session: {session_id}', level=1)
                
                # Session summary
                session_start = min(chat.get('timestamp', datetime.now()) for chat in chats)
                session_end = max(chat.get('timestamp', datetime.now()) for chat in chats)
                
                doc.add_paragraph(f"Session Duration: {session_start} to {session_end}")
                doc.add_paragraph(f"Total Interactions: {len(chats)}")
                doc.add_paragraph("")
                
                # Individual chats
                for i, chat in enumerate(chats, 1):
                    self._add_chat_interaction(doc, chat, i)
                
                # Page break between sessions
                if len(sessions) > 1 and session_id != list(sessions.keys())[-1]:
                    doc.add_page_break()
            
            doc_bytes = self._document_to_bytes(doc)
            logger.info(f"Exported {len(chat_history)} chat records to Word document")
            return doc_bytes
            
        except Exception as e:
            logger.error(f"Failed to export chat history to Word: {str(e)}")
            raise e
    
    def export_agent_simulation_to_word(self, simulation_data: Dict[str, Any]) -> bytes:
        """
        Export agent simulation results to a Word document.
        
        Args:
            simulation_data: Dictionary containing simulation results
        
        Returns:
            bytes: Word document content
        """
        try:
            doc = Document()
            self._setup_document_styles(doc)
            
            # Title
            title = doc.add_heading('AI Agent Simulation Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Metadata
            doc.add_paragraph(f"Export Date: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S UTC')}")
            if simulation_data.get('session_id'):
                doc.add_paragraph(f"Session ID: {simulation_data['session_id']}")
            doc.add_paragraph("")
            
            # Simulation overview
            doc.add_heading('Simulation Overview', level=1)
            
            simulation_type = simulation_data.get('type', 'Unknown')
            doc.add_paragraph(f"Simulation Type: {simulation_type}")
            
            if 'query' in simulation_data:
                doc.add_heading('Original Query/Content', level=2)
                query_para = doc.add_paragraph(simulation_data['query'])
                query_para.style = 'Quote'
            
            # Agents involved
            if 'agents' in simulation_data:
                doc.add_heading('Participating Agents', level=2)
                agents_table = doc.add_table(rows=1, cols=3)
                agents_table.style = 'Table Grid'
                
                # Header
                header_cells = agents_table.rows[0].cells
                header_cells[0].text = 'Agent Name'
                header_cells[1].text = 'Model'
                header_cells[2].text = 'Role'
                
                for agent in simulation_data['agents']:
                    row_cells = agents_table.add_row().cells
                    row_cells[0].text = agent.get('name', 'Unknown')
                    row_cells[1].text = agent.get('model_name', 'Unknown')
                    row_cells[2].text = agent.get('role', 'Analysis')
            
            # Results section
            doc.add_heading('Simulation Results', level=1)
            
            if 'agent_responses' in simulation_data:
                self._add_agent_responses(doc, simulation_data['agent_responses'])
            elif 'debate_chain' in simulation_data:
                self._add_debate_chain(doc, simulation_data['debate_chain'])
            elif 'details' in simulation_data:
                self._add_compliance_details(doc, simulation_data['details'])
            
            # Performance metrics if available
            if simulation_data.get('response_time_ms'):
                doc.add_heading('Performance Metrics', level=2)
                doc.add_paragraph(f"Total Response Time: {simulation_data['response_time_ms']/1000:.2f} seconds")
            
            doc_bytes = self._document_to_bytes(doc)
            logger.info("Exported agent simulation to Word document")
            return doc_bytes
            
        except Exception as e:
            logger.error(f"Failed to export agent simulation to Word: {str(e)}")
            raise e
    
    def export_rag_assessment_to_word(self, assessment_data: Dict[str, Any]) -> bytes:
        """
        Export RAG assessment results to a Word document.
        
        Args:
            assessment_data: Dictionary containing assessment results
        
        Returns:
            bytes: Word document content
        """
        try:
            doc = Document()
            self._setup_document_styles(doc)
            
            # Title
            title = doc.add_heading('RAG Assessment Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Metadata
            doc.add_paragraph(f"Export Date: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S UTC')}")
            if assessment_data.get('performance_metrics', {}).get('session_id'):
                doc.add_paragraph(f"Session ID: {assessment_data['performance_metrics']['session_id']}")
            doc.add_paragraph("")
            
            # Original query
            if 'performance_metrics' in assessment_data and 'query' in assessment_data['performance_metrics']:
                doc.add_heading('Original Query', level=1)
                query_para = doc.add_paragraph(assessment_data['performance_metrics']['query'])
                query_para.style = 'Quote'
            
            # Generated response
            if 'response' in assessment_data:
                doc.add_heading('Generated Response', level=1)
                response_para = doc.add_paragraph(assessment_data['response'])
                response_para.style = 'Quote'
            
            # Performance metrics
            if 'performance_metrics' in assessment_data:
                doc.add_heading('Performance Metrics', level=1)
                self._add_performance_metrics_table(doc, assessment_data['performance_metrics'])
            
            # Quality assessment
            if 'quality_assessment' in assessment_data and assessment_data['quality_assessment']:
                doc.add_heading('Quality Assessment', level=1)
                self._add_quality_assessment_table(doc, assessment_data['quality_assessment'])
            
            # Alignment assessment
            if 'alignment_assessment' in assessment_data and assessment_data['alignment_assessment']:
                doc.add_heading('Alignment Assessment', level=1)
                self._add_alignment_assessment_table(doc, assessment_data['alignment_assessment'])
            
            # Classification metrics
            if 'classification_metrics' in assessment_data and assessment_data['classification_metrics']:
                doc.add_heading('Classification Metrics', level=1)
                self._add_classification_metrics_table(doc, assessment_data['classification_metrics'])
            
            doc_bytes = self._document_to_bytes(doc)
            logger.info("Exported RAG assessment to Word document")
            return doc_bytes
            
        except Exception as e:
            logger.error(f"Failed to export RAG assessment to Word: {str(e)}")
            raise e
    
    def export_reconstructed_document_to_word(self, reconstructed: Dict[str, Any]) -> bytes:
        """
        Export a reconstructed document (text + optional images) to a Word document.

        Supports position-aware image placement by parsing markdown image syntax ![alt](url)
        and inserting images inline at their correct positions.
        """
        try:
            import re
            import tempfile
            import requests
            from docx.shared import Inches

            doc = Document()
            self._setup_document_styles(doc)

            title_text = sanitize_for_xml(reconstructed.get('document_name') or 'Reconstructed Document')
            title = doc.add_heading(title_text, 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Get reconstructed content with position-aware images
            # Sanitize content to remove invalid XML characters (NULL bytes, control chars)
            raw_content = reconstructed.get('reconstructed_content', '') or ''
            content = sanitize_for_xml(raw_content)

            if len(content) != len(raw_content):
                logger.warning(f"Sanitized {len(raw_content) - len(content)} invalid XML characters from content")

            # Debug: Log content stats
            total_lines = len(content.splitlines())
            logger.info(f"Reconstructed content: {len(content)} chars, {total_lines} lines")

            # Debug: Count how many image markers exist
            import re as re_module
            image_count = len(re_module.findall(r'!\[', content))
            logger.info(f"Found {image_count} image markers (![) in reconstructed content")

            # Use correct image storage path (FastAPI service, not ChromaDB)
            IMAGES_DIR = os.getenv("IMAGES_STORAGE_DIR", os.path.join(os.getcwd(), "stored_images"))

            # Parse markdown handling inline images
            # Pattern for markdown images: ![alt text](url)
            # IMPORTANT: Use re.DOTALL to match across newlines (in case image markdown spans lines)
            image_pattern = r'!\[([^\]]*)\]\(([^\)]+)\)'

            images_inserted = 0
            images_failed = 0

            # First, find and process all images in the content
            # Replace each image with a unique placeholder to preserve position
            image_data = []
            for match in re.finditer(image_pattern, content, re.DOTALL):
                alt_text = sanitize_for_xml(match.group(1).strip())
                image_url = match.group(2).strip()
                placeholder = f"<<<IMAGE_{len(image_data)}>>>"
                image_data.append({
                    "alt_text": alt_text,
                    "image_url": image_url,
                    "placeholder": placeholder
                })
                logger.info(f"Found image {len(image_data)}: {alt_text[:50]}...")

            # Replace images with placeholders
            for img in image_data:
                # Build the original pattern to replace
                original = f"![{img['alt_text']}]({img['image_url']})"
                content = content.replace(original, img['placeholder'])

            # Now process line by line with placeholders
            for line in content.splitlines():
                l = (line or '').strip()
                if not l:
                    doc.add_paragraph("")
                    continue

                # Check if line contains an image placeholder
                image_match = None
                for idx, img in enumerate(image_data):
                    if img['placeholder'] in l:
                        image_match = (idx, img)
                        break

                if image_match:
                    idx, img = image_match
                    alt_text = img['alt_text']
                    image_url = img['image_url']
                    placeholder = img['placeholder']

                    # Extract filename from URL
                    filename = image_url.split('/')[-1] if '/' in image_url else image_url

                    # Find placeholder position in line
                    placeholder_pos = l.index(placeholder)

                    # Add text before image placeholder (if any)
                    text_before = l[:placeholder_pos].strip()
                    if text_before:
                        doc.add_paragraph(text_before)

                    # Try to add the image
                    try:
                        image_path = os.path.join(IMAGES_DIR, filename)
                        logger.info(f" Inserting image {idx+1}/{len(image_data)}: {filename}")

                        if os.path.exists(image_path):
                            doc.add_picture(image_path, width=Inches(5.5))
                            images_inserted += 1
                            # Add caption if present
                            if alt_text:
                                para = doc.add_paragraph(alt_text)
                                para.style = 'Intense Quote'
                        else:
                            logger.warning(f" Image not found at path: {image_path}")
                            doc.add_paragraph(f"[Image: {alt_text or filename}]")
                            images_failed += 1
                    except Exception as e:
                        logger.error(f" Failed to insert image {filename}: {e}")
                        doc.add_paragraph(f"[Image: {alt_text or filename}]")
                        images_failed += 1

                    # Add text after image placeholder (if any)
                    text_after = l[placeholder_pos + len(placeholder):].strip()
                    if text_after:
                        doc.add_paragraph(text_after)

                # Handle regular markdown formatting
                elif l.startswith('### '):
                    doc.add_heading(l[4:].strip(), level=3)
                elif l.startswith('## '):
                    doc.add_heading(l[3:].strip(), level=2)
                elif l.startswith('# '):
                    doc.add_heading(l[2:].strip(), level=1)
                elif l.startswith(('-', '*', '•')) and not l.startswith('*'):
                    doc.add_paragraph(l.lstrip('-*• ').strip(), style='List Bullet')
                elif l.startswith('*') and l.endswith('*') and not l.startswith('**') and len(l) > 2:
                    # Italic caption line (from markdown)
                    # Don't add if it's just after an image (avoid duplicate captions)
                    pass  # Skip these as they're handled in alt_text
                else:
                    # Regular paragraph - handle bold formatting
                    if '**' in l and l.count('**') >= 2:
                        # Has bold text
                        para = doc.add_paragraph()
                        parts = l.split('**')
                        for i, part in enumerate(parts):
                            if part:
                                run = para.add_run(part)
                                if i % 2 == 1:  # Odd indices are bold
                                    run.bold = True
                    else:
                        doc.add_paragraph(l)

            logger.info(f"Word export complete: {images_inserted} images inserted, {images_failed} failed")
            return self._document_to_bytes(doc)
        except Exception as e:
            logger.error(f"Failed to export reconstructed document to Word: {str(e)}")
            raise e

    def export_legal_research_to_word(self, research_data: Dict[str, Any]) -> bytes:
        """
        Export legal research results to a Word document.

        Args:
            research_data: Dictionary containing legal research results

        Returns:
            bytes: Word document content
        """
        try:
            doc = Document()
            self._setup_document_styles(doc)

            # Title
            title = doc.add_heading('Legal Research Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Metadata
            doc.add_paragraph(f"Export Date: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S UTC')}")
            doc.add_paragraph("")

            # Query
            if research_data.get('query'):
                doc.add_heading('Research Query', level=1)
                query_para = doc.add_paragraph(research_data['query'])
                query_para.style = 'Quote'
                doc.add_paragraph("")

            # Metrics
            doc.add_heading('Research Summary', level=1)
            table = doc.add_table(rows=4, cols=2)
            table.style = 'Table Grid'

            table.cell(0, 0).text = 'Internal Citations'
            table.cell(0, 1).text = str(len(research_data.get('internal_citations', [])))
            table.cell(1, 0).text = 'External Citations'
            table.cell(1, 1).text = str(len(research_data.get('external_citations', [])))
            table.cell(2, 0).text = 'Collections Searched'
            table.cell(2, 1).text = str(len(research_data.get('collections_searched', [])))
            table.cell(3, 0).text = 'Response Time'
            table.cell(3, 1).text = f"{research_data.get('response_time_ms', 0)}ms"

            doc.add_paragraph("")

            # Main response
            if research_data.get('response'):
                doc.add_heading('Research Response', level=1)
                response_para = doc.add_paragraph(research_data['response'])
                doc.add_paragraph("")

            # Agent analyses
            if research_data.get('agent_analyses'):
                doc.add_heading('Agent Analyses', level=1)
                for idx, analysis in enumerate(research_data['agent_analyses'], 1):
                    doc.add_heading(f"{analysis.get('agent_name', f'Agent {idx}')} ({analysis.get('agent_model', 'Unknown')})", level=2)
                    doc.add_paragraph(analysis.get('analysis', 'No analysis provided'))
                    if analysis.get('response_time_ms'):
                        doc.add_paragraph(f"Response Time: {analysis['response_time_ms']}ms")
                    doc.add_paragraph("")

            # Internal citations
            if research_data.get('internal_citations'):
                doc.add_heading(f'Internal Citations ({len(research_data["internal_citations"])})', level=1)
                for i, citation in enumerate(research_data['internal_citations'], 1):
                    doc.add_heading(f'[{i}] {citation["document_name"]}', level=2)

                    # Citation metadata table
                    cite_table = doc.add_table(rows=3, cols=2)
                    cite_table.style = 'Table Grid'
                    cite_table.cell(0, 0).text = 'Collection'
                    cite_table.cell(0, 1).text = citation.get('collection_name', 'N/A')
                    cite_table.cell(1, 0).text = 'Page'
                    cite_table.cell(1, 1).text = str(citation.get('page_number', 'N/A'))
                    cite_table.cell(2, 0).text = 'Relevance Score'
                    cite_table.cell(2, 1).text = f"{citation.get('relevance_score', 0):.2f}"

                    # Excerpt
                    if citation.get('excerpt'):
                        doc.add_paragraph("")
                        doc.add_paragraph("Excerpt:", style='Heading 3')
                        excerpt_para = doc.add_paragraph(citation['excerpt'])
                        excerpt_para.style = 'Quote'

                    doc.add_paragraph("")

            # External citations
            if research_data.get('external_citations'):
                doc.add_heading(f'External Citations ({len(research_data["external_citations"])})', level=1)
                start_idx = len(research_data.get('internal_citations', [])) + 1
                for i, citation in enumerate(research_data['external_citations'], start_idx):
                    doc.add_heading(f'[{i}] {citation["document_name"]}', level=2)

                    # Citation info
                    if citation.get('citation_format'):
                        doc.add_paragraph(f"Citation: {citation['citation_format']}")
                    if citation.get('url'):
                        doc.add_paragraph(f"URL: {citation['url']}")
                    if citation.get('metadata'):
                        metadata = citation['metadata']
                        if metadata.get('court'):
                            doc.add_paragraph(f"Court: {metadata['court']}")
                        if metadata.get('date'):
                            doc.add_paragraph(f"Date: {metadata['date']}")
                    if citation.get('relevance_score'):
                        doc.add_paragraph(f"Relevance: {citation['relevance_score']*100:.0f}%")

                    # Excerpt
                    if citation.get('excerpt'):
                        doc.add_paragraph("")
                        doc.add_paragraph("Excerpt:", style='Heading 3')
                        excerpt_para = doc.add_paragraph(citation['excerpt'])
                        excerpt_para.style = 'Quote'

                    doc.add_paragraph("")

            doc_bytes = self._document_to_bytes(doc)
            logger.info("Exported legal research to Word document")
            return doc_bytes

        except Exception as e:
            logger.error(f"Failed to export legal research to Word: {str(e)}")
            raise e

    def export_markdown_to_word(self, title: str, markdown_content: str) -> bytes:
        """Export generic markdown-like content to a Word document."""
        try:
            doc = Document()
            self._setup_document_styles(doc)

            # Title
            if title:
                t = doc.add_heading(title, 0)
                t.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Basic markdown handling
            for line in (markdown_content or '').splitlines():
                l = (line or '').strip()
                if not l:
                    doc.add_paragraph("")
                    continue
                if l.startswith('### '):
                    doc.add_heading(l[4:].strip(), level=3)
                elif l.startswith('## '):
                    doc.add_heading(l[3:].strip(), level=2)
                elif l.startswith('# '):
                    doc.add_heading(l[2:].strip(), level=1)
                elif l.startswith(('-', '*', '•')):
                    doc.add_paragraph(l.lstrip('-*• ').strip(), style='List Bullet')
                elif l[:2].isdigit() and len(l) > 2 and l[2] in ('.', ')'):
                    doc.add_paragraph(l, style='List Number')
                else:
                    # Handle inline bold via **text**
                    if '**' in l:
                        parts = l.split('**')
                        p = doc.add_paragraph()
                        bold = False
                        for part in parts:
                            run = p.add_run(part)
                            if bold:
                                run.bold = True
                            bold = not bold
                    else:
                        doc.add_paragraph(l)

            return self._document_to_bytes(doc)
        except Exception as e:
            logger.error(f"Failed to export markdown to Word: {str(e)}")
            raise e
    
    def _setup_document_styles(self, doc: Document):
        """Set up custom styles for professional-looking document."""
        try:
            # Set document-wide font defaults
            styles = doc.styles

            # Configure Normal style (base for all paragraphs)
            normal_style = styles['Normal']
            normal_font = normal_style.font
            normal_font.name = 'Calibri'
            normal_font.size = Pt(11)
            normal_style.paragraph_format.space_after = Pt(10)
            normal_style.paragraph_format.line_spacing = 1.15

            # Configure heading styles
            for level in range(1, 4):
                heading_style = styles[f'Heading {level}']
                heading_font = heading_style.font
                heading_font.name = 'Calibri'
                heading_font.bold = True
                heading_font.color.rgb = None  # Keep default blue

                if level == 1:
                    heading_font.size = Pt(16)
                    heading_style.paragraph_format.space_before = Pt(18)
                    heading_style.paragraph_format.space_after = Pt(12)
                elif level == 2:
                    heading_font.size = Pt(14)
                    heading_style.paragraph_format.space_before = Pt(12)
                    heading_style.paragraph_format.space_after = Pt(6)
                elif level == 3:
                    heading_font.size = Pt(12)
                    heading_style.paragraph_format.space_before = Pt(10)
                    heading_style.paragraph_format.space_after = Pt(6)

            # Quote/Caption style for image captions
            if 'Quote' not in [s.name for s in styles]:
                quote_style = styles.add_style('Quote', WD_STYLE_TYPE.PARAGRAPH)
            else:
                quote_style = styles['Quote']

            quote_font = quote_style.font
            quote_font.name = 'Calibri'
            quote_font.italic = True
            quote_font.size = Pt(10)
            quote_font.color.theme_color = 8  # Gray
            quote_style.paragraph_format.left_indent = Inches(0.25)
            quote_style.paragraph_format.space_after = Pt(12)

            # Intense Quote for image captions (darker, centered)
            if 'Intense Quote' not in [s.name for s in styles]:
                intense_quote_style = styles.add_style('Intense Quote', WD_STYLE_TYPE.PARAGRAPH)
            else:
                intense_quote_style = styles['Intense Quote']

            intense_quote_font = intense_quote_style.font
            intense_quote_font.name = 'Calibri'
            intense_quote_font.italic = True
            intense_quote_font.size = Pt(10)
            intense_quote_font.color.theme_color = 8
            intense_quote_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            intense_quote_style.paragraph_format.space_after = Pt(12)

            # Set document margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)

        except Exception as e:
            logger.warning(f"Failed to set up document styles: {e}")
            pass  # Ignore styling errors
    
    def _setup_document_styles_optimized(self, doc: Document):
        """Set up custom styles for the document with caching."""
        doc_id = id(doc)
        
        if doc_id in self._doc_cache:
            logger.debug("Using cached document styles")
            return
        
        try:
            # Create custom styles
            styles = doc.styles
            
            # Quote style
            if 'Quote' not in [style.name for style in styles]:
                quote_style = styles.add_style('Quote', WD_STYLE_TYPE.PARAGRAPH)
                quote_font = quote_style.font
                quote_font.italic = True
                quote_font.size = Pt(10)
                quote_style.paragraph_format.left_indent = Inches(0.5)
                quote_style.paragraph_format.right_indent = Inches(0.5)
                
                # Cache the setup
                self._doc_cache[doc_id] = True
                logger.debug("Document styles cached")
        except Exception as e:
            logger.warning(f"Failed to set up document styles: {e}")
    
    def _add_agent_basic_info_optimized(self, doc: Document, agent: Dict[str, Any]):
        """Add basic agent information as a table (optimized version)."""
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Table Grid'
        
        # Prepare data more efficiently
        created_at = agent.get('created_at', 'Unknown')
        if isinstance(created_at, str) and len(created_at) > 10:
            created_at = created_at[:10]  # Just date part
        
        rows_data = [
            ('Agent ID', str(agent.get('id', 'N/A'))),
            ('Model', agent.get('model_name', 'Unknown')),
            ('Created', created_at),
            ('Status', 'Active' if agent.get('is_active', True) else 'Inactive'),
            ('Temperature', str(agent.get('temperature', 'N/A'))),
            ('Max Tokens', str(agent.get('max_tokens', 'N/A')))
        ]
        
        # Batch fill table cells
        for i, (label, value) in enumerate(rows_data):
            row = table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
        
        doc.add_paragraph("")
    
    def _add_agent_performance_metrics_optimized(self, doc: Document, agent: Dict[str, Any]):
        """Add agent performance metrics as a table (optimized version)."""
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'
        
        # Pre-calculate metrics
        success_rate = agent.get('success_rate', 0)
        success_rate_str = f"{success_rate*100:.1f}%" if success_rate else 'N/A'
        
        rows_data = [
            ('Total Queries', str(agent.get('total_queries', 0))),
            ('Average Response Time', f"{agent.get('avg_response_time_ms', 0):.0f}ms"),
            ('Success Rate', success_rate_str),
            ('Chain Type', agent.get('chain_type', 'basic'))
        ]
        
        # Batch fill table cells
        for i, (label, value) in enumerate(rows_data):
            row = table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
        
        doc.add_paragraph("")
    
    def _add_agent_basic_info(self, doc: Document, agent: Dict[str, Any]):
        """Add basic agent information as a table."""
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Table Grid'
        
        rows_data = [
            ('Agent ID', str(agent.get('id', 'N/A'))),
            ('Model', agent.get('model_name', 'Unknown')),
            ('Created', agent.get('created_at', 'Unknown')[:10] if agent.get('created_at') else 'Unknown'),
            ('Status', 'Active' if agent.get('is_active', True) else 'Inactive'),
            ('Temperature', str(agent.get('temperature', 'N/A'))),
            ('Max Tokens', str(agent.get('max_tokens', 'N/A')))
        ]
        
        for i, (label, value) in enumerate(rows_data):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = value
        
        doc.add_paragraph("")
    
    def _add_agent_performance_metrics(self, doc: Document, agent: Dict[str, Any]):
        """Add agent performance metrics as a table."""
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'
        
        rows_data = [
            ('Total Queries', str(agent.get('total_queries', 0))),
            ('Average Response Time', f"{agent.get('avg_response_time_ms', 0):.0f}ms"),
            ('Success Rate', f"{agent.get('success_rate', 0)*100:.1f}%" if agent.get('success_rate') else 'N/A'),
            ('Chain Type', agent.get('chain_type', 'basic'))
        ]
        
        for i, (label, value) in enumerate(rows_data):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = value
        
        doc.add_paragraph("")
    
    def _add_chat_interaction(self, doc: Document, chat: Dict[str, Any], interaction_num: int):
        """Add a single chat interaction to the document."""
        doc.add_heading(f'Interaction {interaction_num}', level=2)
        
        # Timestamp
        timestamp = chat.get('timestamp', datetime.now())
        doc.add_paragraph(f"Time: {timestamp}")
        
        # Model used
        if chat.get('model_used'):
            doc.add_paragraph(f"Model: {chat['model_used']}")
        
        # Query type
        if chat.get('query_type'):
            doc.add_paragraph(f"Type: {chat['query_type']}")
        
        # User query
        doc.add_heading('User Query:', level=3)
        query_para = doc.add_paragraph(chat.get('user_query', 'No query recorded'))
        query_para.style = 'Quote'
        
        # Response
        doc.add_heading('AI Response:', level=3)
        response_para = doc.add_paragraph(chat.get('response', 'No response recorded'))
        response_para.style = 'Quote'
        
        # Response time
        if chat.get('response_time_ms'):
            doc.add_paragraph(f"Response Time: {chat['response_time_ms']/1000:.2f} seconds")
        
        doc.add_paragraph("")
    
    def _add_agent_responses(self, doc: Document, agent_responses: Dict[str, Any]):
        """Add agent responses to the document."""
        for i, (agent_name, response) in enumerate(agent_responses.items(), 1):
            doc.add_heading(f'{i}. {agent_name}', level=2)
            response_para = doc.add_paragraph(response)
            response_para.style = 'Quote'
            doc.add_paragraph("")
    
    def _add_debate_chain(self, doc: Document, debate_chain: List[Dict[str, Any]]):
        """Add debate chain results to the document."""
        for i, round_result in enumerate(debate_chain, 1):
            agent_name = round_result.get('agent_name', f'Agent {i}')
            doc.add_heading(f'Round {i}: {agent_name}', level=2)
            
            response = round_result.get('response', 'No response')
            response_para = doc.add_paragraph(response)
            response_para.style = 'Quote'
            
            if round_result.get('agent_id'):
                doc.add_paragraph(f"Agent ID: {round_result['agent_id']}")
            
            doc.add_paragraph("")
    
    def _add_compliance_details(self, doc: Document, details: Dict[str, Any]):
        """Add compliance check details to the document."""
        for idx, analysis in details.items():
            agent_name = analysis.get('agent_name', f'Agent {idx}')
            doc.add_heading(f'{agent_name}', level=2)
            
            reason = analysis.get('reason', analysis.get('raw_text', 'No analysis'))
            reason_para = doc.add_paragraph(reason)
            reason_para.style = 'Quote'
            doc.add_paragraph("")
    
    def _add_performance_metrics_table(self, doc: Document, metrics: Dict[str, Any]):
        """Add performance metrics as a table."""
        table = doc.add_table(rows=7, cols=2)
        table.style = 'Table Grid'
        
        rows_data = [
            ('Total Time', f"{metrics.get('total_time_ms', 0):.1f}ms"),
            ('Retrieval Time', f"{metrics.get('retrieval_time_ms', 0):.1f}ms"),
            ('Generation Time', f"{metrics.get('generation_time_ms', 0):.1f}ms"),
            ('Documents Retrieved', str(metrics.get('documents_retrieved', 0))),
            ('Documents Used', str(metrics.get('documents_used', 0))),
            ('Relevance Score', f"{metrics.get('relevance_score', 0):.2f}"),
            ('Success', 'Yes' if metrics.get('success', False) else 'No')
        ]
        
        for i, (label, value) in enumerate(rows_data):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = value
        
        doc.add_paragraph("")
    
    def _add_quality_assessment_table(self, doc: Document, quality: Dict[str, Any]):
        """Add quality assessment as a table."""
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Table Grid'
        
        rows_data = [
            ('Overall Quality', f"{quality.get('overall_quality', 0):.2f}"),
            ('Relevance', f"{quality.get('relevance_score', 0):.2f}"),
            ('Coherence', f"{quality.get('coherence_score', 0):.2f}"),
            ('Factual Accuracy', f"{quality.get('factual_accuracy', 0):.2f}"),
            ('Completeness', f"{quality.get('completeness_score', 0):.2f}"),
            ('Context Utilization', f"{quality.get('context_utilization', 0):.2f}")
        ]
        
        for i, (label, value) in enumerate(rows_data):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = value
        
        doc.add_paragraph("")
    
    def _add_alignment_assessment_table(self, doc: Document, alignment: Dict[str, Any]):
        """Add alignment assessment as a table."""
        table = doc.add_table(rows=8, cols=2)
        table.style = 'Table Grid'
        
        rows_data = [
            ('Intent Alignment', f"{alignment.get('intent_alignment_score', 0):.2f}"),
            ('Query Coverage', f"{alignment.get('query_coverage_score', 0):.2f}"),
            ('Instruction Adherence', f"{alignment.get('instruction_adherence_score', 0):.2f}"),
            ('Expected Answer Type', alignment.get('expected_answer_type', 'Unknown')),
            ('Actual Answer Type', alignment.get('answer_type_classification', 'Unknown')),
            ('Answer Type Match', 'Yes' if alignment.get('answer_type_match', False) else 'No'),
            ('Tone Consistency', f"{alignment.get('tone_consistency_score', 0):.2f}"),
            ('Scope Accuracy', f"{alignment.get('scope_accuracy_score', 0):.2f}")
        ]
        
        for i, (label, value) in enumerate(rows_data):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = value
        
        # Missing elements
        if alignment.get('missing_elements'):
            doc.add_paragraph("")
            doc.add_paragraph(f"Missing Elements: {', '.join(alignment['missing_elements'])}")
        
        # Extra elements
        if alignment.get('extra_elements'):
            doc.add_paragraph(f"Extra Elements: {', '.join(alignment['extra_elements'])}")
        
        doc.add_paragraph("")
    
    def _add_classification_metrics_table(self, doc: Document, classification: Dict[str, Any]):
        """Add classification metrics as a table."""
        table = doc.add_table(rows=8, cols=2)
        table.style = 'Table Grid'
        
        rows_data = [
            ('Query Classification', classification.get('query_classification', 'Unknown')),
            ('Response Classification', classification.get('response_classification', 'Unknown')),
            ('Classification Confidence', f"{classification.get('classification_confidence', 0):.2f}"),
            ('Domain Relevance', classification.get('domain_relevance', 'Unknown')),
            ('Complexity Level', classification.get('complexity_level', 'Unknown')),
            ('Information Density', f"{classification.get('information_density', 0):.2f}"),
            ('Actionability Score', f"{classification.get('actionability_score', 0):.2f}"),
            ('Specificity Score', f"{classification.get('specificity_score', 0):.2f}")
        ]
        
        for i, (label, value) in enumerate(rows_data):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = value
        
        doc.add_paragraph("")

    def export_markdown_to_word_with_pandoc(
        self,
        title: str,
        markdown_content: str,
        reference_docx: Optional[str] = None,
        include_toc: bool = True,
        number_sections: bool = True
    ) -> bytes:
        """
        Export markdown to Word using Pandoc for professional formatting.
        Based on mil_test_plan_gen.ipynb's write_docx_with_pandoc.

        Args:
            title: Document title
            markdown_content: Markdown content to convert
            reference_docx: Path to reference .docx for styling (optional)
            include_toc: Whether to include table of contents
            number_sections: Whether to number sections automatically

        Returns:
            bytes: Word document content

        Raises:
            RuntimeError: If Pandoc is not available
        """
        try:
            logger.info(f"Exporting '{title}' to Word using Pandoc")

            # Ensure pandoc is available
            self._ensure_pandoc()

            # Sanitize markdown first
            try:
                from services.markdown_sanitization_service import MarkdownSanitizationService
                markdown_content = MarkdownSanitizationService.prepare_for_pandoc(markdown_content)
                logger.debug("Markdown sanitized for Pandoc")
            except Exception as e:
                logger.warning(f"Markdown sanitization failed, using raw content: {e}")

            # Add title if not present
            if not markdown_content.startswith('# '):
                markdown_content = f"# {title}\n\n{markdown_content}"

            # Create temp files
            with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False, encoding='utf-8') as tmp_md:
                tmp_md.write(markdown_content)
                tmp_md_path = tmp_md.name

            with tempfile.NamedTemporaryFile(mode='wb', suffix='.docx', delete=False) as tmp_docx:
                tmp_docx_path = tmp_docx.name

            try:
                # Build pandoc command
                pandoc_args = [
                    'pandoc',
                    tmp_md_path,
                    '-f', 'gfm+pipe_tables+autolink_bare_uris',  # GitHub-flavored markdown
                    '-t', 'docx',
                    '-o', tmp_docx_path,
                ]

                if include_toc:
                    pandoc_args.extend(['--toc', '--toc-depth=3'])

                if number_sections:
                    pandoc_args.append('--number-sections')

                if reference_docx and os.path.exists(reference_docx):
                    pandoc_args.extend(['--reference-doc', reference_docx])
                    logger.info(f"Using reference document: {reference_docx}")

                # Run pandoc
                logger.debug(f"Running pandoc command: {' '.join(pandoc_args)}")
                result = subprocess.run(
                    pandoc_args,
                    capture_output=True,
                    text=True,
                    timeout=120
                )

                if result.returncode != 0:
                    logger.error(f"Pandoc failed: {result.stderr}")
                    raise RuntimeError(f"Pandoc conversion failed: {result.stderr}")

                # Read generated docx
                with open(tmp_docx_path, 'rb') as f:
                    docx_bytes = f.read()

                logger.info(f"Pandoc export successful: {len(docx_bytes)} bytes, TOC={include_toc}, numbered={number_sections}")
                return docx_bytes

            finally:
                # Cleanup temp files
                try:
                    os.unlink(tmp_md_path)
                    os.unlink(tmp_docx_path)
                except Exception as e:
                    logger.warning(f"Failed to cleanup temp files: {e}")

        except RuntimeError as e:
            # Pandoc not available or conversion failed
            logger.error(f"Pandoc export failed: {e}")
            # Fallback to python-docx
            logger.info("Falling back to python-docx export")
            return self.export_markdown_to_word(title, markdown_content)
        except Exception as e:
            logger.error(f"Unexpected error in Pandoc export: {e}")
            # Fallback to python-docx
            return self.export_markdown_to_word(title, markdown_content)

    def _ensure_pandoc(self):
        """Ensure pandoc is installed and available"""
        try:
            result = subprocess.run(['pandoc', '--version'], capture_output=True, timeout=5)
            if result.returncode == 0:
                version = result.stdout.decode().splitlines()[0]
                logger.info(f"Pandoc available: {version}")
                return
        except FileNotFoundError:
            logger.error("Pandoc executable not found")
        except Exception as e:
            logger.error(f"Pandoc check failed: {e}")

        raise RuntimeError(
            "Pandoc is not installed or not in PATH. Please install via: "
            "brew install pandoc (macOS), apt-get install pandoc (Linux), "
            "or download from https://pandoc.org/installing.html"
        )

    def export_test_cards_to_word(self, test_cards: List[Dict[str, Any]], test_plan_title: str = "Test Plan") -> bytes:
        """
        Export test cards to a Word document with proper table formatting.

        Args:
            test_cards: List of test card dictionaries
            test_plan_title: Title of the test plan

        Returns:
            bytes: Word document content
        """
        try:
            doc = Document()
            self._setup_document_styles(doc)

            # Title
            title = doc.add_heading(f'Test Cards: {test_plan_title}', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Metadata
            doc.add_paragraph(f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S UTC')}")
            doc.add_paragraph(f"Total Test Cards: {len(test_cards)}")
            doc.add_paragraph("")

            if not test_cards:
                doc.add_paragraph("No test cards available.")
                return self._document_to_bytes(doc)

            # Create table with headers
            # Simplified columns to match actual test card data structure
            # Columns: Test ID | Test Title | Requirement ID | Requirement | Test Procedures | Status | Pass | Fail | Notes
            table = doc.add_table(rows=1, cols=9)
            table.style = 'Table Grid'

            # Header row
            headers = ['Test ID', 'Test Title', 'Requirement ID', 'Requirement', 'Test Procedures',
                       'Status', 'Pass', 'Fail', 'Notes']
            header_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                header_cells[i].text = header
                # Bold header text
                for paragraph in header_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

            # Add test cards
            for card in test_cards:
                metadata = card.get('metadata', {})
                content = card.get('content', '')

                row_cells = table.add_row().cells

                # Test ID
                row_cells[0].text = metadata.get('test_id', 'N/A')

                # Test Title (use document_name which is what we actually store)
                row_cells[1].text = metadata.get('document_name', metadata.get('test_title', 'N/A'))

                # Requirement ID
                row_cells[2].text = metadata.get('requirement_id', 'N/A')

                # Requirement Text
                requirement_text = metadata.get('requirement_text', 'N/A')
                if requirement_text and len(requirement_text) > 300:
                    requirement_text = requirement_text[:297] + "..."
                row_cells[3].text = requirement_text

                # Test Procedures - parse markdown table to extract just the procedures column
                procedures_text = self._extract_procedures_from_markdown_table(content) if content else 'N/A'
                if len(procedures_text) > 500:
                    procedures_text = procedures_text[:497] + "..."
                row_cells[4].text = procedures_text

                # Status (execution_status as text, not checkbox)
                execution_status = metadata.get('execution_status', 'not_executed')
                row_cells[5].text = execution_status.replace('_', ' ').title()

                # Pass checkbox
                passed = metadata.get('passed', 'false')
                row_cells[6].text = '☑' if str(passed).lower() == 'true' else '☐'

                # Fail checkbox
                failed = metadata.get('failed', 'false')
                row_cells[7].text = '☑' if str(failed).lower() == 'true' else '☐'

                # Notes
                row_cells[8].text = metadata.get('notes', '')

            # Add table borders
            self._set_table_borders(table)

            doc.add_paragraph("")
            doc.add_paragraph("Instructions: Update the Status field, fill in Pass/Fail checkboxes (☐/☑), and add notes during test execution.")
            doc.add_paragraph("Status values: Not Executed, In Progress, Completed, Failed")

            return self._document_to_bytes(doc)

        except Exception as e:
            logger.error(f"Failed to export test cards to Word: {str(e)}")
            raise e

    def _extract_procedures_from_markdown_table(self, markdown_table: str) -> str:
        """
        Extract test procedures text from markdown table.
        Parses the markdown table and extracts the Procedures column content.

        Table format: | Test ID | Test Title | Procedures | Expected Results | Acceptance Criteria | Dependencies | Executed | Pass | Fail | Notes |

        Args:
            markdown_table: Markdown table string

        Returns:
            Procedures text with <br> converted to newlines
        """
        try:
            if not markdown_table or not markdown_table.strip():
                return 'N/A'

            lines = markdown_table.strip().split('\n')

            # Need at least 3 lines (header, separator, data)
            if len(lines) < 3:
                return markdown_table  # Return as-is if not a table

            # Find the data row (skip header line 0, separator line 1)
            data_row = lines[2] if len(lines) > 2 else ""

            # Split by pipe and clean up
            cells = [cell.strip() for cell in data_row.split('|')]

            # Table format: | Test ID | Test Title | Procedures | Expected Results | Acceptance Criteria | Dependencies | Executed | Pass | Fail | Notes |
            # Cells indices: 0=empty, 1=Test ID, 2=Test Title, 3=Procedures, ...
            if len(cells) > 3:
                procedures = cells[3]  # Procedures column

                # Convert <br> to newlines for better readability
                procedures = procedures.replace('<br>', '\n').replace('<br/>', '\n')

                # Remove any remaining HTML-like tags
                import re
                procedures = re.sub(r'<[^>]+>', '', procedures)

                return procedures.strip()

            return 'N/A'

        except Exception as e:
            logger.warning(f"Failed to parse markdown table for procedures: {e}")
            return markdown_table  # Return original content if parsing fails

    def _set_table_borders(self, table):
        """Set borders for all cells in a table."""
        tbl = table._tbl
        tbl_pr = tbl.tblPr

        # Remove existing borders
        for el in tbl_pr.findall(qn('w:tblBorders')):
            tbl_pr.remove(el)

        # Add new borders
        tbl_borders = OxmlElement('w:tblBorders')
        for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            element = OxmlElement(f'w:{edge}')
            element.set(qn('w:val'), 'single')
            element.set(qn('w:sz'), '12')  # Border size
            element.set(qn('w:space'), '0')
            element.set(qn('w:color'), '000000')  # Black borders
            tbl_borders.append(element)
        tbl_pr.append(tbl_borders)

    def _document_to_bytes(self, doc: Document) -> bytes:
        """Convert Document object to bytes."""
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        return doc_io.getvalue()

    def cleanup_temp_files(self):
        """Clean up temporary files."""
        try:
            import shutil
            if os.path.exists(self.temp_dir):
                shutil.rmtree(self.temp_dir)
                os.makedirs(self.temp_dir, exist_ok=True)
        except Exception as e:
            logger.warning(f"Failed to cleanup temp files: {str(e)}")
