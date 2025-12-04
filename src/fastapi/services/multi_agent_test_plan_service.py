# services/multi_agent_test_plan_service.py
"""
Multi-Agent Test Plan Generation Service - Based on mil_test_plan_gen.ipynb

Architecture:
1. Multiple Actor Agents per section using GPT-4 (can scale by adding more agents)
2. Critic Agent synthesizes actor outputs per section using GPT-4 
3. Final Critic Agent consolidates all sections using GPT-4
4. Redis Pipeline for intermediate storage and scaling
5. ChromaDB integration for section retrieval
"""

import json
import logging
import os
import re
import uuid
from typing import Dict, List, Optional, Any
from dataclasses import dataclass
from datetime import datetime
import redis
import requests
from docx import Document
from docx.shared import Pt, RGBColor
import base64
import io
from concurrent.futures import ThreadPoolExecutor, as_completed
import time
import asyncio
from string import Formatter

# Import LLMInvoker for direct invocation with system prompt support
from services.llm_invoker import LLMInvoker

from services.llm_service import LLMService
from config.agent_registry import get_agent_registry
from repositories.agent_set_repository import AgentSetRepository
from repositories.test_plan_agent_repository import TestPlanAgentRepository
from core.database import get_db

logger = logging.getLogger(__name__)

@dataclass
class ActorResult:
    """Result from a single actor agent"""
    agent_id: str
    model_name: str
    section_title: str
    rules_extracted: str
    processing_time: float

@dataclass
class CriticResult:
    """Result from critic agent synthesis"""
    section_title: str
    synthesized_rules: str
    dependencies: List[str]
    conflicts: List[str]
    test_procedures: List[Dict[str, Any]]
    actor_count: int

@dataclass
class FinalTestPlan:
    """Final consolidated test plan"""
    title: str
    pipeline_id: str
    total_sections: int
    total_requirements: int
    total_test_procedures: int
    consolidated_markdown: str
    processing_status: str
    sections: List[CriticResult]

class MultiAgentTestPlanService:
    def __init__(self, llm_service: LLMService, chroma_url: str, fastapi_url: str = None):
        self.llm_service = llm_service
        self.chroma_url = chroma_url.rstrip("/")
        # Use FastAPI URL for vectordb endpoints if provided, otherwise fall back to chroma_url
        self.fastapi_url = (fastapi_url or chroma_url).rstrip("/")

        # Redis setup for pipeline
        redis_host = os.getenv("REDIS_HOST", "redis")
        redis_port = int(os.getenv("REDIS_PORT", 6379))
        self.redis_client = redis.Redis(host=redis_host, port=redis_port, decode_responses=True)

        # ===== AGENT REGISTRY INTEGRATION =====
        # Load agent configuration from database-backed registry
        # This provides database-first loading with ENV var overrides
        self.agent_registry = get_agent_registry()

        # Get agent models from registry (already loaded from DB or hardcoded defaults)
        self.actor_models = self.agent_registry.actor_models
        self.critic_model = self.agent_registry.critic_model
        self.final_critic_model = self.agent_registry.final_critic_model

        # Keep originals for potential fallback logging
        self._original_actor_models = list(self.actor_models)
        self._original_critic_model = self.critic_model

        # Log source of agent configuration
        db_source = "database" if self.agent_registry.is_using_database() else "hardcoded/ENV"
        logger.info(f"Agents loaded from: {db_source}")

        # Pipeline retention (seconds). Keep progress so UI can re-open later.
        try:
            self.pipeline_ttl_seconds = int(os.getenv("PIPELINE_TTL_SECONDS", str(60 * 60 * 24 * 7)))  # 7 days
        except ValueError:
            self.pipeline_ttl_seconds = 60 * 60 * 24 * 7
        
        logger.info(f"MultiAgentTestPlanService initialized with {len(self.actor_models)} GPT-4 actor agents")
        self._test_redis_connection()
    
    def _test_redis_connection(self):
        """Test Redis connection and setup pipeline keys"""
        try:
            self.redis_client.ping()
            logger.info("Redis connection successful")
        except Exception as e:
            logger.error(f"Redis connection failed: {e}")

    def _calculate_safe_max_tokens(
        self,
        model_name: str,
        system_prompt: str,
        user_prompt: str,
        requested_max_tokens: int
    ) -> int:
        """
        Calculate safe max_tokens based on model context window and input size.

        Prevents context length errors by dynamically adjusting max_tokens
        based on the actual input size.

        Args:
            model_name: Name of the LLM model
            system_prompt: System prompt text
            user_prompt: User prompt text
            requested_max_tokens: Original requested max_tokens

        Returns:
            Adjusted max_tokens that won't exceed context window
        """
        # Model context windows (tokens)
        MODEL_CONTEXT_LIMITS = {
            'gpt-4': 8192,
            'gpt-4-0613': 8192,
            'gpt-4-32k': 32768,
            'gpt-4-32k-0613': 32768,
            'gpt-4-turbo': 128000,
            'gpt-4-turbo-preview': 128000,
            'gpt-4-1106-preview': 128000,
            'gpt-4o': 128000,
            'gpt-3.5-turbo': 16385,
            'gpt-3.5-turbo-16k': 16385,
        }

        # Get model context limit (default to 8192 for unknown models)
        context_limit = MODEL_CONTEXT_LIMITS.get(model_name, 8192)

        # Estimate token count (rough approximation: 4 characters per token)
        # For production, consider using tiktoken for accurate counting
        try:
            import tiktoken
            encoding = tiktoken.encoding_for_model(model_name if model_name in MODEL_CONTEXT_LIMITS else 'gpt-4')
            input_tokens = len(encoding.encode(system_prompt + user_prompt))
        except (ImportError, Exception) as e:
            # Fallback to character-based estimation
            total_chars = len(system_prompt) + len(user_prompt)
            input_tokens = total_chars // 4
            logger.debug(f"Using character-based token estimation: {input_tokens} tokens (tiktoken not available)")

        # Reserve safety margin (100 tokens for overhead)
        safety_margin = 100
        available_tokens = context_limit - input_tokens - safety_margin

        # Return the minimum of requested tokens and available tokens
        safe_max_tokens = min(requested_max_tokens, max(available_tokens, 100))

        logger.debug(
            f"Token calculation: model={model_name}, context_limit={context_limit}, "
            f"input_tokens={input_tokens}, requested={requested_max_tokens}, "
            f"safe={safe_max_tokens}"
        )

        return safe_max_tokens

    def _load_agent_set_configuration(self, agent_set_id: int) -> Optional[Dict[str, Any]]:
        """
        Load agent set configuration from database

        Args:
            agent_set_id: ID of the agent set to load

        Returns:
            Agent set configuration dict or None if not found/error
        """
        db = None
        try:
            # Get database session
            db = next(get_db())
            repo = AgentSetRepository()

            # Load agent set
            agent_set = repo.get_by_id(agent_set_id, db)

            if not agent_set:
                logger.error(f"Agent set ID {agent_set_id} not found")
                return None

            if not agent_set.is_active:
                logger.warning(f"Agent set ID {agent_set_id} is inactive, using anyway")

            logger.info(f"Loaded agent set: {agent_set.name} (ID: {agent_set_id})")
            logger.info(f"Set type: {agent_set.set_type}, Stages: {len(agent_set.set_config.get('stages', []))}")

            # Increment usage count
            repo.increment_usage_count(agent_set_id, db)

            return agent_set.set_config

        except Exception as e:
            logger.error(f"Failed to load agent set configuration: {e}")
            return None
        finally:
            # CRITICAL: Close database session to prevent connection pool exhaustion
            if db is not None:
                db.close()

    def _execute_agent_by_id(self, agent_id: int, section_title: str, section_content: str, context_vars: Dict[str, str] = None) -> Optional[ActorResult]:
        """
        Execute a single agent by database ID

        Args:
            agent_id: Database ID of the agent to execute
            section_title: Section title for context
            section_content: Section content to process
            context_vars: Dictionary of context variables for prompt formatting

        Returns:
            ActorResult with agent's output or None if failed
        """
        db = None
        try:
            # Load agent from database
            db = next(get_db())
            repo = TestPlanAgentRepository()
            agent = repo.get_by_id(agent_id, db)

            if not agent:
                logger.error(f"Agent ID {agent_id} not found in database")
                return None

            if not agent.is_active:
                logger.warning(f"Agent ID {agent_id} ({agent.name}) is inactive, skipping")
                return None

            # Copy agent data to avoid using it after session close
            agent_name = agent.name
            agent_type = agent.agent_type
            agent_model_name = agent.model_name
            agent_system_prompt = agent.system_prompt
            agent_user_prompt_template = agent.user_prompt_template
            agent_temperature = agent.temperature
            agent_max_tokens = agent.max_tokens

            # Close database session BEFORE making LLM call (which can take a long time)
            db.close()
            db = None

            start_time = time.time()

            # Prepare prompt based on agent's template
            # Build format variables with defaults
            format_vars = {
                'section_title': section_title,
                'section_content': section_content,
                'context': context_vars.get('context', '') if context_vars else '',
                'actor_outputs': context_vars.get('actor_outputs', '') if context_vars else '',
                'critic_output': context_vars.get('critic_output', '') if context_vars else '',
                'actor_outputs_summary': context_vars.get('actor_outputs_summary', '') if context_vars else '',
                'synthesized_rules': context_vars.get('synthesized_rules', '') if context_vars else '',
                'previous_sections_summary': context_vars.get('previous_sections_summary', '') if context_vars else '',
            }

            logger.debug(f"Agent {agent_id} template variables available: {list(format_vars.keys())}")
            logger.debug(f"Agent {agent_id} context_vars keys: {list(context_vars.keys()) if context_vars else 'None'}")

            # Use simple string replacement instead of .format() to avoid issues with JSON in templates
            # Templates may contain JSON examples with curly braces that conflict with .format()
            user_prompt = agent_user_prompt_template
            for key, value in format_vars.items():
                # Replace {key} with the actual value
                user_prompt = user_prompt.replace(f'{{{key}}}', str(value))

            # Execute agent via LLM service
            logger.info(f"Executing agent: {agent_name} (ID: {agent_id}, Type: {agent_type}, Model: {agent_model_name})")

            # Dynamic token limit adjustment to prevent context length errors
            adjusted_max_tokens = self._calculate_safe_max_tokens(
                model_name=agent_model_name,
                system_prompt=agent_system_prompt,
                user_prompt=user_prompt,
                requested_max_tokens=agent_max_tokens
            )

            if adjusted_max_tokens < agent_max_tokens:
                logger.warning(
                    f"Agent {agent_id} max_tokens reduced from {agent_max_tokens} to {adjusted_max_tokens} "
                    f"to prevent context length error (input is large)"
                )

            response = LLMInvoker.invoke(
                model_name=agent_model_name,
                prompt=user_prompt,
                system_prompt=agent_system_prompt,
                temperature=agent_temperature,
                max_tokens=adjusted_max_tokens
            )

            processing_time = time.time() - start_time

            # Return as ActorResult for compatibility
            # LLMInvoker.invoke() returns a string directly
            return ActorResult(
                agent_id=f"agent_{agent_id}_{uuid.uuid4().hex[:8]}",
                model_name=agent_model_name,
                section_title=section_title,
                rules_extracted=response,  # Already a string from LLMInvoker
                processing_time=processing_time
            )

        except Exception as e:
            import traceback
            logger.error(f"Failed to execute agent ID {agent_id}: {e}")
            logger.error(f"Traceback: {traceback.format_exc()}")
            return None
        finally:
            # CRITICAL: Close database session to prevent connection pool exhaustion
            if db is not None:
                db.close()

    def _execute_stage(self, stage: Dict[str, Any], section_title: str, section_content: str, all_stage_outputs: Dict[str, List[ActorResult]] = None) -> List[ActorResult]:
        """
        Execute a single stage from agent set configuration

        Args:
            stage: Stage configuration dict with agent_ids, execution_mode, etc.
            section_title: Section title for context
            section_content: Section content to process
            all_stage_outputs: Dictionary mapping stage names to their outputs (for building context)

        Returns:
            List of ActorResult from this stage
        """
        agent_ids = stage.get('agent_ids', [])
        execution_mode = stage.get('execution_mode', 'parallel')
        stage_name = stage.get('stage_name', 'unnamed_stage')

        logger.info(f"Executing stage '{stage_name}' with {len(agent_ids)} agent(s) in {execution_mode} mode")

        # Build context variables based on previous stage outputs
        context_vars = {}

        if all_stage_outputs:
            # Build actor_outputs from actor stage
            if 'actor' in all_stage_outputs:
                actor_results = all_stage_outputs['actor']
                actor_outputs_text = "\n\n".join([
                    f"## Actor Agent {idx + 1} Output:\n{result.rules_extracted}"
                    for idx, result in enumerate(actor_results)
                ])
                context_vars['actor_outputs'] = actor_outputs_text
                context_vars['actor_outputs_summary'] = actor_outputs_text  # Use same text for summary

            # Build critic_output from critic stage
            if 'critic' in all_stage_outputs:
                critic_results = all_stage_outputs['critic']
                if critic_results:
                    # Typically critic stage has one output
                    context_vars['critic_output'] = critic_results[0].rules_extracted
                    context_vars['synthesized_rules'] = critic_results[0].rules_extracted

            # Build general context string
            all_outputs_text = ""
            for stage_key, outputs in all_stage_outputs.items():
                all_outputs_text += f"\n\n=== {stage_key.upper()} Stage Outputs ===\n\n"
                for idx, result in enumerate(outputs, 1):
                    all_outputs_text += f"Agent {idx} Output:\n{result.rules_extracted}\n\n"
            context_vars['context'] = all_outputs_text
            context_vars['previous_sections_summary'] = ""  # TODO: Track previous sections if needed

        results = []

        if execution_mode == 'parallel':
            # Execute agents in parallel
            with ThreadPoolExecutor(max_workers=len(agent_ids)) as executor:
                futures = []
                for agent_id in agent_ids:
                    future = executor.submit(
                        self._execute_agent_by_id,
                        agent_id, section_title, section_content, context_vars
                    )
                    futures.append(future)

                # Collect results
                for future in as_completed(futures):
                    try:
                        result = future.result(timeout=180)  # 3 minute timeout per agent
                        if result:
                            results.append(result)
                    except Exception as e:
                        logger.error(f"Stage '{stage_name}' agent failed: {e}")

        elif execution_mode == 'sequential':
            # Execute agents one after another
            for agent_id in agent_ids:
                result = self._execute_agent_by_id(agent_id, section_title, section_content, context_vars)
                if result:
                    results.append(result)
                    # Update context vars with latest result for next agent
                    context_vars['context'] = context_vars.get('context', '') + f"\n\nLatest Agent Output:\n{result.rules_extracted}\n\n"

        elif execution_mode == 'batched':
            # Execute in batches (not fully implemented, fallback to parallel)
            logger.warning(f"Batched execution mode not fully implemented for stage '{stage_name}', using parallel")
            return self._execute_stage({**stage, 'execution_mode': 'parallel'}, section_title, section_content, all_stage_outputs)

        logger.info(f"Stage '{stage_name}' completed: {len(results)} successful agent executions")
        return results

    def generate_multi_agent_test_plan(self,
                                     source_collections: List[str],
                                     source_doc_ids: List[str],
                                     doc_title: str = "Test Plan",
                                     agent_set_id: int = None,
                                     pipeline_id: str = None) -> FinalTestPlan:
        """
        Main entry point for multi-agent test plan generation

        Args:
            source_collections: List of ChromaDB collections to query
            source_doc_ids: List of specific document IDs to include
            doc_title: Title for the generated test plan
            agent_set_id: ID of agent set to use for orchestration (required)

        Raises:
            ValueError: If agent_set_id is None or invalid
        """
        logger.info("=== STARTING MULTI-AGENT TEST PLAN GENERATION ===")
        start_time = time.time()

        # Validate agent_set_id is provided
        if agent_set_id is None:
            raise ValueError("agent_set_id is required. Please select an agent set from the Agent Set Manager.")

        # Use provided pipeline_id or generate unique one
        if pipeline_id is None:
            pipeline_id = f"pipeline_{uuid.uuid4().hex[:12]}"

        # Load agent set configuration
        logger.info(f"Using agent set ID: {agent_set_id}")
        agent_set_config = self._load_agent_set_configuration(agent_set_id)
        if not agent_set_config:
            raise ValueError(f"Failed to load agent set {agent_set_id}. Agent set may not exist or is invalid.")

        try:
            # 0. Validate model availability and fallback to llama if needed
            self._maybe_fallback_to_llama(pipeline_id)

            # 1. Extract document sections from ChromaDB
            sections = self._extract_document_sections(source_collections, source_doc_ids)
            
            if not sections:
                logger.error("No sections extracted from ChromaDB")
                return self._create_fallback_test_plan(doc_title, pipeline_id)
            
            logger.info(f"Processing {len(sections)} sections with multi-agent pipeline")
            
            # 2. Initialize Redis pipeline for this run
            self._initialize_pipeline(pipeline_id, sections, doc_title)
            
            # 3. Mark pipeline as processing
            try:
                self._update_pipeline_metadata(pipeline_id, {
                    "status": "PROCESSING"
                })
                self.redis_client.zadd("pipeline:processing", {pipeline_id: time.time()})
            except Exception as e:
                logger.warning(f"Failed to mark pipeline processing: {e}")

            # 4. Deploy actor agents for each section (parallel processing)
            section_results = self._deploy_section_agents(pipeline_id, sections, agent_set_config)
            
            # 5. Deploy final critic agent to consolidate everything
            # If aborted, do not run final critic; return partial/aborted plan
            if self._is_aborted(pipeline_id):
                logger.warning(f"Pipeline {pipeline_id} aborted; skipping final critic")
                self._update_pipeline_metadata(pipeline_id, {
                    "status": "ABORTED",
                    "completed_at": datetime.now().isoformat(),
                })
                self.redis_client.zrem("pipeline:processing", pipeline_id)
                aborted_markdown = f"# {doc_title}\n\nProcess aborted. {len(section_results)} sections completed before abort."
                final_plan = FinalTestPlan(
                    title=doc_title,
                    pipeline_id=pipeline_id,
                    total_sections=len(section_results),
                    total_requirements=sum(len(r.test_procedures) for r in section_results),
                    total_test_procedures=sum(len(r.test_procedures) for r in section_results),
                    consolidated_markdown=aborted_markdown,
                    processing_status="ABORTED",
                    sections=section_results,
                )
                # If purge_on_abort flag set, purge all keys except abort flag
                try:
                    meta = self.redis_client.hgetall(f"pipeline:{pipeline_id}:meta") or {}
                    if meta.get("purge_on_abort") == "1":
                        self._purge_pipeline_keys(pipeline_id)
                except Exception as e:
                    logger.warning(f"Purge on abort failed: {e}")
            else:
                final_plan = self._deploy_final_critic_agent(pipeline_id, section_results, doc_title)
            
            # 6. Mark pipeline for retention (do not hard-delete so UI can view progress)
            self._cleanup_pipeline(pipeline_id)
            # Remove from processing set
            try:
                self.redis_client.zrem("pipeline:processing", pipeline_id)
            except Exception:
                pass
            
            elapsed_time = time.time() - start_time
            logger.info(f"Multi-agent test plan generation completed in {elapsed_time:.2f}s")
            
            return final_plan
            
        except Exception as e:
            logger.error(f"Multi-agent test plan generation failed: {e}")
            self._cleanup_pipeline(pipeline_id)
            try:
                self._update_pipeline_metadata(pipeline_id, {
                    "status": "FAILED",
                    "error": str(e)
                })
                self.redis_client.zrem("pipeline:processing", pipeline_id)
            except Exception:
                pass
            return self._create_fallback_test_plan(doc_title, pipeline_id)
    
    def _extract_document_sections(self, source_collections: List[str], source_doc_ids: List[str]) -> Dict[str, str]:
        """Extract sections from ChromaDB with robust reconstruction fallback.

        Strategy:
        - If explicit document IDs are provided, reconstruct full document(s) and split into natural sections.
        - Otherwise, group by metadata-based 'section_title' or page and combine chunks.
        """
        sections: Dict[str, str] = {}

        # 1) Preferred path: reconstruct by provided document IDs
        if source_doc_ids:
            for collection_name in source_collections:
                for doc_id in source_doc_ids:
                    try:
                        resp = requests.get(
                            f"{self.fastapi_url}/api/vectordb/documents/reconstruct/{doc_id}",
                            params={"collection_name": collection_name},
                            timeout=180,
                        )
                        if not resp.ok:
                            logger.warning(f"Reconstruct failed for doc_id={doc_id} in {collection_name}: {resp.status_code}")
                            continue
                        data = resp.json()
                        content = data.get("reconstructed_content") or ""
                        doc_name = data.get("document_name") or str(doc_id)
                        if content and len(content.strip()) > 50:
                            self._create_document_sections(doc_name, content, sections)
                            logger.info(f"Reconstructed {doc_name}: sections now {len(sections)}")
                    except Exception as e:
                        logger.error(f"Reconstruct error for {doc_id} in {collection_name}: {e}")

            if sections:
                logger.info(f"Extracted {len(sections)} sections via reconstruct path")
                return sections

        # 2) Fallback path: metadata grouping from the collection
        for collection_name in source_collections:
            try:
                response = requests.get(
                    f"{self.fastapi_url}/api/vectordb/documents",
                    params={"collection_name": collection_name},
                    timeout=60
                )

                if not response.ok:
                    logger.error(f"Failed to fetch from collection {collection_name}")
                    continue

                data = response.json()
                docs = data.get("documents", [])
                metas = data.get("metadatas", [])
                ids = data.get("ids", [])

                logger.info(f"Collection {collection_name} has {len(docs)} documents")

                # Group by document and section with improved metadata handling
                document_sections: Dict[str, List[str]] = {}
                for doc_id, doc, meta in zip(ids, docs, metas):
                    doc_name = (
                        (meta.get("document_name") or meta.get("filename") or meta.get("source"))
                        or doc_id
                    )

                    section_title = (
                        meta.get("section_title") or meta.get("heading") or meta.get("title")
                        or f"Section {meta.get('page_number', meta.get('page', 'Unknown'))}"
                    )

                    # Filter by explicit IDs if specified
                    if source_doc_ids:
                        if doc_id not in source_doc_ids and not any(str(x) in str(doc_name) for x in source_doc_ids):
                            continue

                    key = f"{doc_name} - {section_title}"
                    document_sections.setdefault(key, [])
                    document_sections[key].append(doc)

                # Combine chunks for each section
                for section_key, chunks in document_sections.items():
                    combined_content = "\n\n".join(chunk for chunk in chunks if chunk and chunk.strip())
                    if len(combined_content.strip()) > 100:  # Only substantial content
                        sections[section_key] = combined_content

            except Exception as e:
                logger.error(f"Error processing collection {collection_name}: {e}")

        logger.info(f"Extracted {len(sections)} sections for multi-agent processing (fallback path)")
        return sections

    def _clean_section_title(self, title: str) -> str:
        """
        Clean section title by removing PDF filenames and other artifacts.

        Args:
            title: Raw section title (e.g., "disr_ipv6_50.pdf - Introduction")

        Returns:
            Cleaned title (e.g., "Introduction")
        """
        import re

        # Remove PDF filename prefix (e.g., "disr_ipv6_50.pdf - ")
        title = re.sub(r'^[^\s]+\.pdf\s*-\s*', '', title)

        # Remove document name prefix (e.g., "Document Name - ")
        # but preserve section numbers
        if ' - ' in title and not re.match(r'^\d+\.', title.split(' - ')[0]):
            parts = title.split(' - ', 1)
            if len(parts) == 2 and not parts[0].replace('.', '').replace(' ', '').isdigit():
                title = parts[1]

        # Clean up extra whitespace
        title = ' '.join(title.split())

        return title.strip()

    def _normalize_section_content(self, content: str) -> str:
        """
        Normalize section content by removing duplicate section headings and adjusting heading levels.

        When critic agents generate content, they include their own ## headings which conflict
        with the numbered sections we add during assembly. This method:
        1. Removes the first ## heading (which duplicates our numbered section title)
        2. Keeps all other content including Dependencies, Conflicts, Test Procedures

        Args:
            content: Raw content from critic agent (synthesized_rules)

        Returns:
            Content with normalized heading structure
        """
        import re

        lines = content.strip().split('\n')
        output_lines = []
        first_h2_removed = False

        for line in lines:
            # Skip the first ## heading (duplicates our numbered section title)
            if line.strip().startswith('## ') and not first_h2_removed:
                first_h2_removed = True
                continue

            # Skip empty lines immediately after removed heading
            if not first_h2_removed and not line.strip():
                continue

            output_lines.append(line)

        return '\n'.join(output_lines)

    def _create_document_sections(self, doc_name: str, full_document: str, sections: Dict[str, str]):
        """Create logical sections from a reconstructed full document using natural headers.

        - Prefer numbered headers, ALL-CAPS, APPENDIX/CHAPTER/SECTION markers
        - Split very large sections into sub-blocks to keep units testable
        """
        natural_sections = self._extract_natural_sections(full_document)

        # Capture distinct APPENDIX blocks if present
        appendix_pattern = re.compile(r"^APPENDIX\s+[A-Z](?:\s*[-–]\s*.*)?$", re.MULTILINE)
        appendix_matches = list(appendix_pattern.finditer(full_document))
        if appendix_matches:
            for idx, m in enumerate(appendix_matches):
                start = m.start()
                end = appendix_matches[idx + 1].start() if idx + 1 < len(appendix_matches) else len(full_document)
                title = m.group(0).strip()
                body = full_document[start:end].strip()
                if title not in natural_sections:
                    natural_sections[title] = body

        if len(natural_sections) > 1:
            for section_title, section_content in natural_sections.items():
                section_key = f"{doc_name} - {section_title}"
                content = (section_content or "").strip()
                if len(content) > 8000:
                    subsections = self._split_large_section_for_testing(content, section_title)
                    for subsection_title, subsection_content in subsections.items():
                        sections[f"{section_key} - {subsection_title}"] = subsection_content
                else:
                    sections[section_key] = content
        else:
            # Fallback to size-based split if no natural sections found
            if len(full_document) > 6000:
                parts = self._split_large_section_for_testing(full_document, "Complete Document")
                for idx, (t, c) in enumerate(parts.items(), start=1):
                    sections[f"{doc_name} - Part {idx}"] = c
            else:
                sections[f"{doc_name} - Complete Document"] = (full_document or "").strip()

    def _extract_natural_sections(self, document_text: str) -> Dict[str, str]:
        """Heuristic extraction of natural sections from text"""
        sections: Dict[str, str] = {}
        lines = document_text.split('\n')
        current_section_title = "Introduction"
        current_section_content: List[str] = []

        for line in lines:
            line_clean = line.strip()
            is_section_header = False
            header_title = ""

            if line_clean:
                # Markdown headings from reconstruction (##, ###)
                if line_clean.startswith("## ") or line_clean.startswith("### "):
                    is_section_header = True
                    # Keep original heading text
                    header_title = line_clean[3:].strip() if line_clean.startswith("## ") else line_clean[4:].strip()
                # Numbered sections like 1., 1.1, 2.3.4 followed by ALL CAPS words
                elif re.match(r'^\d+(\.\d+)*\.?\s+[A-Z][A-Z\s]+', line_clean):
                    is_section_header = True
                    header_title = line_clean
                # ALL CAPS short headers
                elif (line_clean.isupper() and len(line_clean.split()) <= 8 and len(line_clean) > 5 and not line_clean.endswith('.') and not line_clean.startswith(('THE ', 'THIS ', 'THESE '))):
                    is_section_header = True
                    header_title = line_clean
                # APPENDIX/CHAPTER/SECTION/PART
                elif line_clean.startswith(('APPENDIX', 'CHAPTER', 'SECTION', 'PART')):
                    is_section_header = True
                    header_title = line_clean
                # Short keyword headers
                elif (any(k in line_clean.upper() for k in ['REQUIREMENTS','SPECIFICATIONS','PROCEDURES','TESTING','CONFIGURATION']) and len(line_clean.split()) <= 6 and not line_clean.endswith('.') and (line_clean.startswith(tuple('0123456789')) or line_clean.isupper())):
                    is_section_header = True
                    header_title = line_clean

            if is_section_header and current_section_content:
                sections[current_section_title] = '\n'.join(current_section_content)
                current_section_title = header_title
                current_section_content = [line]
            elif is_section_header and not current_section_content:
                current_section_title = header_title
                current_section_content = [line]
            else:
                current_section_content.append(line)

        if current_section_content:
            sections[current_section_title] = '\n'.join(current_section_content)

        return sections

    def _split_large_section_for_testing(self, content: str, section_title: str) -> Dict[str, str]:
        """Split large sections into smaller units for more focused processing."""
        subsections: Dict[str, str] = {}

        # Numbered subsections first (e.g., 4.1, 4.2)
        subsection_pattern = re.compile(r'^(\d+\.\d+(?:\.\d+)*)\s+(.+)$', re.MULTILINE)
        matches = list(subsection_pattern.finditer(content))
        if len(matches) > 1:
            for i, match in enumerate(matches):
                start = match.start()
                end = matches[i + 1].start() if i + 1 < len(matches) else len(content)
                subsection_num = match.group(1)
                subsection_name = match.group(2)[:50]
                subsection_content = content[start:end].strip()
                if len(subsection_content) > 500:
                    subsections[f"{subsection_num} {subsection_name}"] = subsection_content

        # Paragraph/block split fallback
        if not subsections:
            paragraphs = content.split('\n\n')
            current_block = ""
            block_num = 1
            for paragraph in paragraphs:
                if len(current_block + paragraph) > 5000:
                    if current_block.strip():
                        subsections[f"Block {block_num}"] = current_block.strip()
                        block_num += 1
                        current_block = paragraph + "\n\n"
                    else:
                        current_block += paragraph + "\n\n"
                else:
                    current_block += paragraph + "\n\n"
            if current_block.strip():
                subsections[f"Block {block_num}"] = current_block.strip()

        return subsections
    
    def _initialize_pipeline(self, pipeline_id: str, sections: Dict[str, str], doc_title: str):
        """Initialize Redis pipeline with sections and metadata"""
        # Get existing metadata (if created by API endpoint)
        existing_meta = self.redis_client.hgetall(f"pipeline:{pipeline_id}:meta") or {}

        # Update with initialization data (preserve existing fields like agent_set_name)
        now = datetime.now().isoformat()
        pipeline_data = {
            "pipeline_id": pipeline_id,
            "doc_title": doc_title,
            "status": "INITIALIZING",
            "total_sections": len(sections),
            "sections_processed": 0,
            "created_at": existing_meta.get("created_at", now),
            "last_updated_at": now,
            "progress_message": f"Initializing pipeline with {len(sections)} sections...",
            "actor_agents": len(self.actor_models)
        }

        # Preserve agent_set_name if it was already set
        if "agent_set_name" in existing_meta:
            pipeline_data["agent_set_name"] = existing_meta["agent_set_name"]

        # Store pipeline metadata (merge with existing)
        self.redis_client.hset(f"pipeline:{pipeline_id}:meta", mapping=pipeline_data)
        # Track recent pipelines for quick listing in UI
        try:
            now_ts = time.time()
            self.redis_client.zadd("pipeline:recent", {pipeline_id: now_ts})
        except Exception as e:
            logger.warning(f"Failed to zadd pipeline: {e}")
        
        # Store sections for processing
        for idx, (section_title, section_content) in enumerate(sections.items()):
            section_data = {
                "title": section_title,
                "content": section_content,
                "status": "PENDING",
                "index": idx
            }
            self.redis_client.hset(f"pipeline:{pipeline_id}:section:{idx}", mapping=section_data)
        
        # Initialize result queues
        self.redis_client.delete(f"pipeline:{pipeline_id}:actor_results")
        self.redis_client.delete(f"pipeline:{pipeline_id}:critic_results")
        
        logger.info(f"Pipeline {pipeline_id} initialized with {len(sections)} sections")
    
    def _deploy_section_agents(self, pipeline_id: str, sections: Dict[str, str], agent_set_config: Optional[Dict[str, Any]] = None) -> List[CriticResult]:
        """
        Deploy multiple agents per section with Redis coordination

        Args:
            pipeline_id: Unique pipeline identifier
            sections: Dictionary of section title -> content
            agent_set_config: Optional agent set configuration. If None, uses default orchestration.

        Returns:
            List of CriticResult objects for each section
        """
        if agent_set_config:
            logger.info(f"Deploying agents for {len(sections)} sections using agent set configuration")
            logger.info(f"Agent set has {len(agent_set_config.get('stages', []))} stages")
        else:
            logger.info(f"Deploying agents for {len(sections)} sections using default orchestration")
        
        section_results = []
        
        # Process each section with multiple actor agents + critic
        with ThreadPoolExecutor(max_workers=8) as executor:
            future_to_section = {}
            
            for idx, (section_title, section_content) in enumerate(sections.items()):
                # Respect abort flag: stop submitting new work
                if self._is_aborted(pipeline_id):
                    logger.warning(f"Abort requested for pipeline {pipeline_id}; stopping new submissions at section {idx}")
                    # Mark remaining sections as aborted
                    self.redis_client.hset(f"pipeline:{pipeline_id}:section:{idx}", "status", "ABORTED")
                    break
                future = executor.submit(
                    self._process_section_with_multi_agents,
                    pipeline_id, idx, section_title, section_content, agent_set_config
                )
                future_to_section[future] = section_title
            
            # Collect results as they complete
            for future in as_completed(future_to_section):
                section_title = future_to_section[future]
                try:
                    critic_result = future.result(timeout=300)  # 5 minute timeout per section
                    if critic_result:
                        section_results.append(critic_result)
                        logger.info(f"Section completed: {section_title}")
                    else:
                        logger.warning(f"Section failed or aborted: {section_title}")
                except Exception as e:
                    logger.error(f"Section processing error for '{section_title}': {e}")
        
        logger.info(f"Completed processing {len(section_results)} sections")
        return section_results
    
    def _process_section_with_multi_agents(self,
                                         pipeline_id: str,
                                         section_idx: int,
                                         section_title: str,
                                         section_content: str,
                                         agent_set_config: Dict[str, Any]) -> Optional[CriticResult]:
        """
        Process a single section with agents from agent set configuration

        Args:
            pipeline_id: Unique pipeline identifier
            section_idx: Section index number
            section_title: Title of the section
            section_content: Content of the section
            agent_set_config: Agent set configuration (required)

        Returns:
            CriticResult or None if aborted/failed
        """

        # Respect abort flag early
        if self._is_aborted(pipeline_id):
            self.redis_client.hset(f"pipeline:{pipeline_id}:section:{section_idx}", "status", "ABORTED")
            return None

        # Update section status
        self.redis_client.hset(f"pipeline:{pipeline_id}:section:{section_idx}", "status", "PROCESSING")

        try:
            # Execute agents based on agent_set_config
            if 'stages' not in agent_set_config:
                raise ValueError("Agent set configuration must contain 'stages'")

            if agent_set_config and 'stages' in agent_set_config:
                logger.info(f"Using custom agent set orchestration with {len(agent_set_config['stages'])} stages")

                # Execute stages in sequence, passing context between them
                all_stage_results = []
                all_stage_outputs = {}  # Track outputs by stage name for context building

                for stage_idx, stage in enumerate(agent_set_config['stages']):
                    stage_name = stage.get('stage_name', f'stage_{stage_idx}')
                    logger.info(f"Executing stage {stage_idx + 1}/{len(agent_set_config['stages'])}: {stage_name}")

                    stage_results = self._execute_stage(stage, section_title, section_content, all_stage_outputs)
                    all_stage_results.extend(stage_results)

                    # Track this stage's outputs by name for future stages
                    all_stage_outputs[stage_name] = stage_results

                # Use the last stage's results as final actor results
                actor_results = all_stage_results if all_stage_results else []

                # Store all stage results in Redis
                for result in actor_results:
                    result_key = f"pipeline:{pipeline_id}:actor:{section_idx}:{result.agent_id}"
                    result_data = {
                        "agent_id": result.agent_id,
                        "model_name": result.model_name,
                        "section_title": result.section_title,
                        "rules_extracted": result.rules_extracted,
                        "processing_time": result.processing_time
                    }
                    self.redis_client.hset(result_key, mapping=result_data)

                # For agent sets, use the final stage output as the synthesized result
                # Create a CriticResult from the final outputs
                if actor_results:
                    final_output = "\n\n".join([r.rules_extracted for r in actor_results])
                    critic_result = CriticResult(
                        section_title=section_title,
                        synthesized_rules=final_output,
                        dependencies=[],
                        conflicts=[],
                        test_procedures=[],  # TODO: Parse test procedures from output
                        actor_count=len(actor_results)
                    )
                else:
                    logger.warning(f"No results from agent set stages for section: {section_title}")
                    critic_result = None
            
            # Store critic result in Redis
            if critic_result:
                critic_key = f"pipeline:{pipeline_id}:critic:{section_idx}"
                critic_data = {
                    "section_title": critic_result.section_title,
                    "synthesized_rules": critic_result.synthesized_rules,
                    "dependencies": json.dumps(critic_result.dependencies),
                    "conflicts": json.dumps(critic_result.conflicts),
                    "test_procedures": json.dumps(critic_result.test_procedures),
                    "actor_count": critic_result.actor_count
                }
                self.redis_client.hset(critic_key, mapping=critic_data)
                
                # Update section status
                self.redis_client.hset(f"pipeline:{pipeline_id}:section:{section_idx}", "status", "COMPLETED")
                # Increment processed counter on meta
                try:
                    self.redis_client.hincrby(f"pipeline:{pipeline_id}:meta", "sections_processed", 1)
                except Exception:
                    pass
                
                return critic_result
            
        except Exception as e:
            logger.error(f"Error processing section {section_title}: {e}")
            self.redis_client.hset(f"pipeline:{pipeline_id}:section:{section_idx}", "status", "FAILED")
        
        return None

    def _is_aborted(self, pipeline_id: str) -> bool:
        try:
            return self.redis_client.get(f"pipeline:{pipeline_id}:abort") == "1"
        except Exception:
            return False

    def _update_pipeline_metadata(self, pipeline_id: str, updates: dict):
        """Update pipeline metadata with automatic timestamp update"""
        updates["last_updated_at"] = datetime.now().isoformat()
        self.redis_client.hset(f"pipeline:{pipeline_id}:meta", mapping=updates)

    def _run_actor_agents(self, section_title: str, section_content: str) -> List[ActorResult]:
        """Run multiple GPT-4 actor agents in parallel for a section"""
        actor_results = []
        
        with ThreadPoolExecutor(max_workers=len(self.actor_models)) as executor:
            futures = []
            
            for idx, model in enumerate(self.actor_models):
                agent_id = f"actor_{idx}_{uuid.uuid4().hex[:8]}"
                future = executor.submit(
                    self._run_single_actor, agent_id, model, section_title, section_content
                )
                futures.append(future)
            
            # Collect actor results
            for future in as_completed(futures):
                try:
                    result = future.result(timeout=180)  # 3 minute timeout per GPT-4 actor
                    if result:
                        actor_results.append(result)
                except Exception as e:
                    logger.error(f"Actor agent failed: {e}")
        
        logger.info(f"Completed {len(actor_results)} GPT-4 actor agents for section: {section_title}")
        return actor_results
    
    def _run_single_actor(self, agent_id: str, model: str, section_title: str, section_content: str) -> Optional[ActorResult]:
        """
        Run a single actor agent with database-backed prompts.

        Prompts are loaded from database via agent_registry if available,
        otherwise falls back to hardcoded prompts.
        """
        start_time = time.time()

        try:
            # Try to get prompts from database first
            # agent_id format: "actor_0_b4e879d5" where 0 is the index
            parts = agent_id.split('_')
            actor_index = int(parts[1]) if len(parts) >= 2 and parts[1].isdigit() else 0
            agent_prompts = self.agent_registry.get_actor_agent_prompts(actor_index)

            if agent_prompts:
                # Use database-backed prompts
                system_prompt = agent_prompts['system_prompt']
                user_prompt_template = agent_prompts['user_prompt_template']

                # Replace placeholders
                user_prompt = user_prompt_template.replace('{section_title}', section_title)
                user_prompt = user_prompt.replace('{section_content}', section_content)

                # Combine system and user prompts
                prompt = f"{system_prompt}\n\n{user_prompt}"

                logger.debug(f"Actor {agent_id}: Using database-backed prompts")
            else:
                # Fallback to hardcoded prompt
                prompt = f"""You are a compliance and test planning expert.

Analyze the following section of a military standard and extract EVERY possible testable rule, specification, constraint, or requirement. Rules MUST be extremely detailed, explicit, and step-by-step, and should include measurable criteria, acceptable ranges, and referenced figures or tables if mentioned.

For ambiguous or implicit requirements, describe a specific test strategy.
Generate a short, content-based TITLE for this section (do not use page numbers).

ABSOLUTELY DO NOT REPEAT, DUPLICATE, OR PARAPHRASE THE SAME RULE OR LINE. Each requirement, dependency, and test step must appear ONCE ONLY.

Organize your output as follows, using markdown headings and bolded text:

## [Section Title]
**Dependencies:**
- List detailed dependencies as explicit tests, if any.

**Conflicts:**
- List detected or possible conflicts and provide recommendations or mitigation steps.

**Test Rules:**
1. (Very detailed, step-by-step numbered test rules)

Section Name: {section_title}

Section Text:
{section_content}

---
If you find truly nothing testable, reply: 'No testable rules in this section.'
"""
                logger.debug(f"Actor {agent_id}: Using hardcoded fallback prompts")

            response = self.llm_service.query_direct(
                model_name=model,
                query=prompt
            )[0]

            processing_time = time.time() - start_time

            return ActorResult(
                agent_id=agent_id,
                model_name=model,
                section_title=section_title,
                rules_extracted=response,
                processing_time=processing_time
            )

        except Exception as e:
            logger.error(f"Actor {agent_id} failed: {e}")
            return None
    
    def _run_critic_agent(self, section_title: str, section_content: str, actor_results: List[ActorResult]) -> Optional[CriticResult]:
        """
        Run critic agent with database-backed prompts to synthesize actor outputs.

        Prompts are loaded from database via agent_registry if available,
        otherwise falls back to hardcoded prompts.
        """

        if not actor_results:
            logger.warning(f"No actor results to critique for section: {section_title}")
            return None

        try:
            # Prepare actor outputs for critic
            actor_outputs_text = ""
            for result in actor_results:
                actor_outputs_text += f"\n\nModel {result.model_name} ({result.agent_id}):\n{result.rules_extracted}\n{'='*40}"

            # Try to get prompts from database first
            critic_prompts = self.agent_registry.get_critic_agent_prompts()

            if critic_prompts:
                # Use database-backed prompts
                system_prompt = critic_prompts['system_prompt']
                user_prompt_template = critic_prompts['user_prompt_template']

                # Build context for prompt replacement
                # The template should have placeholders like {section_title}, {section_content}, {actor_outputs}
                user_prompt = user_prompt_template.replace('{section_title}', section_title)
                user_prompt = user_prompt.replace('{section_content}', section_content)
                user_prompt = user_prompt.replace('{actor_outputs}', actor_outputs_text)
                user_prompt = user_prompt.replace('{actor_count}', str(len(actor_results)))

                # Combine system and user prompts
                prompt = f"{system_prompt}\n\n{user_prompt}"

                logger.debug("Critic: Using database-backed prompts")
            else:
                # Fallback to hardcoded prompt
                prompt = f"""You are a senior test planning reviewer (Critic AI).

Given the following section and rules extracted by several different GPT-4 models, do the following:
1. Carefully review and compare the provided rule sets.
2. Synthesize a SINGLE, detailed and explicit set of testable rules.
3. Eliminate redundancies, correct errors, and ensure all requirements are present.
4. Ensure the final test plan is step-by-step, detailed, and well organized.

NEVER simply combine all lines verbatim—synthesize, deduplicate, and streamline the content into a concise, non-repetitive format. If a rule, step, or line has the same or similar meaning as another, KEEP ONLY ONE.

Present your result in markdown format with these headings: '## [Section Title]', '**Dependencies:**', '**Conflicts:**', '**Test Rules:**'

Section Name: {section_title}

Section Text:
{section_content}

---
Actor Outputs from {len(actor_results)} GPT-4 models:
{actor_outputs_text}
"""
                logger.debug("Critic: Using hardcoded fallback prompts")

            response = self.llm_service.query_direct(
                model_name=self.critic_model,
                query=prompt
            )[0]
            
            # Apply deduplication (from notebook)
            deduplicated_response = self._deduplicate_markdown(response)
            
            # Extract structured data from critic response
            dependencies = self._extract_dependencies_from_markdown(deduplicated_response)
            conflicts = self._extract_conflicts_from_markdown(deduplicated_response)
            test_procedures = self._extract_test_procedures_from_markdown(deduplicated_response)
            
            return CriticResult(
                section_title=section_title,
                synthesized_rules=deduplicated_response,
                dependencies=dependencies,
                conflicts=conflicts,
                test_procedures=test_procedures,
                actor_count=len(actor_results)
            )
            
        except Exception as e:
            logger.error(f"GPT-4 Critic agent failed for section {section_title}: {e}")
            return None
    
    def _deploy_final_critic_agent(self, pipeline_id: str, section_results: List[CriticResult], doc_title: str) -> FinalTestPlan:
        """Deploy final GPT-4 critic agent to consolidate all sections"""
        logger.info("Deploying final GPT-4 critic agent for consolidation")

        try:
            # Prepare all section results for final critic
            sections_summary = []
            all_sections_content = ""

            for result in section_results:
                sections_summary.append({
                    "title": result.section_title,
                    "dependencies_count": len(result.dependencies),
                    "conflicts_count": len(result.conflicts),
                    "test_procedures_count": len(result.test_procedures),
                    "actor_count": result.actor_count
                })

                all_sections_content += f"\n\n## {result.section_title}\n"
                all_sections_content += result.synthesized_rules
                all_sections_content += "\n" + "="*60

            # Check if content is too large for context window
            # Rough estimate: 1 token ≈ 4 characters
            estimated_tokens = len(all_sections_content) / 4
            max_context_tokens = 7000  # Conservative limit for gpt-4 (8192 total - overhead)

            if estimated_tokens > max_context_tokens:
                logger.warning(f"Content too large for final critic ({estimated_tokens:.0f} tokens estimated). Skipping final consolidation and assembling directly from sections.")

                # Assemble document directly from section results without LLM consolidation
                final_markdown = f"# {doc_title}\n\n"

                # Note: Pandoc will auto-generate TOC with --toc flag, so we don't manually create one
                # This prevents numbering conflicts

                final_markdown += "\n---\n\n"

                # Add sections with cleaned titles
                for idx, result in enumerate(section_results, 1):
                    clean_title = self._clean_section_title(result.section_title)
                    final_markdown += f"## {idx}. {clean_title}\n\n"
                    # Normalize heading levels: strip first ## heading from synthesized_rules if present,
                    # and downgrade remaining headings (## → ###, ### → ####)
                    content = self._normalize_section_content(result.synthesized_rules)
                    final_markdown += content
                    final_markdown += "\n\n---\n\n"

                final_markdown += "## Summary & Recommendations\n\n"
                final_markdown += f"This test plan covers {len(section_results)} sections with comprehensive test procedures and requirements.\n\n"
                final_markdown += "**Note**: Document assembled directly from section results due to size. Each section has been individually synthesized by multiple AI agents and reviewed by a critic agent.\n"

            else:
                # Content fits in context window - use final critic for consolidation
                logger.info(f"Content size acceptable ({estimated_tokens:.0f} tokens estimated). Running final critic consolidation.")

                # Final critic prompt (based on notebook's final_test_plan_docx logic)
                prompt = f"""You are a final Critic AI creating a comprehensive military/technical standard test plan.

Given the following detailed section-by-section test procedures (each synthesized from multiple GPT-4 actor agents), combine them into a single, fully ordered, professional test plan document:

DOCUMENT STRUCTURE:
1. Title Page: '{doc_title}'
2. Executive Summary: Brief overview of test scope and objectives
3. For each section: Include the detailed TEST PROCEDURES (not requirements tables) with CLEAN section titles
4. Summary & Recommendations: Synthesize critical points and compliance strategy

NOTE: Do NOT manually create a "Table of Contents" - Pandoc will auto-generate it from section headings.

CRITICAL REQUIREMENTS:
- PRESERVE ORIGINAL REQUIREMENT IDs from source document (e.g., 4.2.1, REQ-01, etc.)
- DO NOT generate requirements tables - include TEST PROCEDURES only
- Each test procedure must have: Requirement ID, Objective, Setup, Steps, Expected Results, Pass/Fail Criteria
- Use hierarchical numbering for TEST PLAN sections: 1, 2, 3, ... (sub-sections: 1.1, 1.2, 2.1, etc.)
- DO NOT include Table of Contents from the source documents
- CLEAN section titles: Remove PDF filenames (e.g., "disr_ipv6_50.pdf - Introduction" becomes "Introduction")
- Ensure test procedures are executable by engineers
- Use BULLET POINTS (-, *) for all lists within test procedures (not numbered lists 1., 2., 3.) to prevent enumeration conflicts

FORMAT:
- Only main test plan section titles in TOC (not 'Dependencies', 'Test Rules', etc.)
- Preserve markdown formatting for DOCX conversion
- Ensure continuous numbering with no gaps

SECTIONS SUMMARY:
{json.dumps(sections_summary, indent=2)}

DETAILED SECTIONS:
{all_sections_content}

Create a comprehensive markdown document that consolidates all {len(section_results)} sections into a cohesive test plan.
"""

                response = self.llm_service.query_direct(
                    model_name=self.final_critic_model,
                    query=prompt
                )[0]

                final_markdown = response

            # Apply final deduplication
            final_markdown = self._final_global_deduplicate(final_markdown)

            # REMOVED: _add_structured_tables() - this was generating unwanted requirements tables
            # The actor/critic agents now generate proper test procedures with original requirement IDs
            # No need for post-processing table generation

            # Calculate totals
            total_requirements = sum(len(result.test_procedures) for result in section_results)
            total_test_procedures = total_requirements  # Each requirement becomes a test procedure
            
            # Store final result in Redis
            final_result_key = f"pipeline:{pipeline_id}:final_result"
            final_data = {
                "title": doc_title,
                "consolidated_markdown": final_markdown,
                "total_sections": len(section_results),
                "total_requirements": total_requirements,
                "total_test_procedures": total_test_procedures,
                "processing_status": "COMPLETED",
                "completed_at": datetime.now().isoformat()
            }
            self.redis_client.hset(final_result_key, mapping=final_data)
            # Update pipeline meta processing status (NOT completed - that's done by background task)
            try:
                self._update_pipeline_metadata(pipeline_id, {
                    "status": "processing",  # Keep as processing; background task will set to "completed"
                    "progress_message": "Test plan generated, preparing export...",
                    "last_updated_at": datetime.now().isoformat()
                })
                self.redis_client.zadd("pipeline:recent", {pipeline_id: time.time()})
            except Exception as e:
                logger.warning(f"Failed to update pipeline meta: {e}")
            
            return FinalTestPlan(
                title=doc_title,
                pipeline_id=pipeline_id,
                total_sections=len(section_results),
                total_requirements=total_requirements,
                total_test_procedures=total_test_procedures,
                consolidated_markdown=final_markdown,
                processing_status="COMPLETED",
                sections=section_results
            )
            
        except Exception as e:
            logger.error(f"Final GPT-4 critic agent failed: {e}")
            return FinalTestPlan(
                title=doc_title,
                pipeline_id=pipeline_id,
                total_sections=len(section_results),
                total_requirements=0,
                total_test_procedures=0,
                consolidated_markdown=f"# {doc_title}\n\nFinal critic agent failed: {str(e)}",
                processing_status="FAILED",
                sections=section_results
            )
    
    def _deduplicate_markdown(self, text: str) -> str:
        """Deduplicate sentences within markdown sections (from notebook)"""
        output = []
        section_boundary = lambda l: l.startswith("## ") or (l.startswith("**") and l.endswith("**"))
        
        def process_block(block):
            local_seen = set()
            for sentence in re.split(r'(?<=[.!?]) +', block):
                sent = sentence.strip()
                norm = re.sub(r'\s+', ' ', sent.lower())
                if not sent or norm in local_seen:
                    continue
                output.append(sent)
                local_seen.add(norm)
        
        current_block = []
        for line in text.split('\n'):
            if section_boundary(line):
                process_block(' '.join(current_block))
                current_block = []
                output.append(line)
            elif line.strip() == "":
                process_block(' '.join(current_block))
                current_block = []
                output.append(line)
            else:
                current_block.append(line.strip())
        
        process_block(' '.join(current_block))
        return '\n'.join(output)
    
    def _final_global_deduplicate(self, text: str) -> str:
        """Remove duplicate lines globally (from notebook)"""
        seen = set()
        out = []
        
        for line in text.split('\n'):
            sentences = re.split(r'(?<=[.!?]) +', line) if len(line) > 120 else [line]
            unique_sentences = []
            
            for s in sentences:
                norm = re.sub(r'\s+', ' ', s.strip().lower())
                if norm and norm not in seen:
                    unique_sentences.append(s)
                    seen.add(norm)
                elif not s.strip():
                    unique_sentences.append(s)
            
            joined = ' '.join(unique_sentences).strip()
            if joined or not line.strip():
                out.append(joined)
        
        return '\n'.join(out)

    def _add_structured_tables(self, markdown: str) -> str:
        """
        Add structured Requirements and Test Procedures tables to the test plan.

        Extracts requirements and test procedures from narrative text and formats them
        into structured markdown tables with proper headers and IDs.

        Args:
            markdown: The consolidated test plan markdown

        Returns:
            Enhanced markdown with structured tables
        """
        import re
        from typing import List, Dict, Any

        logger.info("Adding structured Requirements and Test Procedures tables...")

        # Split into sections
        sections = re.split(r'\n(?=#+\s+\d+\.)', markdown)

        enhanced_sections = []
        req_counter = 1
        test_counter = 1

        for section in sections:
            if not section.strip():
                continue

            # Extract section header
            lines = section.split('\n')
            section_header = lines[0] if lines else ""
            section_content = '\n'.join(lines[1:]) if len(lines) > 1 else ""

            # Extract requirements from section
            requirements = self._extract_requirements_from_section(section_content, req_counter)
            req_counter += len(requirements)

            # Extract test procedures from section
            test_procedures = self._extract_test_procedures_from_section(section_content, test_counter)
            test_counter += len(test_procedures)

            # Build enhanced section
            enhanced_section = section_header + '\n\n'

            # Add Requirements table if any found
            if requirements:
                enhanced_section += "### Requirements\n\n"
                enhanced_section += "| Req ID | Requirement Description | Source | Priority | Testable |\n"
                enhanced_section += "|--------|------------------------|--------|----------|----------|\n"
                for req in requirements:
                    enhanced_section += f"| {req['id']} | {req['description'][:100]} | {req['source']} | {req['priority']} | {req['testable']} |\n"
                enhanced_section += "\n"

            # Add original section content
            enhanced_section += section_content + '\n\n'

            # Add Test Procedures table if any found
            if test_procedures:
                enhanced_section += "### Test Procedures\n\n"
                enhanced_section += "| Test ID | Test Description | Steps | Expected Result | Acceptance Criteria | Req ID |\n"
                enhanced_section += "|---------|-----------------|-------|-----------------|---------------------|--------|\n"
                for test in test_procedures:
                    steps = '; '.join(test['steps'][:3]) if test['steps'] else 'N/A'
                    enhanced_section += f"| {test['id']} | {test['description'][:80]} | {steps[:60]} | {test['expected_result'][:60]} | {test['acceptance_criteria'][:60]} | {test['req_id']} |\n"
                enhanced_section += "\n"

            enhanced_sections.append(enhanced_section)

        result = '\n'.join(enhanced_sections)

        # Add summary tables at the beginning
        all_requirements = []
        all_tests = []
        req_counter = 1
        test_counter = 1

        for section in sections:
            if not section.strip():
                continue
            lines = section.split('\n')
            section_content = '\n'.join(lines[1:]) if len(lines) > 1 else ""

            reqs = self._extract_requirements_from_section(section_content, req_counter)
            req_counter += len(reqs)
            all_requirements.extend(reqs)

            tests = self._extract_test_procedures_from_section(section_content, test_counter)
            test_counter += len(tests)
            all_tests.extend(tests)

        # Prepend summary
        summary = "# Requirements and Test Procedures Summary\n\n"
        summary += f"**Total Requirements:** {len(all_requirements)}\n\n"
        summary += f"**Total Test Procedures:** {len(all_tests)}\n\n"
        summary += "---\n\n"

        logger.info(f"Added {len(all_requirements)} requirements and {len(all_tests)} test procedures to tables")

        return summary + result

    def _extract_requirements_from_section(self, content: str, start_id: int) -> List[Dict[str, Any]]:
        """Extract requirements from section content"""
        requirements = []

        # Pattern matching for requirement keywords
        requirement_patterns = [
            r'(?i)(must|shall|should|required?)\s+(.{20,200}?)(?:\.|;|\n)',
            r'(?i)requirement:?\s*(.{20,200}?)(?:\.|;|\n)',
            r'(?i)the\s+(?:system|device|implementation|software)\s+(must|shall|should)\s+(.{20,200}?)(?:\.|;|\n)',
        ]

        for pattern in requirement_patterns:
            matches = re.finditer(pattern, content, re.IGNORECASE | re.MULTILINE)
            for match in matches:
                if len(match.groups()) >= 2:
                    desc = match.group(2).strip() if len(match.group(2)) > len(match.group(1)) else match.group(1).strip()
                else:
                    desc = match.group(1).strip()

                # Clean up description
                desc = desc.replace('\n', ' ').strip()
                if len(desc) < 20 or len(desc) > 300:
                    continue

                # Check if already added
                if any(req['description'] == desc for req in requirements):
                    continue

                # Determine priority based on keywords
                priority = "High" if "must" in match.group(0).lower() or "shall" in match.group(0).lower() else "Medium"

                requirements.append({
                    'id': f'REQ-{start_id + len(requirements):03d}',
                    'description': desc,
                    'source': 'Section Analysis',
                    'priority': priority,
                    'testable': 'Yes'
                })

                if len(requirements) >= 10:  # Limit per section
                    break

        return requirements

    def _extract_test_procedures_from_section(self, content: str, start_id: int) -> List[Dict[str, Any]]:
        """Extract test procedures from section content"""
        test_procedures = []

        # Pattern matching for test procedures
        test_patterns = [
            r'(?i)test(?:\s+procedure)?:?\s*(.{20,200}?)(?:\.|;|\n)',
            r'(?i)verify(?:\s+that)?\s+(.{20,200}?)(?:\.|;|\n)',
            r'(?i)(?:check|confirm|ensure|validate)\s+(?:that\s+)?(.{20,200}?)(?:\.|;|\n)',
            r'\d+\.\s+(.{20,200}?)(?:\.|;|\n)',  # Numbered lists
        ]

        for pattern in test_patterns:
            matches = re.finditer(pattern, content, re.IGNORECASE | re.MULTILINE)
            for match in matches:
                desc = match.group(1).strip()

                # Clean up description
                desc = desc.replace('\n', ' ').strip()
                if len(desc) < 15 or len(desc) > 300:
                    continue

                # Check if already added
                if any(test['description'] == desc for test in test_procedures):
                    continue

                # Extract steps if bulleted/numbered list follows
                steps = []
                # Simple step extraction (can be enhanced)
                step_matches = re.findall(r'(?:^|\n)\s*[-•]\s*(.+?)(?=\n|$)', content[match.end():match.end()+500])
                steps = [s.strip() for s in step_matches[:5]]

                test_procedures.append({
                    'id': f'TC-{start_id + len(test_procedures):03d}',
                    'description': desc,
                    'steps': steps if steps else ['Setup test environment', 'Execute test', 'Verify results'],
                    'expected_result': 'Test passes all criteria',
                    'acceptance_criteria': 'All requirements met',
                    'req_id': f'REQ-{start_id:03d}'  # Link to first requirement in section
                })

                if len(test_procedures) >= 10:  # Limit per section
                    break

        return test_procedures

    def _extract_dependencies_from_markdown(self, markdown: str) -> List[str]:
        """Extract dependencies from markdown format"""
        dependencies = []
        in_dependencies = False
        
        for line in markdown.split('\n'):
            if line.strip().startswith('**Dependencies:**'):
                in_dependencies = True
                continue
            elif line.strip().startswith('**') and in_dependencies:
                break
            elif in_dependencies and line.strip().startswith('- '):
                dependencies.append(line.strip()[2:])
        
        return dependencies
    
    def _extract_conflicts_from_markdown(self, markdown: str) -> List[str]:
        """Extract conflicts from markdown format"""
        conflicts = []
        in_conflicts = False
        
        for line in markdown.split('\n'):
            if line.strip().startswith('**Conflicts:**'):
                in_conflicts = True
                continue
            elif line.strip().startswith('**') and in_conflicts:
                break
            elif in_conflicts and line.strip().startswith('- '):
                conflicts.append(line.strip()[2:])
        
        return conflicts
    
    def _extract_test_procedures_from_markdown(self, markdown: str) -> List[Dict[str, Any]]:
        """Extract test procedures from markdown format"""
        procedures = []
        in_test_rules = False
        
        for line in markdown.split('\n'):
            if line.strip().startswith('**Test Rules:**'):
                in_test_rules = True
                continue
            elif line.strip().startswith('**') and in_test_rules:
                break
            elif in_test_rules and re.match(r'^\d+\.', line.strip()):
                procedures.append({
                    "id": f"test_{len(procedures)+1}",
                    "description": line.strip(),
                    "type": "functional"
                })
        
        return procedures
    
    def _cleanup_pipeline(self, pipeline_id: str):
        """Mark pipeline keys with an expiration instead of hard deletion so UI can inspect later."""
        try:
            pattern = f"pipeline:{pipeline_id}:*"
            keys = self.redis_client.keys(pattern)
            for k in keys:
                try:
                    self.redis_client.expire(k, self.pipeline_ttl_seconds)
                except Exception:
                    pass
            # Also keep the meta record updated and expiring
            self.redis_client.expire(f"pipeline:{pipeline_id}:meta", self.pipeline_ttl_seconds)
            logger.info(f"Retained {len(keys)} Redis keys for pipeline {pipeline_id} with TTL={self.pipeline_ttl_seconds}s")
        except Exception as e:
            logger.error(f"Error retaining pipeline {pipeline_id}: {e}")

    def _purge_pipeline_keys(self, pipeline_id: str):
        try:
            pattern = f"pipeline:{pipeline_id}:*"
            keys = self.redis_client.keys(pattern) or []
            keep_key = f"pipeline:{pipeline_id}:abort"
            to_delete = [k for k in keys if k != keep_key]
            if to_delete:
                self.redis_client.delete(*to_delete)
            # Remove from index sets
            try:
                self.redis_client.zrem("pipeline:recent", pipeline_id)
                self.redis_client.zrem("pipeline:processing", pipeline_id)
            except Exception:
                pass
            logger.info(f"Purged pipeline {pipeline_id} keys: {len(to_delete)} deleted")
        except Exception as e:
            logger.error(f"Error purging pipeline {pipeline_id}: {e}")

    # ===========================
    # Model availability + fallback
    # ===========================
    def _maybe_fallback_to_llama(self, pipeline_id: str):
        """If OpenAI models are configured but unavailable or quota-exceeded, switch to llama models.

        Strategy:
        - If any actor/critic model starts with 'gpt' but OPEN_AI_API_KEY is missing -> fallback to llama
        - Else, attempt a very small direct query with critic model; on exception -> fallback to llama
        - Record fallback in Redis pipeline meta
        """
        try:
            needs_openai = any(str(m).lower().startswith("gpt") for m in (self.actor_models + [self.critic_model, self.final_critic_model]))
            if not needs_openai:
                return

            openai_key = os.getenv("OPENAI_API_KEY")
            if not openai_key:
                self._apply_llama_fallback(pipeline_id, reason="missing_openai_key")
                return

            # Quick probe with critic model
            try:
                # Keep it very short; do not log history
                _ = self.llm_service.query_direct(self.critic_model, "Return OK.", session_id=None, log_history=False)
                # If success, keep current models
                return
            except Exception as e:
                # Any failure here triggers fallback
                self._apply_llama_fallback(pipeline_id, reason=f"probe_failed: {str(e)[:120]}")
        except Exception as e:
            # Never fail generation due to fallback logic
            logger.warning(f"Model fallback check error: {e}")

    def _apply_llama_fallback(self, pipeline_id: str, reason: str = "unavailable"):
        try:
            # Use Llama 3.1 8B via Ollama for local CPU-based fallback
            fallback_model = "llama3.1:8b"
            logger.warning(f"Falling back to {fallback_model} models due to: {reason}")
            self.actor_models = [fallback_model for _ in self.actor_models]
            self.critic_model = fallback_model
            self.final_critic_model = fallback_model
            # Record in Redis
            self._update_pipeline_metadata(pipeline_id, {
                "model_fallback": fallback_model,
                "fallback_reason": reason,
                "original_actor_models": json.dumps(self._original_actor_models),
                "original_critic_model": self._original_critic_model,
                "actor_agents": len(self.actor_models),
            })
        except Exception as e:
            logger.warning(f"Failed to record model fallback: {e}")
    
    def _create_fallback_test_plan(self, doc_title: str, pipeline_id: str) -> FinalTestPlan:
        """Create fallback test plan"""
        fallback_markdown = f"""# {doc_title}

## Fallback Mode Notice
This test plan was generated in fallback mode due to section extraction issues.

## Multi-Agent Architecture
- 3x GPT-4 Actor Agents per section
- 1x GPT-4 Critic Agent per section  
- 1x GPT-4 Final Critic Agent for consolidation
- Redis pipeline for scaling

## Next Steps
1. Fix ChromaDB section extraction
2. Verify collection names and document IDs
3. Re-run multi-agent pipeline with GPT-4 agents
"""
        
        return FinalTestPlan(
            title=doc_title,
            pipeline_id=pipeline_id,
            total_sections=0,
            total_requirements=0,
            total_test_procedures=0,
            consolidated_markdown=fallback_markdown,
            processing_status="FALLBACK",
            sections=[]
        )
    
    def export_to_word(self, test_plan: FinalTestPlan) -> str:
        """Export test plan to Word document"""
        doc = Document()
        doc.add_heading(test_plan.title, level=0)
        
        # Add multi-agent architecture info
        doc.add_heading('Multi-Agent Architecture', level=2)
        doc.add_paragraph("This test plan was generated using:")
        # Use Word list formatting only; do not include literal bullet characters
        doc.add_paragraph("3x GPT-4 Actor Agents per section", style='List Bullet')
        doc.add_paragraph("1x GPT-4 Critic Agent per section for synthesis", style='List Bullet')  
        doc.add_paragraph("1x GPT-4 Final Critic Agent for consolidation", style='List Bullet')
        doc.add_paragraph("Redis pipeline for scalable processing", style='List Bullet')
        
        # Convert markdown to Word
        self._convert_markdown_to_docx(test_plan.consolidated_markdown, doc)
        
        # Add statistics
        doc.add_heading('Test Plan Statistics', level=2)
        doc.add_paragraph(f"Total sections processed: {test_plan.total_sections}")
        doc.add_paragraph(f"Total requirements: {test_plan.total_requirements}")
        doc.add_paragraph(f"Total test procedures: {test_plan.total_test_procedures}")
        doc.add_paragraph(f"Processing status: {test_plan.processing_status}")
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        # Serialize and encode
        buf = io.BytesIO()
        doc.save(buf)
        return base64.b64encode(buf.getvalue()).decode("utf-8")
    
    def _convert_markdown_to_docx(self, markdown_content: str, doc: Document):
        """Convert markdown to Word document with proper formatting"""
        lines = markdown_content.split('\n')
        i = 0
        in_table = False
        table_lines = []
        in_code_block = False
        code_lines = []

        while i < len(lines):
            line = lines[i]
            l = line.strip()

            # Handle code blocks
            if l.startswith('```'):
                if in_code_block:
                    # End of code block - add as formatted text
                    if code_lines:
                        code_para = doc.add_paragraph('\n'.join(code_lines))
                        code_para.style = 'Normal'
                        for run in code_para.runs:
                            run.font.name = 'Courier New'
                            run.font.size = Pt(9)
                    code_lines = []
                    in_code_block = False
                else:
                    in_code_block = True
                i += 1
                continue

            if in_code_block:
                code_lines.append(line)
                i += 1
                continue

            # Handle tables
            if '|' in l and l.count('|') >= 2:
                if not in_table:
                    in_table = True
                    table_lines = []
                table_lines.append(l)
                i += 1
                continue
            elif in_table:
                # Process accumulated table
                self._add_markdown_table_to_doc(doc, table_lines)
                table_lines = []
                in_table = False
                # Don't increment i, process this line as normal

            # Skip empty lines
            if not l:
                i += 1
                continue

            # Handle headings
            if l.startswith('#'):
                heading_level = 0
                while heading_level < len(l) and l[heading_level] == '#':
                    heading_level += 1
                heading_text = l[heading_level:].strip()
                if heading_level == 1:
                    # Skip main title as already added
                    pass
                elif heading_level == 2:
                    doc.add_heading(heading_text, level=1)
                elif heading_level == 3:
                    doc.add_heading(heading_text, level=2)
                elif heading_level == 4:
                    doc.add_heading(heading_text, level=3)
                else:
                    doc.add_heading(heading_text, level=4)
                i += 1
                continue

            # Handle bold-only lines as headings
            if l.startswith("**") and l.endswith("**") and l.count("**") == 2:
                doc.add_heading(l.replace("*", "").strip(), level=3)
                i += 1
                continue

            # Handle bullet lists
            if l.startswith(("-", "*", "•")) and len(l) > 1 and l[1] == ' ':
                text = l[2:].strip()
                self._add_formatted_paragraph(doc, text, style='List Bullet')
                i += 1
                continue

            # Handle numbered lists
            if re.match(r'^\d+[\.\)]\s', l):
                text = re.sub(r'^\d+[\.\)]\s+', '', l)
                self._add_formatted_paragraph(doc, text, style='List Number')
                i += 1
                continue

            # Regular paragraph with inline formatting
            self._add_formatted_paragraph(doc, l)
            i += 1

        # Process any remaining table
        if in_table and table_lines:
            self._add_markdown_table_to_doc(doc, table_lines)

    def _add_formatted_paragraph(self, doc: Document, text: str, style='Normal'):
        """Add a paragraph with proper inline markdown formatting (bold, italic, links)"""
        p = doc.add_paragraph(style=style)

        # Parse inline markdown: **bold**, *italic*, [link](url)
        i = 0
        while i < len(text):
            # Check for bold **text**
            if text[i:i+2] == '**':
                end = text.find('**', i+2)
                if end != -1:
                    run = p.add_run(text[i+2:end])
                    run.bold = True
                    i = end + 2
                    continue

            # Check for italic *text* (but not part of **)
            if text[i] == '*' and (i == 0 or text[i-1] != '*') and (i+1 < len(text) and text[i+1] != '*'):
                end = text.find('*', i+1)
                if end != -1 and (end+1 >= len(text) or text[end+1] != '*'):
                    run = p.add_run(text[i+1:end])
                    run.italic = True
                    i = end + 1
                    continue

            # Check for links [text](url)
            if text[i] == '[':
                link_end = text.find('](', i)
                if link_end != -1:
                    url_end = text.find(')', link_end+2)
                    if url_end != -1:
                        link_text = text[i+1:link_end]
                        # Just add the link text, not the URL (Word doc limitation)
                        run = p.add_run(link_text)
                        run.font.color.rgb = RGBColor(0, 0, 255)
                        run.underline = True
                        i = url_end + 1
                        continue

            # Regular character
            p.add_run(text[i])
            i += 1

    def _add_markdown_table_to_doc(self, doc: Document, table_lines: list):
        """Convert markdown table to Word table"""
        if not table_lines:
            return

        # Filter out separator lines (---|---|---)
        data_lines = [line for line in table_lines if not re.match(r'^\s*\|[\s\-:|]+\|\s*$', line)]

        if not data_lines:
            return

        # Parse table rows
        rows = []
        for line in data_lines:
            # Split by | and clean up
            cells = [cell.strip() for cell in line.split('|')]
            # Remove empty first/last cells if present
            if cells and not cells[0]:
                cells = cells[1:]
            if cells and not cells[-1]:
                cells = cells[:-1]
            if cells:
                rows.append(cells)

        if not rows:
            return

        # Create Word table
        num_cols = len(rows[0])
        table = doc.add_table(rows=len(rows), cols=num_cols)
        table.style = 'Light Grid Accent 1'

        # Populate table
        for i, row_data in enumerate(rows):
            row = table.rows[i]
            for j, cell_text in enumerate(row_data):
                if j < len(row.cells):
                    cell = row.cells[j]
                    # Remove markdown formatting from cell text
                    clean_text = re.sub(r'\*\*(.+?)\*\*', r'\1', cell_text)  # Remove bold
                    clean_text = re.sub(r'\*(.+?)\*', r'\1', clean_text)  # Remove italic
                    cell.text = clean_text
                    # Make header row bold
                    if i == 0:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
    
    def save_to_chromadb(self, test_plan: FinalTestPlan, session_id: str, pipeline_id: Optional[str] = None) -> Dict[str, Any]:
        """Save generated test plan to ChromaDB (single collection) with idempotency per pipeline.

        If a generated_document_id is already recorded in pipeline meta, re-use it and do not create a duplicate.
        """
        try:
            # First, ensure the target collection exists
            target_collection = os.getenv("GENERATED_TESTPLAN_COLLECTION", "generated_test_plan")
            logger.info(f"Attempting to save test plan to collection: {target_collection}")
            self._ensure_collection_exists(target_collection)
            logger.info(f"Collection {target_collection} verified/created")

            # Determine pipeline id
            pid = pipeline_id or getattr(test_plan, 'pipeline_id', None)
            logger.info(f"Saving test plan with pipeline_id: {pid}, session_id: {session_id}")

            # If a document was already saved for this pipeline, return that
            if pid:
                try:
                    meta = self.redis_client.hgetall(f"pipeline:{pid}:meta") or {}
                    existing_id = meta.get("generated_document_id")
                    if existing_id:
                        logger.info(f"Test plan already saved for pipeline {pid}, reusing document_id: {existing_id}")
                        return {
                            "document_id": existing_id,
                            "collection_name": meta.get("collection", target_collection),
                            "saved": True,
                            "generated_at": meta.get("completed_at") or datetime.now().isoformat(),
                            "architecture": "multi_agent_gpt4",
                            "reused": True
                        }
                except Exception as e:
                    logger.warning(f"Error checking for existing document: {e}")

                # Acquire a short-lived lock to avoid concurrent double-save
                try:
                    lock_key = f"pipeline:{pid}:saving_lock"
                    if not self.redis_client.set(lock_key, "1", nx=True, ex=60):
                        # Someone else is saving; wait briefly then return recorded id if present
                        time.sleep(1)
                        meta = self.redis_client.hgetall(f"pipeline:{pid}:meta") or {}
                        existing_id = meta.get("generated_document_id")
                        if existing_id:
                            return {
                                "document_id": existing_id,
                                "collection_name": meta.get("collection", target_collection),
                                "saved": True,
                                "generated_at": meta.get("completed_at") or datetime.now().isoformat(),
                                "architecture": "multi_agent_gpt4",
                                "reused": True
                            }
                except Exception:
                    pass

            # Create document content
            doc_content = f"Title: {test_plan.title}\n\n"
            doc_content += f"Architecture: Multi-Agent GPT-4 Pipeline\n"
            doc_content += f"Sections: {test_plan.total_sections}\n"
            doc_content += f"Requirements: {test_plan.total_requirements}\n" 
            doc_content += f"Test Procedures: {test_plan.total_test_procedures}\n\n"
            doc_content += test_plan.consolidated_markdown
            
            # Generate unique document ID first (needed for metadata)
            doc_id = f"testplan_multiagent_{(pid or session_id)}_{uuid.uuid4().hex[:8]}"
            logger.info(f"Generated document ID: {doc_id}")

            # Prepare metadata (include document_id for UI compatibility)
            metadata = {
                "title": test_plan.title,
                "type": "generated_test_plan",
                "architecture": "multi_agent_gpt4",
                "session_id": session_id,
                "generated_at": datetime.now().isoformat(),
                "total_sections": test_plan.total_sections,
                "total_requirements": test_plan.total_requirements,
                "total_test_procedures": test_plan.total_test_procedures,
                "processing_status": test_plan.processing_status,
                "agent_types": "3x_gpt4_actors_1x_gpt4_critic_1x_final_critic",
                "word_count": len(doc_content.split()),
                "char_count": len(doc_content),
                # UI compatibility fields
                "document_id": doc_id,
                "document_name": test_plan.title
            }

            # Save to ChromaDB
            payload = {
                "collection_name": target_collection,
                "documents": [doc_content],
                "metadatas": [metadata],
                "ids": [doc_id]
            }

            logger.info(f"Sending save request to ChromaDB: {self.fastapi_url}/api/vectordb/documents/add")
            logger.info(f"Payload metadata: {metadata}")

            response = requests.post(
                f"{self.fastapi_url}/api/vectordb/documents/add",
                json=payload,
                timeout=30
            )

            logger.info(f"ChromaDB response status: {response.status_code}")

            if response.ok:
                logger.info(f"Saved multi-agent test plan to ChromaDB ({target_collection}): {doc_id}")
                # Record generated doc in pipeline meta if pipeline_id provided
                try:
                    if pid:
                        self._update_pipeline_metadata(pid, {
                            "generated_document_id": doc_id,
                            "collection": target_collection,
                        })
                        self.redis_client.zadd("pipeline:recent", {pid: time.time()})
                except Exception as e:
                    logger.warning(f"Failed to record generated doc in pipeline meta: {e}")

                return {
                    "document_id": doc_id,
                    "collection_name": target_collection,
                    "saved": True,
                    "generated_at": metadata["generated_at"],
                    "architecture": "multi_agent_gpt4"
                }
            else:
                error_msg = f"Failed to save to ChromaDB: {response.status_code} - {response.text}"
                logger.error(error_msg)
                logger.error(f"Failed payload collection: {target_collection}, doc_id: {doc_id}")
                return {"saved": False, "error": response.text, "status_code": response.status_code}

        except requests.exceptions.RequestException as e:
            error_msg = f"Network error saving to ChromaDB: {e}"
            logger.error(error_msg)
            import traceback
            logger.error(f"Traceback: {traceback.format_exc()}")
            return {"saved": False, "error": str(e), "error_type": "network"}
        except Exception as e:
            error_msg = f"Unexpected error saving to ChromaDB: {e}"
            logger.error(error_msg)
            import traceback
            logger.error(f"Traceback: {traceback.format_exc()}")
            return {"saved": False, "error": str(e), "error_type": "unexpected"}
    
    def _ensure_generated_documents_collection_exists(self):
        """Deprecated no-op: use _ensure_collection_exists with GENERATED_TESTPLAN_COLLECTION instead."""
        try:
            self._ensure_collection_exists(os.getenv("GENERATED_TESTPLAN_COLLECTION", "generated_test_plan"))
        except Exception:
            pass

    def _ensure_collection_exists(self, name: str):
        """Ensure an arbitrary collection exists."""
        try:
            list_response = requests.get(f"{self.fastapi_url}/api/vectordb/collections", timeout=10)
            if list_response.ok:
                collections = list_response.json().get("collections", [])
                if name not in collections:
                    logger.info(f"Collection '{name}' does not exist. Creating it...")
                    create_response = requests.post(
                        f"{self.fastapi_url}/api/vectordb/collection/create",
                        params={"collection_name": name},
                        timeout=10
                    )
                    if create_response.ok:
                        logger.info(f"Successfully created collection '{name}'")
                    else:
                        logger.error(f"Failed to create collection '{name}': {create_response.status_code} - {create_response.text}")
                        raise Exception(f"Collection creation failed: {create_response.text}")
                else:
                    logger.debug(f"Collection '{name}' already exists")
            else:
                logger.error(f"Failed to list collections: {list_response.status_code} - {list_response.text}")
                raise Exception(f"Failed to verify collections: {list_response.text}")
        except Exception as e:
            logger.error(f"Unable to ensure collection {name}: {e}")
            raise
