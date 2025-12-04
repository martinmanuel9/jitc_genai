"""
Test Card Generation Service
Converts test procedures into executable test cards with pass/fail tracking.
Enhanced with notebook features: explicit prompts and parallel processing.
"""

from typing import List, Dict, Any, Optional
from dataclasses import dataclass, asdict
from concurrent.futures import ThreadPoolExecutor, as_completed
import re
import logging
import json

logger = logging.getLogger(__name__)

# Parallel processing configuration (from notebook)
MAX_WORKERS = 8  # Maximum concurrent test card generations

@dataclass
class TestCard:
    """Represents an executable test card"""
    test_id: str
    test_title: str
    procedures: List[str]  # Step-by-step procedures
    dependencies: List[str]
    expected_results: str
    acceptance_criteria: str
    section_title: str
    # Execution tracking (populated by users)
    executed: bool = False
    passed: bool = False
    failed: bool = False
    notes: str = ""

    def to_dict(self):
        """Convert to dictionary for JSON serialization"""
        return asdict(self)


class TestCardService:
    """
    Service for generating test cards from test plan sections.
    Based on mil_test_plan_gen.ipynb's build_test_card_for_section.
    """

    def __init__(self, llm_service):
        self.llm_service = llm_service

    def generate_test_card_from_rules(
        self,
        section_title: str,
        rules_markdown: str,
        format: str = "markdown_table"
    ) -> str:
        """
        Generate test card from test rules markdown.
        Enhanced with notebook-style explicit prompts.

        Args:
            section_title: Title of the section
            rules_markdown: Markdown containing test rules
            format: Output format ("markdown_table", "json", "docx_table")

        Returns:
            Test card in requested format
        """
        try:
            logger.info(f"Generating test card for section: {section_title} (format: {format})")

            # Extract test procedures from markdown
            test_procedures = self._extract_test_procedures_from_markdown(rules_markdown)

            if not test_procedures:
                logger.warning(f"No test procedures found in section: {section_title}")
                return self._create_empty_test_card(section_title, format)

            # Generate test card using LLM with format-specific prompts
            test_card_content = self._generate_test_card_with_llm(
                section_title,
                rules_markdown,
                test_procedures,
                format  # Pass format to LLM method
            )

            # Format output
            if format == "markdown_table":
                # For markdown, LLM returns table directly
                if isinstance(test_card_content, str):
                    return test_card_content
                return self._format_as_markdown_table(test_card_content)
            elif format == "json":
                if isinstance(test_card_content, str):
                    return test_card_content  # Already JSON string
                return self._format_as_json(test_card_content)
            elif format == "docx_table":
                return self._format_as_docx_table(test_card_content)
            else:
                raise ValueError(f"Unsupported format: {format}")

        except Exception as e:
            logger.error(f"Test card generation failed for {section_title}: {e}")
            # Return empty test card on error
            return self._create_empty_test_card(section_title, format)

    def _generate_test_card_with_llm(
        self,
        section_title: str,
        rules_markdown: str,
        test_procedures: List[Dict[str, Any]],
        format: str = "json"
    ):
        """
        Use LLM to transform test rules into structured test cards.
        Enhanced with notebook-style explicit prompts for better output quality.

        Args:
            section_title: Section title
            rules_markdown: Test rules in markdown
            test_procedures: Extracted test procedures
            format: Output format (markdown_table or json)

        Returns:
            For markdown_table: str (raw markdown table)
            For json: List[TestCard] or str (JSON string)
        """
        try:
            # Choose prompt based on format
            if format == "markdown_table":
                prompt = self._create_markdown_table_prompt(section_title, rules_markdown)
            else:  # json or docx_table
                prompt = self._create_json_prompt(section_title, rules_markdown)

            logger.info(f"Using {format} prompt for test card generation")

            response = self.llm_service.query_direct(
                model_name="gpt-4",
                query=prompt
            )[0]

            # Return based on format
            if format == "markdown_table":
                # Return raw table (notebook-style)
                logger.info(f"Generated markdown table test card")
                return response  # Raw markdown table string

            else:  # JSON format
                # Parse JSON response
                test_cards_data = self._parse_json_response(response)

                if not test_cards_data:
                    # Fallback to markdown extraction
                    logger.warning("LLM JSON parsing failed, using fallback extraction")
                    return self._extract_test_cards_from_markdown(section_title, rules_markdown)

                # Convert to TestCard objects
                test_cards = []
                for tc_data in test_cards_data:
                    test_cards.append(TestCard(
                        test_id=tc_data.get("test_id", f"TC-{len(test_cards)+1:03d}"),
                        test_title=tc_data.get("test_title", "Untitled Test"),
                        procedures=tc_data.get("procedures", []),
                        dependencies=tc_data.get("dependencies", []),
                        expected_results=tc_data.get("expected_results", ""),
                        acceptance_criteria=tc_data.get("acceptance_criteria", ""),
                        section_title=tc_data.get("section_title", section_title)
                    ))

                logger.info(f"Generated {len(test_cards)} test cards using LLM")
                return test_cards

        except Exception as e:
            logger.error(f"LLM test card generation failed: {e}")
            # Fallback: extract from markdown
            if format == "markdown_table":
                # Return empty table
                return self._create_empty_test_card(section_title, format)
            else:
                return self._extract_test_cards_from_markdown(section_title, rules_markdown)

    def _create_markdown_table_prompt(self, section_title: str, rules_markdown: str) -> str:
        """
        Create notebook-style explicit markdown table prompt.
        This produces cleaner, better-formatted tables.
        """
        # Filter out Table of Contents and non-testable sections
        filtered_markdown = self._filter_testable_content(rules_markdown)

        return f"""You are a QA test documentation assistant specializing in military/technical standards testing.

From the following test procedures (Markdown), generate a single Markdown pipe table named 'Test Card'
that lists one row per test procedure. Do NOT include any text before or after the table.

CRITICAL REQUIREMENTS:
- PRESERVE ORIGINAL REQUIREMENT IDs from the source (e.g., 4.2.1, REQ-01, etc.)
- Each test card row represents ONE requirement's test procedure
- Derive tests from the 'Test Procedures' sections (ignore TOC, introductions, or non-testable content)
- Group related test steps for the same requirement into ONE test card row

Columns: Req ID | Test Title | Procedures | Expected Results | Acceptance Criteria | Dependencies | Executed | Pass | Fail | Notes

Format Requirements:
- 'Req ID': Use ORIGINAL requirement ID from source document (e.g., 4.2.1, REQ-01)
- 'Test Title': Brief description of what requirement is being tested
- 'Procedures': Concise numbered steps separated by <br> (e.g., '1) Setup equipment<br>2) Configure parameters<br>3) Execute test')
- 'Expected Results': Specific measurable outcomes with values/ranges
- 'Acceptance Criteria': Explicit pass/fail thresholds with numbers
- 'Dependencies': Prerequisites, tools, equipment (leave empty if none)
- Leave 'Executed', 'Pass', and 'Fail' empty with checkbox (). Do NOT tick anything.
- 'Notes': Leave empty for execution notes

Output ONLY the table in GitHub-style pipe-table format.

=== SECTION NAME ===
{section_title}

=== TEST PROCEDURES (MARKDOWN) ===
{filtered_markdown}

=== END ==="""

    def _create_json_prompt(self, section_title: str, rules_markdown: str) -> str:
        """
        Create enhanced JSON prompt with explicit requirements.
        More detailed than original to match notebook quality.
        """
        # Filter out Table of Contents and non-testable sections
        filtered_markdown = self._filter_testable_content(rules_markdown)

        return f"""You are a QA test documentation assistant specializing in military/technical standards testing.

From the following test procedures (Markdown), generate test cards that group all requirements for comprehensive testing.

CRITICAL REQUIREMENTS:
- PRESERVE ORIGINAL REQUIREMENT IDs from source (e.g., 4.2.1, REQ-01, etc.)
- Each test card represents ONE requirement from the standard
- Group all related test steps for a requirement into ONE test card
- Derive tests from 'Test Procedures' sections (ignore TOC, introductions, non-testable content)

Requirements:
- Each test must be independently executable by an engineer
- 'test_id': Use ORIGINAL requirement ID from source document (e.g., "4.2.1", "REQ-01")
- 'test_title': Brief description of what requirement is being tested
- 'procedures': Detailed numbered steps with explicit actions (e.g., '1) Power off equipment and disconnect from power source', '2) Connect calibrated multimeter to input terminals')
- 'expected_results': Specific measurable outcomes with values and ranges
- 'acceptance_criteria': Explicit pass/fail thresholds with specific numbers
- 'dependencies': Prerequisites, tools, equipment from Dependencies section
- 'section_title': Keep as provided

Section Name: {section_title}

Test Procedures (Markdown):
{filtered_markdown}

Output format: Return ONLY a JSON array of test card objects with this exact structure:
[
  {{
    "test_id": "4.2.1",
    "test_title": "Brief descriptive title of requirement",
    "procedures": ["Step 1 description", "Step 2 description", "Step 3 description"],
    "expected_results": "Specific expected outcome with values and ranges",
    "acceptance_criteria": "Measurable pass/fail threshold with specific numbers",
    "dependencies": ["Specific tool or prerequisite"],
    "section_title": "{section_title}"
  }}
]

IMPORTANT:
- Return ONLY the JSON array, no additional text or formatting
- Use original requirement IDs from the source document
- One test card per requirement (group all test steps for that requirement)"""

    def _parse_json_response(self, response: str) -> Optional[List[Dict[str, Any]]]:
        """Parse JSON from LLM response, handling various formats"""
        try:
            # Try direct JSON parsing first
            return json.loads(response)
        except json.JSONDecodeError:
            pass

        # Try to extract JSON from markdown code blocks
        json_patterns = [
            r'```(?:json)?\s*(\[.*?\])\s*```',  # ```json [...] ```
            r'```(?:json)?\s*(\{.*?\})\s*```',  # ```json {...} ```
            r'\[.*?\]',  # Raw array
        ]

        for pattern in json_patterns:
            match = re.search(pattern, response, re.DOTALL)
            if match:
                try:
                    json_str = match.group(1) if match.lastindex else match.group(0)
                    data = json.loads(json_str)
                    if isinstance(data, list):
                        return data
                    elif isinstance(data, dict):
                        return [data]
                except json.JSONDecodeError:
                    continue

        return None

    def _format_as_markdown_table(self, test_cards: List[TestCard]) -> str:
        """
        Format test cards as readable text sections for Pandoc export.
        Uses structured text format instead of markdown tables for better readability in Word.
        """
        if not test_cards:
            return "*No testable procedures found in this section.*\n\n"

        result = ""
        for tc in test_cards:
            result += f"#### {tc.test_id}: {tc.test_title}\n\n"

            # Test Procedures
            result += "**Test Procedures:**\n\n"
            for i, proc in enumerate(tc.procedures, 1):
                result += f"{i}. {proc}\n"
            result += "\n"

            # Expected Results
            result += f"**Expected Results:** {tc.expected_results}\n\n"

            # Acceptance Criteria
            result += f"**Acceptance Criteria:** {tc.acceptance_criteria}\n\n"

            # Dependencies
            if tc.dependencies:
                dependencies_str = ", ".join(tc.dependencies)
                result += f"**Dependencies:** {dependencies_str}\n\n"

            # Execution tracking
            result += "**Execution Status:** [ ] Not Executed | [ ] Pass | [ ] Fail\n\n"
            result += "**Notes:** _____________________________________________\n\n"
            result += "---\n\n"

        return result

    def _format_as_json(self, test_cards: List[TestCard]) -> str:
        """Format test cards as JSON"""
        return json.dumps([tc.to_dict() for tc in test_cards], indent=2)

    def _format_as_docx_table(self, test_cards: List[TestCard]) -> Any:
        """
        Format test cards as python-docx Table object.
        Returns table object that can be added to a Document.
        """
        from docx import Document
        from docx.shared import Inches, Pt

        doc = Document()
        table = doc.add_table(rows=1, cols=9)
        table.style = 'Light Grid Accent 1'

        # Header row
        header_cells = table.rows[0].cells
        headers = ['Test ID', 'Test Title', 'Procedures', 'Expected Results',
                   'Acceptance Criteria', 'Executed', 'Pass', 'Fail', 'Notes']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].font.bold = True

        # Data rows
        for tc in test_cards:
            row_cells = table.add_row().cells
            row_cells[0].text = tc.test_id
            row_cells[1].text = tc.test_title
            row_cells[2].text = '\n'.join([f"{i+1}) {p}" for i, p in enumerate(tc.procedures)])
            row_cells[3].text = tc.expected_results
            row_cells[4].text = tc.acceptance_criteria
            row_cells[5].text = ''
            row_cells[6].text = ''
            row_cells[7].text = ''
            row_cells[8].text = ''

        return table

    def _filter_testable_content(self, markdown: str) -> str:
        """
        Filter out non-testable sections from markdown.
        Removes: Table of Contents, Introduction, Executive Summary, etc.
        Keeps: Test Procedures, Test Rules, Requirements sections
        """
        import re

        # Patterns for non-testable sections
        skip_patterns = [
            r'##\s*Table of Contents',
            r'##\s*Introduction',
            r'##\s*Executive Summary',
            r'##\s*Background',
            r'##\s*Overview',
            r'##\s*Purpose',
            r'##\s*Scope',
        ]

        lines = markdown.split('\n')
        filtered_lines = []
        skip_section = False

        for line in lines:
            # Check if we're entering a non-testable section
            if any(re.match(pattern, line, re.IGNORECASE) for pattern in skip_patterns):
                skip_section = True
                continue

            # Check if we're entering a testable section (end skip)
            if line.startswith('##') and not skip_section:
                # New section - check if testable
                if any(keyword in line.lower() for keyword in ['test', 'procedure', 'requirement', 'verification']):
                    skip_section = False
            elif line.startswith('##'):
                skip_section = False  # Any new section ends skip

            if not skip_section:
                filtered_lines.append(line)

        return '\n'.join(filtered_lines)

    def _extract_test_procedures_from_markdown(self, markdown: str) -> List[Dict[str, Any]]:
        """Extract test procedures from markdown (fallback method)"""
        procedures = []

        # Updated to look for Test Procedures sections (new format)
        in_test_procedures = False
        current_procedure = None

        for line in markdown.split('\n'):
            line_stripped = line.strip()

            # Look for Test Procedures section
            if line_stripped.startswith('**Test Procedures:**') or line_stripped.startswith('### Test Procedure'):
                in_test_procedures = True

                # Extract requirement ID from heading if present
                req_id_match = re.search(r'Test Procedure\s+([\d.]+|REQ-\d+)', line_stripped)
                if req_id_match:
                    current_procedure = {
                        "id": req_id_match.group(1),
                        "description": "",
                        "type": "functional"
                    }
                continue

            # Look for old-style Test Rules
            elif line_stripped.startswith('**Test Rules:**'):
                in_test_procedures = True
                continue
            elif line_stripped.startswith('**') and in_test_procedures:
                # Hit another section
                if current_procedure and current_procedure["description"]:
                    procedures.append(current_procedure)
                    current_procedure = None
                continue
            elif in_test_procedures and re.match(r'^\d+\.', line_stripped):
                # Numbered test step
                if current_procedure:
                    current_procedure["description"] += " " + line_stripped
                else:
                    procedures.append({
                        "id": f"test_{len(procedures)+1}",
                        "description": line_stripped,
                        "type": "functional"
                    })

        # Add last procedure if exists
        if current_procedure and current_procedure["description"]:
            procedures.append(current_procedure)

        return procedures

    def _extract_test_cards_from_markdown(
        self,
        section_title: str,
        markdown: str
    ) -> List[TestCard]:
        """Fallback: Extract test cards directly from markdown without LLM"""
        test_cards = []
        procedures = self._extract_test_procedures_from_markdown(markdown)

        for i, proc in enumerate(procedures, 1):
            # Try to extract acceptance criteria from the description
            description = proc["description"]

            test_cards.append(TestCard(
                test_id=f"TC-{i:03d}",
                test_title=f"Test Procedure {i}",
                procedures=[description],
                dependencies=[],
                expected_results="Test completes without errors",
                acceptance_criteria="All steps pass successfully",
                section_title=section_title
            ))

        if not test_cards:
            logger.warning(f"No test cards extracted from markdown for: {section_title}")

        return test_cards

    def _create_empty_test_card(self, section_title: str, format: str) -> str:
        """Create empty test card when no procedures found"""
        if format == "markdown_table":
            return f"*No testable procedures found in this section.*\n\n"
        elif format == "json":
            return "[]"
        else:
            return ""

    def generate_test_cards_for_pipeline(
        self,
        pipeline_id: str,
        redis_client,
        format: str = "markdown_table",
        max_workers: int = MAX_WORKERS
    ) -> Dict[str, str]:
        """
        Generate test cards for all sections in a pipeline with parallel processing.
        Enhanced with notebook-style parallel generation for 4-8x speedup.

        Args:
            pipeline_id: Redis pipeline ID
            redis_client: Redis client instance
            format: Output format
            max_workers: Maximum concurrent workers (default: 8)

        Returns:
            Dictionary mapping section titles to test card content
        """
        test_cards = {}

        try:
            # Get all section critic results from Redis
            pattern = f"pipeline:{pipeline_id}:critic:*"
            critic_keys = redis_client.keys(pattern)

            logger.info(f"Generating test cards for {len(critic_keys)} sections in pipeline {pipeline_id} (parallel: {max_workers} workers)")

            # Prepare tasks
            tasks = []
            for key in critic_keys:
                try:
                    critic_data = redis_client.hgetall(key)
                    section_title = critic_data.get("section_title", "")
                    synthesized_rules = critic_data.get("synthesized_rules", "")

                    if section_title and synthesized_rules:
                        tasks.append((key, section_title, synthesized_rules))
                except Exception as e:
                    logger.error(f"Failed to read data for key {key}: {e}")
                    continue

            if not tasks:
                logger.warning("No valid tasks found for test card generation")
                return {}

            # Process in parallel
            with ThreadPoolExecutor(max_workers=max_workers) as executor:
                # Submit all tasks
                future_to_section = {
                    executor.submit(
                        self.generate_test_card_from_rules,
                        section_title,
                        synthesized_rules,
                        format
                    ): (key, section_title)
                    for key, section_title, synthesized_rules in tasks
                }

                # Collect results as they complete
                completed = 0
                total = len(future_to_section)
                for future in as_completed(future_to_section):
                    key, section_title = future_to_section[future]
                    completed += 1
                    try:
                        test_card = future.result()
                        test_cards[section_title] = test_card
                        logger.info(f" [{completed}/{total}] Generated test card for: {section_title}")
                    except Exception as e:
                        logger.error(f" [{completed}/{total}] Failed to generate test card for {section_title}: {e}")
                        continue

            logger.info(f"Successfully generated {len(test_cards)} test cards in parallel")
            return test_cards

        except Exception as e:
            logger.error(f"Pipeline test card generation failed: {e}")
            return {}

    # =========================================================================
    # NEW: Test Card as Separate Documents (ChromaDB Integration)
    # =========================================================================

    def generate_test_cards_from_test_plan(
        self,
        test_plan_id: str,
        test_plan_content: str,
        test_plan_title: str,
        format: str = "markdown_table"
    ) -> List[Dict[str, Any]]:
        """
        Parse a test plan and generate individual test card documents.
        Each test procedure becomes a separate test card document.

        Args:
            test_plan_id: ID of the test plan document
            test_plan_content: Full test plan markdown content
            test_plan_title: Title of the test plan
            format: Output format for test cards

        Returns:
            List of test card documents ready to save to ChromaDB
        """
        import uuid
        import os
        import requests
        from datetime import datetime

        try:
            logger.info(f"Generating individual test cards from test plan: {test_plan_id}")

            # Parse test plan into sections
            sections = self._parse_test_plan_into_sections(test_plan_content)
            logger.info(f"Parsed test plan into {len(sections)} sections")

            # Generate test cards for each section
            all_test_cards = []
            test_card_counter = 1

            for section in sections:
                section_title = section.get('title', 'Unknown Section')
                section_index = section.get('index', 0)
                section_content = section.get('content', '')

                # Extract individual test procedures from section
                test_procedures = self._extract_individual_tests(section_content)
                logger.info(f"Section '{section_title}': extracted {len(test_procedures)} individual tests")

                for test_proc in test_procedures:
                    # Generate unique test card ID
                    test_id = f"TC-{test_card_counter:03d}"
                    card_id = f"testcard_{test_plan_id}_{test_id}_{uuid.uuid4().hex[:8]}"

                    # Create test card document
                    test_card_doc = {
                        "document_id": card_id,
                        "document_name": f"{test_id} {test_proc.get('title', 'Test')}",
                        "content": self._format_test_card_content(test_proc, test_id, format),
                        "metadata": {
                            # Link to test plan
                            "test_plan_id": test_plan_id,
                            "test_plan_title": test_plan_title,
                            "section_title": section_title,
                            "section_index": section_index,

                            # Test identification
                            "test_id": test_id,
                            "test_title": test_proc.get('title', 'Untitled Test'),
                            "requirement_id": test_proc.get('requirement_id', ''),
                            "requirement_text": test_proc.get('requirement_text', ''),

                            # Test details
                            "procedures": json.dumps(test_proc.get('procedures', [])),
                            "expected_results": test_proc.get('expected_results', ''),
                            "acceptance_criteria": test_proc.get('acceptance_criteria', ''),
                            "dependencies": json.dumps(test_proc.get('dependencies', [])),

                            # Execution tracking
                            "execution_status": "not_executed",
                            "executed_by": "",
                            "executed_at": "",
                            "execution_duration_minutes": 0,
                            "passed": "false",
                            "failed": "false",
                            "actual_results": "",
                            "notes": "",

                            # Categorization
                            "test_type": test_proc.get('test_type', 'functional'),
                            "test_category": test_proc.get('category', ''),
                            "priority": test_proc.get('priority', 'medium'),
                            "estimated_duration_minutes": test_proc.get('estimated_duration', 30),

                            # Timestamps
                            "created_at": datetime.now().isoformat(),
                            "updated_at": datetime.now().isoformat(),
                            "type": "test_card",
                            "format": format
                        }
                    }

                    all_test_cards.append(test_card_doc)
                    test_card_counter += 1

            logger.info(f"Generated {len(all_test_cards)} individual test card documents")
            return all_test_cards

        except Exception as e:
            logger.error(f"Failed to generate test cards from test plan: {e}")
            raise

    def save_test_cards_to_chromadb(
        self,
        test_cards: List[Dict[str, Any]],
        collection_name: str = "test_cards"
    ) -> Dict[str, Any]:
        """
        Save individual test cards to ChromaDB.

        Args:
            test_cards: List of test card documents
            collection_name: ChromaDB collection name

        Returns:
            Result dictionary with saved test card IDs
        """
        import os
        import requests

        try:
            fastapi_url = os.getenv("FASTAPI_URL", "http://fastapi:9020")
            logger.info(f"Saving {len(test_cards)} test cards to ChromaDB collection: {collection_name}")

            # Ensure collection exists
            try:
                response = requests.get(f"{fastapi_url}/api/vectordb/collections")
                collections = response.json() if response.ok else []
                if collection_name not in collections:
                    logger.info(f"Creating collection: {collection_name}")
                    requests.post(
                        f"{fastapi_url}/api/vectordb/collection/create",
                        params={"collection_name": collection_name}
                    )
            except Exception as e:
                logger.warning(f"Could not verify/create collection: {e}")

            # Prepare batch for ChromaDB
            ids = []
            documents = []
            metadatas = []

            for card in test_cards:
                ids.append(card["document_id"])
                documents.append(card["content"])
                metadatas.append(card["metadata"])

            # Save to ChromaDB
            payload = {
                "collection_name": collection_name,
                "documents": documents,
                "metadatas": metadatas,
                "ids": ids
            }

            response = requests.post(
                f"{fastapi_url}/api/vectordb/documents/add",
                json=payload,
                timeout=60
            )

            if response.ok:
                logger.info(f"Successfully saved {len(test_cards)} test cards to ChromaDB")
                return {
                    "saved": True,
                    "count": len(test_cards),
                    "test_card_ids": ids,
                    "collection": collection_name
                }
            else:
                logger.error(f"Failed to save test cards: {response.status_code} - {response.text}")
                return {
                    "saved": False,
                    "error": response.text
                }

        except Exception as e:
            logger.error(f"Error saving test cards to ChromaDB: {e}")
            return {
                "saved": False,
                "error": str(e)
            }

    def _parse_test_plan_into_sections(self, content: str) -> List[Dict[str, Any]]:
        """
        Parse test plan markdown into sections (optimized for large documents).

        Args:
            content: Test plan markdown content

        Returns:
            List of sections with title, index, and content (max 50 sections)
        """
        sections = []
        lines = content.split('\n')
        current_section = None
        section_index = 0

        # Limit total sections to prevent excessive processing
        MAX_SECTIONS = 50

        for line in lines:
            # Check for section headers (## Section Title or # Section Title)
            if line.startswith('## ') or line.startswith('# '):
                if current_section:
                    sections.append(current_section)

                    # Stop if we've reached max sections
                    if len(sections) >= MAX_SECTIONS:
                        logger.warning(f"Reached maximum section limit ({MAX_SECTIONS}), stopping parse")
                        break

                section_index += 1
                current_section = {
                    'title': line.replace('##', '').replace('#', '').strip()[:200],  # Limit title length
                    'index': section_index,
                    'content': ''
                }
            elif current_section:
                current_section['content'] += line + '\n'

        # Add last section
        if current_section and len(sections) < MAX_SECTIONS:
            sections.append(current_section)

        logger.info(f"Parsed {len(sections)} sections from test plan")
        return sections

    def _extract_individual_tests(self, section_content: str) -> List[Dict[str, Any]]:
        """
        Extract individual test procedures from section content.
        Looks for **Test Rules:** section and extracts each numbered rule as a test.

        Args:
            section_content: Section markdown content

        Returns:
            List of individual test procedures (max 20 per section)
        """
        tests = []
        MAX_TESTS_PER_SECTION = 20

        # Limit content size
        if len(section_content) > 50000:
            logger.warning(f"Section content too large ({len(section_content)} chars), truncating to 50KB")
            section_content = section_content[:50000]

        # Find the **Test Rules:** section
        test_rules_match = re.search(r'\*\*Test Rules:\*\*\s*\n(.*?)(?=\n\*\*|\n##|\Z)', section_content, re.DOTALL | re.IGNORECASE)

        if not test_rules_match:
            # Fallback: look for any numbered list
            logger.info("No '**Test Rules:**' section found, looking for any numbered list")
            test_rules_content = section_content
        else:
            test_rules_content = test_rules_match.group(1)

        # Extract numbered test rules (1. , 2. , etc.)
        lines = test_rules_content.split('\n')
        current_rule_num = None
        current_rule_lines = []

        for line in lines:
            # Match numbered items: "1. ", "2. ", etc.
            match = re.match(r'^(\d+)\.\s+(.+)', line)
            if match:
                # Save previous rule
                if current_rule_num is not None and current_rule_lines:
                    rule_text = '\n'.join(current_rule_lines)
                    tests.append(self._create_test_from_rule(current_rule_num, rule_text))

                    if len(tests) >= MAX_TESTS_PER_SECTION:
                        logger.warning(f"Reached maximum test limit ({MAX_TESTS_PER_SECTION}) for section")
                        break

                # Start new rule
                current_rule_num = int(match.group(1))
                current_rule_lines = [line]
            elif current_rule_num is not None:
                # Continue current rule
                current_rule_lines.append(line)

        # Add last rule
        if current_rule_num is not None and current_rule_lines and len(tests) < MAX_TESTS_PER_SECTION:
            rule_text = '\n'.join(current_rule_lines)
            tests.append(self._create_test_from_rule(current_rule_num, rule_text))

        # If no tests found, create a single test from section
        if not tests:
            logger.warning("No numbered test rules found in section")
            tests.append({
                'title': 'Section Review',
                'procedures': ['Review section content', 'Verify compliance'],
                'expected_results': 'Section requirements are met',
                'acceptance_criteria': 'All requirements verified',
                'dependencies': [],
                'requirement_id': 'REQ-001',
                'requirement_text': section_content[:300] if section_content else 'Section content',
                'test_type': 'functional',
                'category': 'compliance',
                'priority': 'medium',
                'estimated_duration': 30
            })

        return tests

    def _create_test_from_rule(self, rule_num: int, rule_text: str) -> Dict[str, Any]:
        """
        Create a test card from a test rule.
        Parses the rule to extract title, procedures, expected results, etc.

        Args:
            rule_num: Rule number (e.g., 1, 2, 3)
            rule_text: Full rule text including numbered line

        Returns:
            Test card dict
        """
        # Extract title from first line (remove numbering)
        first_line = rule_text.split('\n')[0] if '\n' in rule_text else rule_text
        title = re.sub(r'^\d+\.\s+', '', first_line).strip()

        # Limit title length
        if len(title) > 100:
            title = title[:97] + "..."

        # Extract procedures (look for sub-numbered items like a), b), 1), 2), etc.)
        procedures = self._extract_procedures(rule_text)

        # Extract expected results (look for keywords)
        expected_results = self._extract_expected_results(rule_text)

        # Extract acceptance criteria (look for specific values/ranges)
        acceptance_criteria = self._extract_acceptance_criteria(rule_text)

        # Extract dependencies
        dependencies = self._extract_dependencies(rule_text)

        # Create requirement ID
        requirement_id = f"REQ-{rule_num:03d}"

        return {
            'title': title or f"Test Rule {rule_num}",
            'procedures': procedures,
            'expected_results': expected_results,
            'acceptance_criteria': acceptance_criteria,
            'dependencies': dependencies,
            'requirement_id': requirement_id,
            'requirement_text': rule_text.strip()[:500],  # First 500 chars of rule
            'test_type': 'functional',
            'category': 'compliance',
            'priority': 'high' if any(kw in rule_text.lower() for kw in ['critical', 'mandatory', 'must']) else 'medium',
            'estimated_duration': 30 + (len(procedures) * 10)  # Base 30 min + 10 min per procedure
        }

    def _extract_procedures(self, text: str) -> List[str]:
        """
        Extract step-by-step procedures from test text.
        Looks for sub-numbered items (a), b), 1), 2), etc.) or bullet points.
        """
        procedures = []

        # Pattern 1: Sub-numbered items like "a)", "b)", "1)", "2)"
        sub_pattern = re.compile(r'^\s*([a-z]\)|\d+\))\s+(.+?)$', re.MULTILINE | re.IGNORECASE)
        matches = sub_pattern.findall(text)
        if matches:
            procedures = [m[1].strip() for m in matches if m[1].strip()]

        # Pattern 2: Bullet points or dashes
        if not procedures:
            bullet_pattern = re.compile(r'^\s*[-•]\s+(.+?)$', re.MULTILINE)
            matches = bullet_pattern.findall(text)
            if matches:
                procedures = [m.strip() for m in matches if m.strip()]

        # Pattern 3: Lines with procedure keywords
        if not procedures:
            for line in text.split('\n'):
                line = line.strip()
                if any(kw in line.lower() for kw in ['verify', 'test', 'measure', 'check', 'ensure', 'configure', 'setup', 'execute']):
                    if len(line) > 10 and len(line) < 300:
                        procedures.append(line)

        # Fallback: Extract key sentences
        if not procedures:
            sentences = re.split(r'[.!]\s+', text)
            procedures = [s.strip() for s in sentences if 20 < len(s.strip()) < 300][:5]

        # Default if nothing found
        if not procedures:
            procedures = ["Execute test as per requirement specification"]

        return procedures[:10]  # Limit to 10 procedures

    def _extract_expected_results(self, text: str) -> str:
        """
        Extract expected results from test text.
        Looks for requirement keywords and specific outcomes.
        """
        # Look for explicit expected results sections
        expected_match = re.search(r'(?:expected|result|outcome|output):\s*(.+?)(?:\n|$)', text, re.IGNORECASE)
        if expected_match:
            return expected_match.group(1).strip()[:300]

        # Look for "shall", "must", "should" statements
        requirement_keywords = ['shall', 'must', 'should', 'will', 'required to']
        for keyword in requirement_keywords:
            pattern = re.compile(rf'\b{keyword}\b.+?[.!]', re.IGNORECASE)
            match = pattern.search(text)
            if match:
                result = match.group(0).strip()
                if len(result) < 500:
                    return result

        # Look for statements with "meets", "complies", "within", "between"
        outcome_keywords = ['meets', 'complies', 'within', 'between', 'exceeds', 'achieves']
        for keyword in outcome_keywords:
            if keyword in text.lower():
                sentences = re.split(r'[.!?]\s+', text)
                for sent in sentences:
                    if keyword in sent.lower() and len(sent.strip()) > 20:
                        return sent.strip()[:300]

        # Look for sentences with specific numeric values
        numeric_pattern = re.compile(r'[^.!?]*\b\d+(?:\.\d+)?\s*(?:±|%|dB|Hz|V|A|W|Ω|°C|°F|mm|cm|m|kg|g|lb)\b[^.!?]*[.!?]')
        match = numeric_pattern.search(text)
        if match:
            return match.group(0).strip()[:300]

        return "System meets all specified requirements and performance criteria"

    def _extract_acceptance_criteria(self, text: str) -> str:
        """
        Extract acceptance criteria from test text.
        Looks for pass/fail thresholds, numeric ranges, and measurable criteria.
        """
        # Look for explicit acceptance/pass/fail criteria
        criteria_match = re.search(r'(?:acceptance|pass|fail|criteria):\s*(.+?)(?:\n\n|\Z)', text, re.IGNORECASE | re.DOTALL)
        if criteria_match:
            criteria = criteria_match.group(1).strip()
            if len(criteria) < 500:
                return criteria

        # Look for numeric ranges and thresholds
        # Pattern: "± X", "< X", "> X", "X to Y", "between X and Y", "within X%"
        range_patterns = [
            r'(?:±|plus or minus)\s*[\d.]+\s*[%°CFHzdBVAWΩmkglb]*',
            r'(?:less than|<|≤)\s*[\d.]+\s*[%°CFHzdBVAWΩmkglb]+',
            r'(?:greater than|>|≥)\s*[\d.]+\s*[%°CFHzdBVAWΩmkglb]+',
            r'[\d.]+\s*(?:to|through|-)\s*[\d.]+\s*[%°CFHzdBVAWΩmkglb]+',
            r'between\s+[\d.]+\s+and\s+[\d.]+\s*[%°CFHzdBVAWΩmkglb]*',
            r'within\s+[\d.]+\s*[%°CFHzdBVAWΩmkglb]+',
            r'(?:minimum|maximum|min|max)\s*[:=]?\s*[\d.]+\s*[%°CFHzdBVAWΩmkglb]+'
        ]

        found_criteria = []
        for pattern in range_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            found_criteria.extend(matches[:2])  # Limit matches per pattern

        if found_criteria:
            return "Pass if: " + "; ".join(found_criteria[:5])

        # Look for "shall not exceed", "must be", "should be" statements
        constraint_keywords = ['shall not exceed', 'must not exceed', 'must be', 'should be', 'required to be']
        for keyword in constraint_keywords:
            pattern = re.compile(rf'\b{keyword}\b.+?[.!]', re.IGNORECASE)
            match = pattern.search(text)
            if match:
                return match.group(0).strip()[:300]

        # Look for percentages, ranges, or specific values
        value_pattern = re.compile(r'\b\d+(?:\.\d+)?\s*(?:%|±|dB|Hz|V|A|W|Ω|°C|°F|mm|cm|m|kg|g|lb)\b')
        if value_pattern.search(text):
            return "Measured values must meet specification limits (see requirement details)"

        return "All test procedures complete successfully with results within specified tolerance"

    def _extract_dependencies(self, text: str) -> List[str]:
        """Extract test dependencies from text"""
        dependencies = []
        # Look for equipment mentions
        equipment_keywords = ['calibrated', 'meter', 'analyzer', 'scope', 'generator', 'equipment', 'chamber']
        for keyword in equipment_keywords:
            if keyword in text.lower():
                # Extract surrounding context
                idx = text.lower().find(keyword)
                context = text[max(0, idx-20):min(len(text), idx+50)]
                dependencies.append(context.strip())
        return dependencies[:3]

    def _format_test_card_content(self, test_proc: Dict[str, Any], test_id: str, format: str) -> str:
        """Format test card content based on format type"""
        if format == "markdown_table":
            procedures_text = "<br>".join([f"{i+1}) {p}" for i, p in enumerate(test_proc.get('procedures', []))])
            deps_text = "<br>".join(test_proc.get('dependencies', [])) if test_proc.get('dependencies') else "None"

            return f"""| Test ID | Test Title | Procedures | Expected Results | Acceptance Criteria | Dependencies | Executed | Pass | Fail | Notes |
|---------|------------|------------|------------------|---------------------|--------------|----------|------|------|-------|
| {test_id} | {test_proc.get('title', 'Test')} | {procedures_text} | {test_proc.get('expected_results', '')} | {test_proc.get('acceptance_criteria', '')} | {deps_text} | () | () | () | |
"""
        elif format == "json":
            return json.dumps(test_proc, indent=2)
        else:
            return str(test_proc)
